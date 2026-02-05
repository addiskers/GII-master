from flask import Flask, request, jsonify, send_file, render_template, redirect, url_for, session, flash, g
import io
import os
import glob
import json
import sqlite3
from datetime import datetime, timezone, timedelta
from docx import Document
from scrapy.crawler import CrawlerProcess
from scrapy.utils.project import get_project_settings
from multiprocessing import Process
from multi_scraper.spiders.data import MarketResearchSpider, generate_docx_from_data
from multi_scraper.toc import export_to_word, add_bullet_point_text, transform_market_data, generate_segmental_analysis, title_h1
from openai import OpenAI
from dotenv import load_dotenv
import re
import tempfile
import secrets

app = Flask(__name__, template_folder='templates', static_folder='templates', static_url_path='/templates')

# Load environment variables first
load_dotenv()

# Use strong random secret key if not set in environment
app.secret_key = os.getenv('FLASK_SECRET') or secrets.token_hex(32)
if not os.getenv('FLASK_SECRET'):
    app.logger.warning('FLASK_SECRET not set! Using random key. Sessions will be invalidated on restart. Set FLASK_SECRET in .env for production.')

# IST Timezone (UTC+5:30)
IST = timezone(timedelta(hours=5, minutes=30))

def now_ist():
    """Return current time in IST"""
    return datetime.now(IST).replace(tzinfo=None).isoformat()

# Initialize OpenAI client (new style – picks up API key from env)
openai_client = OpenAI()

# Ensure the scraped_json directory exists
os.makedirs('scraped_json', exist_ok=True)

# Helper function to flatten nested AI segments into a list of tuples
def flatten_ai_segments(segments_data):
    """
    Convert AI-generated nested segment structure to flat list with levels.
    
    Input format:
    {
      "segments": [
        {
          "name": "Product Type",
          "subsegments": [
            {"name": "Convoluted Air Springs", "subsegments": []},
            {"name": "Rolling Lobe Air Springs", "subsegments": []}
          ]
        }
      ]
    }
    
    Output format: [(name, level), ...]
    Example: [("Product Type", 0), ("Convoluted Air Springs", 1), ...]
    """
    result = []
    
    def traverse(segments, level=0):
        for segment in segments:
            if isinstance(segment, dict):
                result.append((segment.get('name', ''), level))
                if segment.get('subsegments'):
                    traverse(segment['subsegments'], level + 1)
    
    if isinstance(segments_data, dict) and 'segments' in segments_data:
        traverse(segments_data['segments'])
    elif isinstance(segments_data, list):
        traverse(segments_data)
    
    return result

# Jinja filter: format ISO timestamp to "dd mon yyyy HH:MM" (e.g., 26 nov 2025 08:31)
@app.template_filter('fmt_ts')
def fmt_ts(value):
    try:
        if not value:
            return '—'
        s = str(value).strip()
        # Normalize common variants
        s = s.rstrip('Z')
        # Try ISO parse
        try:
            dt = datetime.fromisoformat(s)
        except Exception:
            # Fallback: match "YYYY-MM-DD[ T]HH:MM"
            m = re.match(r"^(\d{4})-(\d{2})-(\d{2})[ T](\d{2}):(\d{2})", s)
            if not m:
                return s
            y, mo, d, hh, mm = m.groups()
            dt = datetime(int(y), int(mo), int(d), int(hh), int(mm))
        out = dt.strftime('%Y-%m-%d %I:%M %p')
        return out
    except Exception:
        return '—'

# ==== Auth & User Management (merged from auth_app.py) ====
DB_PATH = os.path.join(os.path.dirname(__file__), 'auth_users.db')

def get_db():
    db = getattr(g, '_database', None)
    if db is None:
        db = g._database = sqlite3.connect(DB_PATH)
        db.row_factory = sqlite3.Row
    return db

def init_db():
    db = get_db()
    
    # Drop old tables if they exist (migration - one time only)
    try:
        db.execute('DROP TABLE IF EXISTS submissions')
        db.execute('DROP TABLE IF EXISTS chat_messages')
        # Note: rd_submissions is NOT dropped - data persists
    except Exception:
        pass
    
    # Users table - only admin and researcher roles
    db.execute('''
    CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT UNIQUE NOT NULL,
        password_hash TEXT NOT NULL,
        role TEXT NOT NULL CHECK(role IN ('admin', 'researcher')),
        created_by TEXT,
        created_at TEXT,
        last_login TEXT
    )
    ''')
    
    # RD Submissions table - tracks all Full RD saves with complete data
    db.execute('''
    CREATE TABLE IF NOT EXISTS rd_submissions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        market_name TEXT NOT NULL,
        researcher_username TEXT NOT NULL,
        json_path TEXT NOT NULL,
        file_path TEXT,
        submitted_at TEXT,
        timestamp TEXT,
        
        -- Industry Classification
        sector TEXT,
        industry_group TEXT,
        industry TEXT,
        sub_industry TEXT,
        
        -- Market Inputs
        value_unit TEXT,
        cagr REAL,
        market_size_2024 REAL,
        market_size_2025 REAL,
        projected_size_2033 REAL,
        
        -- Data arrays stored as JSON
        segments TEXT,
        ai_gen_seg TEXT,
        companies TEXT,
        
        created_by TEXT,
        version TEXT,
        
        -- Download tracking
        downloaded INTEGER DEFAULT 0,
        last_downloaded_at TEXT
    )
    ''')
    
    # Clean up old creator users (migration)
    try:
        db.execute("DELETE FROM users WHERE role = 'creator'")
    except Exception:
        pass
    
    db.commit()

from werkzeug.security import generate_password_hash, check_password_hash

def create_user(username, password, role, created_by=None):
    db = get_db()
    # Use pbkdf2:sha256 to avoid environments without hashlib.scrypt
    password_hash = generate_password_hash(password, method='pbkdf2:sha256', salt_length=16)
    now = now_ist()
    try:
        db.execute('INSERT INTO users (username, password_hash, role, created_by, created_at) VALUES (?,?,?,?,?)',
                   (username, password_hash, role, created_by, now))
        db.commit()
        return True
    except sqlite3.IntegrityError:
        return False

def get_user_by_username(username):
    db = get_db()
    cur = db.execute('SELECT * FROM users WHERE username = ?', (username,))
    return cur.fetchone()

def update_last_login(username):
    db = get_db()
    now = now_ist()
    db.execute('UPDATE users SET last_login = ? WHERE username = ?', (now, username))
    db.commit()

def list_users():
    db = get_db()
    cur = db.execute('SELECT username, role, created_by, created_at, last_login FROM users ORDER BY role, username')
    return cur.fetchall()

@app.before_request
def before_request():
    init_db()

@app.teardown_appcontext
def close_connection(exception):
    db = getattr(g, '_database', None)
    if db is not None:
        db.close()

# Create default admin if none exists
with app.app_context():
    init_db()
    db = get_db()
    cur = db.execute('SELECT COUNT(*) as cnt FROM users WHERE role = ?', ('admin',))
    r = cur.fetchone()
    # Use strong default password from env, or generate a random one
    admin_pass = os.getenv('DEFAULT_ADMIN_PASSWORD')
    if not admin_pass:
        admin_pass = secrets.token_urlsafe(16)
        app.logger.warning(f'DEFAULT_ADMIN_PASSWORD not set! Generated random password: {admin_pass}')
        app.logger.warning('Set DEFAULT_ADMIN_PASSWORD in .env for production.')
    if not r or r['cnt'] == 0:
        # Create default admin using pbkdf2
        create_user('admin', admin_pass, 'admin', created_by='system')
        app.logger.info('Default admin created with username `admin`. Change the password immediately.')
    else:
        # If existing admin uses scrypt, migrate to pbkdf2 using DEFAULT_ADMIN_PASSWORD
        cur = db.execute('SELECT username, password_hash FROM users WHERE role = ? LIMIT 1', ('admin',))
        admin_row = cur.fetchone()
        if admin_row and isinstance(admin_row['password_hash'], str) and admin_row['password_hash'].startswith('scrypt:'):
            new_hash = generate_password_hash(admin_pass, method='pbkdf2:sha256', salt_length=16)
            db.execute('UPDATE users SET password_hash = ? WHERE username = ?', (new_hash, admin_row['username']))
            db.commit()
            app.logger.warning('Admin password hash migrated from scrypt to pbkdf2. Use DEFAULT_ADMIN_PASSWORD to log in.')

def login_required(role=None):
    """Require login, optionally restricting by role or roles.

    role can be a string (single role) or an iterable of roles.
    """
    def decorator(func):
        def wrapper(*args, **kwargs):
            # Helper: respond with JSON for API routes to avoid HTML parsing errors
            def api_unauthorized():
                return jsonify({'error': 'Unauthorized'}), 401

            # Not logged in
            if 'username' not in session:
                if request.path.startswith('/api/'):
                    return api_unauthorized()
                return redirect(url_for('login'))

            # Role check
            if role:
                allowed = role
                if isinstance(allowed, str):
                    allowed = (allowed,)
                if session.get('role') not in allowed:
                    flash('Unauthorized', 'danger')
                    if request.path.startswith('/api/'):
                        return api_unauthorized()
                    return redirect(url_for('login'))

            return func(*args, **kwargs)
        wrapper.__name__ = func.__name__
        return wrapper
    return decorator

def generate_ai_segments(title):
    """Generate segments and subsegments using OpenAI GPT-5-mini based on the title"""
    try:
        if not os.getenv("OPENAI_API_KEY"):
            return {'error': 'OpenAI API key not configured'}
        
        prompt = f"""
            You are a market research expert who creates concise and logical market segmentation hierarchies.

            Based on the following title: "{title}"

            Generate a clear and relevant **market segmentation structure** suitable for a professional market research report.

            ### Guidelines:
            - Create **4 to 5 main segments (Level 1)** depending on the market scope.
            - Each main segment may include **2–3 subsegments (Level 2)**.
            - Add deeper levels (**Level 2 or Level 3**) **only if necessary and meaningful** — do not force full depth if the topic is narrow and in each level.
            - Each level should become **more specific and detailed** than the previous one.
            - Avoid unnecessary repetition or overly granular splits.
            - Keep the structure **realistic, business-oriented, and readable**.
            - keep the names in maximum 3-4 words dont exceed it.
            - Dont include geographic or regional segments.
            - and at every level dont give only one point give atlest 2 points or more.

            ### Format (Strict JSON):
            {{
            "segments": [
                {{
                "name": "Level 1 Segment Name",
                "subsegments": [
                    {{
                    "name": "Level 2 Sub-segment Name",
                    "subsegments": [
                        {{
                        "name": "Level 3 Sub-sub-segment Name",
                        "subsegments": [
                            {{
                            "name": "Level 4 Sub-sub-sub-segment Name",
                            "subsegments": [] ← Must be an empty list at the deepest level
                            }}
                        ]
                        }}
                    ]
                    }}
                ]
                }}
            ]
            }}

            ### Output Rules:
            - Respond **only with valid JSON** (no explanations, text, or notes).
            - Structure depth and number of segments should **match the complexity** of "{title}".
            - Be relevant, realistic, and concise.
            """

        # Updated API call (model configurable via env)
        model_name = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
        response = openai_client.chat.completions.create(
            model=model_name,
            messages=[
                {"role": "system", "content": "You are a market research expert who creates detailed market segmentation structures. Always respond with valid JSON only."},
                {"role": "user", "content": prompt}
            ]
        )
        
        # Extract JSON safely
        content = response.choices[0].message.content.strip()
        print(f"DEBUG: AI Raw Response: {content}")
        
        json_match = re.search(r'\{.*\}', content, re.DOTALL)
        if json_match:
            json_str = json_match.group()
            parsed_result = json.loads(json_str)
            print(f"DEBUG: Parsed AI Result: {json.dumps(parsed_result, indent=2)}")
            return parsed_result
        else:
            parsed_result = json.loads(content)
            print(f"DEBUG: Parsed AI Result (direct): {json.dumps(parsed_result, indent=2)}")
            return parsed_result
            
    except json.JSONDecodeError as e:
        return {'error': f'Failed to parse AI response: {str(e)}'}
    except Exception as e:
        return {'error': f'AI generation failed: {str(e)}'}

def analyze_ai_segments(title, scraped_segments):
    """Analyze scraped segmentation entries and generate a consolidated hierarchy using OpenAI"""
    try:
        if not os.getenv("OPENAI_API_KEY"):
            return {'error': 'OpenAI API key not configured'}

        if not scraped_segments or not isinstance(scraped_segments, list):
            return {'error': 'scraped_segments must be a non-empty list'}

        # Deduplicate and sanitize scraped segments
        cleaned_segments = []
        seen = set()
        for seg in scraped_segments:
            if not isinstance(seg, str):
                continue
            # Strip numbering like "1.", "1.2.", etc.
            cleaned = re.sub(r"^\d+(?:\.\d+)*\.\s*", "", seg).strip()
            if cleaned and cleaned not in seen:
                seen.add(cleaned)
                cleaned_segments.append(cleaned)

        # Build the prompt using scraped segments as context
        segments_text = "\n".join(f"- {s}" for s in cleaned_segments[:300])  # safety cap

        prompt = f"""
            You are a market research expert. Multiple reports were scraped and produced the following segmentation entries for the topic:
            Title: "{title or 'N/A'}"

            Scraped segmentation entries (unnormalized, mixed levels):
            {segments_text}

            Task: Merge, normalize, and organize these entries into a clean hierarchical market segmentation suitable for a professional report.

            Rules:
            - Create 2–5 Level 1 segments that cover the space comprehensively.
            - Add 1–4 subsegments per parent as needed; go deeper only when meaningful (up to Level 4).
            - Deduplicate synonyms and overlapping items; prefer common industry naming.
            - Keep names concise (max 3–4 words), business-oriented, and readable.
            - Do NOT include regions/geography unless they appear explicitly and are relevant.
            - Avoid empty parents; if a segment has no children, set its subsegments to [].

            Output JSON (strict):
            {{
              "segments": [
                {{ "name": "Level 1 Segment", "subsegments": [
                  {{ "name": "Level 2 Subsegment", "subsegments": [
                    {{ "name": "Level 3 Subsegment", "subsegments": [
                      {{ "name": "Level 4 Subsegment", "subsegments": [] }}
                    ] }}
                  ] }}
                ] }}
              ]
            }}

            Respond ONLY with valid JSON.
        """

        model_name = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
        response = openai_client.chat.completions.create(
            model=model_name,
            messages=[
                {"role": "system", "content": "You are a market research expert who creates detailed, normalized segmentation hierarchies. Always respond with valid JSON only."},
                {"role": "user", "content": prompt}
            ]
        )

        content = response.choices[0].message.content.strip()
        json_match = re.search(r'\{.*\}', content, re.DOTALL)
        if json_match:
            json_str = json_match.group()
            parsed_result = json.loads(json_str)
            return parsed_result
        else:
            return json.loads(content)

    except json.JSONDecodeError as e:
        return {'error': f'Failed to parse AI response: {str(e)}'}
    except Exception as e:
        return {'error': f'AI analysis failed: {str(e)}'}

def validate_and_generate_companies(market_name, scraped_companies=None):
    """
    Validate scraped company names and auto-generate missing ones to reach exactly 20 unique, 
    active companies. Prioritizes scraped companies as top selections. Uses GPT to validate 
    and generate relevant companies.
    
    Args:
        market_name: Name of the market (used for context in generation)
        scraped_companies: Optional list of scraped company names to validate and prioritize
    
    Returns:
        Dictionary with 'COMPANY_PROFILES' list of exactly 20 company names
    """
    try:
        if not os.getenv("OPENAI_API_KEY"):
            return {'error': 'OpenAI API key not configured'}
        
        # Deduplicate and clean scraped companies
        valid_scraped = []
        if scraped_companies and isinstance(scraped_companies, list):
            seen = set()
            for company in scraped_companies:
                if not isinstance(company, str):
                    continue
                clean = company.strip()
                if clean and clean.lower() not in seen:
                    seen.add(clean.lower())
                    valid_scraped.append(clean)
        
        # Build validation and generation prompt
        scraped_text = ""
        if valid_scraped:
            scraped_text = f"""
Scraped company names from market research (prioritize these if valid):
{chr(10).join(f"- {c}" for c in valid_scraped[:100])}

IMPORTANT: Validate the scraped companies first:
- REJECT if bankrupt, defunct, or no longer operating
- REJECT if acquired or merged into another company
- REJECT if it's only a subsidiary (keep parent companies)
- ACCEPT if currently active and independently operating
- Include ACCEPTED scraped companies in the top positions

"""
        
        prompt = f"""
You are a market research expert specializing in company validation and database creation.

Task: Create a ranked list of exactly 20 unique, relevant, and currently active company names for the {market_name} market.

{scraped_text}

Requirements:
1. Output ONLY company names (no descriptions, explanations, or labels)
2. Each company must be:
   - Currently active (operating as of 2026)
   - NOT bankrupt or defunct
   - NOT merged or acquired by another company
3. All 20 companies must be unique (no duplicates or subsidiaries of the same parent)
4. Companies should be relevant to: {market_name}
5. If scraped companies were provided, validate and include them in top positions
6. Fill remaining slots with relevant companies from industry knowledge

Output format (exactly 20 companies, one per line, names only):
1. Company Name A
2. Company Name B
3. Company Name C
... (continue to exactly 20)

DO NOT:
- Add descriptions or explanations
- Include company status labels (e.g., "leading", "top")
- Exceed or fall below 20 companies
- Include duplicate companies or subsidiaries of the same parent
- Add anything after the company names
"""

        model_name = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
        response = openai_client.chat.completions.create(
            model=model_name,
            messages=[
                {"role": "system", "content": "You are a market research expert. Return ONLY a numbered list of exactly 20 company names, one per line. No other text."},
                {"role": "user", "content": prompt}
            ]
        )
        
        content = response.choices[0].message.content.strip()
        
        # Parse the numbered list
        companies = []
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            # Remove numbering (e.g., "1.", "1)", "1 -")
            clean = re.sub(r'^[\d]+[\.\)\-\s]+', '', line).strip()
            if clean and clean.lower() not in [c.lower() for c in companies]:
                companies.append(clean)
        
        # Ensure exactly 20 companies
        if len(companies) < 20:
            print(f"Warning: Generated only {len(companies)} companies, expected 20")
            # Pad with generic names if needed (shouldn't happen with good prompt)
            for i in range(len(companies), 20):
                companies.append(f"Market Player {i+1}")
        elif len(companies) > 20:
            companies = companies[:20]
        
        return {'COMPANY_PROFILES': companies}
        
    except Exception as e:
        print(f"Error in validate_and_generate_companies: {e}")
        return {'error': f'Company generation failed: {str(e)}'}

def run_spider(urls, output_dir):
    """Run Scrapy crawler in a separate process"""
    settings = get_project_settings()
    settings.set('LOG_LEVEL', 'ERROR', priority='cmdline')
    settings.set('USER_AGENT', 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36', priority='cmdline')
    process = CrawlerProcess(settings)
    process.crawl(MarketResearchSpider, urls=','.join(urls), no_docx=True)
    process.start()

@app.route('/')
def index():
    # Default to login if not authenticated; otherwise send to role home
    if 'username' in session:
        role = session.get('role')
        if role == 'admin':
            return redirect(url_for('admin_dashboard'))
        if role == 'researcher':
            return redirect(url_for('tool'))
        if role == 'creator':
            return redirect(url_for('creator_dashboard'))
    return redirect(url_for('login'))

@app.route('/tool')
@login_required(role=('researcher','admin'))
def tool():
    # Market Research Automation UI for researchers
    return render_template('index.html')

@app.route('/auth/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '')
        
        # Basic input validation
        if not username or not password:
            flash('Username and password are required', 'danger')
            return render_template('auth/login.html')
        
        # TODO: Add rate limiting here (e.g., Flask-Limiter) to prevent brute force
        user = get_user_by_username(username)
        if user and check_password_hash(user['password_hash'], password):
            session['username'] = user['username']
            session['role'] = user['role']
            update_last_login(user['username'])
            flash('Logged in', 'success')
            if user['role'] == 'admin':
                return redirect(url_for('admin_dashboard'))
            if user['role'] == 'researcher':
                return redirect(url_for('tool'))
            if user['role'] == 'creator':
                return redirect(url_for('creator_dashboard'))
            return redirect(url_for('index'))
        else:
            flash('Invalid credentials', 'danger')
    return render_template('auth/login.html')

@app.route('/auth/logout')
def logout():
    session.clear()
    flash('Logged out', 'info')
    return redirect(url_for('login'))

@app.route('/auth/admin')
@login_required(role='admin')
def admin_dashboard():
    users = list_users()
    db = get_db()
    
    # Get all RD submissions
    cur = db.execute('''
        SELECT id, market_name, researcher_username, json_path, file_path,
               submitted_at, cagr, value_unit, market_size_2024, 
               segments, companies, downloaded, last_downloaded_at
        FROM rd_submissions
        ORDER BY submitted_at DESC
    ''')
    submissions_raw = cur.fetchall()
    submissions = []
    for row in submissions_raw:
        sub = dict(row)
        # Parse JSON and get counts
        try:
            segments_list = json.loads(sub['segments']) if sub['segments'] else []
            companies_list = json.loads(sub['companies']) if sub['companies'] else []
            sub['segment_count'] = len(segments_list)
            sub['company_count'] = len(companies_list)
            sub['currency'] = sub['value_unit'] or 'USD Million'
            sub['value_2024'] = sub['market_size_2024']
        except Exception:
            sub['segment_count'] = 0
            sub['company_count'] = 0
            sub['currency'] = 'USD Million'
            sub['value_2024'] = 0
        submissions.append(sub)

    # Aggregate metrics
    total_submissions = len(submissions)
    
    # Active researchers count
    from datetime import datetime, timedelta
    def parse_iso(ts):
        try:
            dt = datetime.fromisoformat(ts) if ts else None
            # Make timezone-naive if timezone-aware
            if dt and dt.tzinfo is not None:
                dt = dt.replace(tzinfo=None)
            return dt
        except Exception:
            return None
    
    now = datetime.now(IST).replace(tzinfo=None)
    cutoff = now - timedelta(days=1)
    active_researchers_today = len({s['researcher_username'] for s in submissions
                                    if (parse_iso(s.get('submitted_at')) or now - timedelta(days=365)) >= cutoff})

    # Markets per researcher
    markets_by_researcher = {}
    for s in submissions:
        r = s['researcher_username']
        m = (s.get('market_name') or '').strip()
        if r not in markets_by_researcher:
            markets_by_researcher[r] = set()
        if m:
            markets_by_researcher[r].add(m)

    # Role counts
    researchers_count = sum(1 for u in users if u['role'] == 'researcher')

    return render_template(
        'auth/admin_dashboard.html',
        users=users,
        submissions=submissions,
        total_submissions=total_submissions,
        researchers_count=researchers_count,
        active_researchers_today=active_researchers_today,
        markets_by_researcher={k: sorted(list(v)) for k, v in markets_by_researcher.items()}
    )

@app.route('/auth/admin/create_user', methods=['GET', 'POST'])
@login_required(role='admin')
def admin_create_user():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        role = request.form.get('role')
        if role not in ('researcher', 'admin'):
            flash('Invalid role - only admin or researcher allowed', 'danger')
            return redirect(url_for('admin_create_user'))
        created_by = session.get('username')
        ok = create_user(username, password, role, created_by)
        if ok:
            flash('User created', 'success')
            return redirect(url_for('admin_dashboard'))
        else:
            flash('User already exists', 'warning')
    return render_template('auth/create_user.html')

# Removed: admin_submissions route (now uses rd_submissions table via admin_dashboard)

# Removed: admin_researcher_detail route (now uses rd_submissions table)

# Removed: admin_creators route (creator role removed)

# Removed: admin_toc route (TOC tab removed from workflow)

# Admin: list researchers
@app.route('/auth/admin/researchers')
@login_required(role='admin')
def admin_researchers():
    users = list_users()
    researchers = [u for u in users if u['role'] == 'researcher']
    return render_template('auth/admin_researchers.html', users=researchers)

@app.route('/auth/admin/submissions')
@login_required(role='admin')
def admin_submissions():
    """Detailed submissions view with all JSON data"""
    db = get_db()
    cur = db.execute('''
        SELECT id, market_name, researcher_username, submitted_at, timestamp,
               sector, industry_group, industry, sub_industry,
               value_unit, cagr, market_size_2024, market_size_2025, projected_size_2033,
               segments, ai_gen_seg, companies, created_by, version,
               downloaded, last_downloaded_at
        FROM rd_submissions
        ORDER BY submitted_at DESC
    ''')
    submissions = []
    for row in cur.fetchall():
        sub = dict(row)
        # Parse JSON fields
        try:
            sub['segments'] = json.loads(sub['segments']) if sub['segments'] else []
            sub['ai_gen_seg'] = json.loads(sub['ai_gen_seg']) if sub['ai_gen_seg'] else []
            sub['companies'] = json.loads(sub['companies']) if sub['companies'] else []
        except Exception:
            sub['segments'] = []
            sub['ai_gen_seg'] = []
            sub['companies'] = []
        submissions.append(sub)
    
    return render_template('auth/admin_submissions.html', submissions=submissions)

@app.route('/auth/researcher')
@login_required(role='researcher')
def researcher_dashboard():
    """Researcher dashboard showing their RD submissions."""
    db = get_db()
    username = session.get('username')
    
    # Get all RD submissions by this researcher
    cur = db.execute('''
        SELECT id, market_name, researcher_username, json_path, file_path,
               submitted_at, cagr, value_unit, market_size_2024,
               segments, companies, downloaded, last_downloaded_at
        FROM rd_submissions
        WHERE researcher_username = ?
        ORDER BY submitted_at DESC
    ''', (username,))
    submissions_raw = cur.fetchall()
    submissions = []
    for row in submissions_raw:
        sub = dict(row)
        # Parse JSON and get counts
        try:
            segments_list = json.loads(sub['segments']) if sub['segments'] else []
            companies_list = json.loads(sub['companies']) if sub['companies'] else []
            sub['segment_count'] = len(segments_list)
            sub['company_count'] = len(companies_list)
            sub['currency'] = sub['value_unit'] or 'USD Million'
            sub['value_2024'] = sub['market_size_2024']
        except Exception:
            sub['segment_count'] = 0
            sub['company_count'] = 0
            sub['currency'] = 'USD Million'
            sub['value_2024'] = 0
        submissions.append(sub)

    metrics = {
        'total_submissions': len(submissions),
    }

    return render_template('auth/researcher_dashboard.html', username=username, submissions=submissions, metrics=metrics)

@app.route('/auth/researcher/notifications')
@login_required(role='researcher')
def researcher_notifications():
    """Notifications page for researchers - currently placeholder since chat is removed."""
    me = session.get('username')
    messages = []
    unread_admin = 0
    return render_template('auth/researcher_notifications.html', username=me, messages=messages, unread_admin=unread_admin)

@app.route('/auth/researcher/submissions')
@login_required(role='researcher')
def researcher_submissions():
    """Detailed submissions view for current researcher with all JSON data"""
    username = session.get('username')
    db = get_db()
    cur = db.execute('''
        SELECT id, market_name, researcher_username, submitted_at, timestamp,
               sector, industry_group, industry, sub_industry,
               value_unit, cagr, market_size_2024, market_size_2025, projected_size_2033,
               segments, ai_gen_seg, companies, created_by, version,
               downloaded, last_downloaded_at
        FROM rd_submissions
        WHERE researcher_username = ?
        ORDER BY submitted_at DESC
    ''', (username,))
    submissions = []
    for row in cur.fetchall():
        sub = dict(row)
        # Parse JSON fields
        try:
            sub['segments'] = json.loads(sub['segments']) if sub['segments'] else []
            sub['ai_gen_seg'] = json.loads(sub['ai_gen_seg']) if sub['ai_gen_seg'] else []
            sub['companies'] = json.loads(sub['companies']) if sub['companies'] else []
        except Exception:
            sub['segments'] = []
            sub['ai_gen_seg'] = []
            sub['companies'] = []
        submissions.append(sub)
    
    return render_template('auth/researcher_submissions.html', username=username, submissions=submissions)


# Removed: creator_notifications route (creator role removed)

@app.route('/auth/profile')
@login_required()
def profile():
    """Simple profile page showing current user's details."""
    row = get_user_by_username(session.get('username'))
    user = dict(row) if row else {'username': session.get('username'), 'role': session.get('role')}
    return render_template('auth/profile.html', user=user)

# Removed: Chat API routes (chat_messages, chat_send, chat_mark_read, chat_unread_count)
# Removed: submission_mark_downloaded route (creator role removed)

# ===== Admin: Delete User =====
@app.route('/auth/admin/delete_user', methods=['POST'])
@login_required(role='admin')
def admin_delete_user():
    """Delete a user and related data.
    Safeguards:
      - admin-only
      - cannot delete yourself
      - cannot delete other admin users
    Body: {"username": "<user>"}
    """
    data = request.json or {}
    target = (data.get('username') or '').strip()
    if not target:
        return jsonify({'error': 'Missing `username`'}), 400
    me = session.get('username')
    if target == me:
        return jsonify({'error': 'You cannot delete your own account'}), 400
    db = get_db()
    # Check role of target
    cur = db.execute('SELECT role FROM users WHERE username = ?', (target,))
    row = cur.fetchone()
    if not row:
        return jsonify({'error': 'User not found'}), 404
    role = row['role']
    if role == 'admin':
        return jsonify({'error': 'Cannot delete another admin'}), 400
    # Delete related RD submissions
    db.execute('DELETE FROM rd_submissions WHERE researcher_username = ?', (target,))
    # Delete user
    db.execute('DELETE FROM users WHERE username = ?', (target,))
    db.commit()
    return jsonify({'status': 'ok', 'deleted': target})

# Removed: creator_dashboard route (creator role removed)

@app.route('/api/scrape', methods=['POST'])
@login_required(role=('researcher','admin'))
def scrape():
    """Endpoint to scrape URLs and return JSON results"""
    urls = request.json.get('urls', [])
    if not urls:
        return jsonify({'error': 'No URLs provided'}), 400

    output_dir = 'scraped_json'
    os.makedirs(output_dir, exist_ok=True)

    # Run Scrapy in a separate process
    p = Process(target=run_spider, args=(urls, output_dir))
    p.start()
    p.join()  # Wait for the process to complete

    # Collect JSON files from scraped_json directory
    json_files = glob.glob(os.path.join(output_dir, '*.json'))
    results = []
    for json_file in json_files:
        try:
            with open(json_file, 'r', encoding='utf-8') as f:
                results.append(json.load(f))
        except Exception as e:
            print(f"Error reading {json_file}: {e}")
        finally:
            try:
                os.remove(json_file)
            except OSError as e:
                print(f"Error removing file {json_file}: {e}")

    return jsonify(results)

# Removed: list_creators_api route (creator role removed)

# Removed: submit_report route (creator role removed)

# Removed: creator_download route (creator role removed)

@app.route('/api/generate-report', methods=['POST'])
@login_required(role=('researcher','admin'))
def generate_report():
    """Endpoint to generate a combined DOCX report from scraped data.

    This endpoint requires 'segments' in the request body and will use those
    segments (exactly as provided by UI) to generate the table of contents used
    for report creation. If 'reports' are provided, they will be used as source
    content but their ToC will be overridden by the provided 'segments'.
    """
    data = request.json or {}
    segments = data.get('segments')
    reports = data.get('reports', [])

    # If segments not at top level, try to extract from reports[0].table_of_contents
    if not isinstance(segments, list) or len(segments) == 0:
        if isinstance(reports, list) and len(reports) > 0:
            first_report = reports[0]
            if isinstance(first_report, dict):
                segments = first_report.get('table_of_contents')
                print(f"DEBUG: Extracted segments from reports[0].table_of_contents: {len(segments) if isinstance(segments, list) else 0} segments")

    # If segments STILL missing, first check autosaved 'latest_segments.json', then saved files
    if not isinstance(segments, list) or len(segments) == 0:
        try:
            current_dir = os.path.dirname(os.path.abspath(__file__))
            saved_data_dir = os.path.join(current_dir, 'saved_data')
            if os.path.exists(saved_data_dir):
                    # Check autosave file first (prefer short_rd_data folder)
                short_rd_dir = os.path.join(saved_data_dir, 'short_rd_data')
                latest_path_legacy = os.path.join(saved_data_dir, 'latest_segments.json')

                # Look for most recent short RD autosave file by market name
                if os.path.exists(short_rd_dir):
                    short_files = glob.glob(os.path.join(short_rd_dir, '*.json'))
                    if short_files:
                        latest_short = max(short_files, key=os.path.getctime)
                        try:
                            with open(latest_short, 'r', encoding='utf-8') as f:
                                latest = json.load(f)
                                latest_segments = latest.get('segments')
                                if isinstance(latest_segments, list) and len(latest_segments) > 0:
                                    segments = latest_segments
                        except Exception:
                            app.logger.exception(f'Failed to read short RD autosave file: {latest_short}')
                elif os.path.exists(latest_path_legacy):
                    # Back-compat: check older autosave location
                    with open(latest_path_legacy, 'r', encoding='utf-8') as f:
                        latest = json.load(f)
                        latest_segments = latest.get('segments')
                        if isinstance(latest_segments, list) and len(latest_segments) > 0:
                            segments = latest_segments

                # Fallback to the most recent full saved JSON file (include full_rd_data folder first)
                if not (isinstance(segments, list) and len(segments) > 0):
                    full_rd_dir = os.path.join(saved_data_dir, 'full_rd_data')
                    json_files = []
                    if os.path.exists(full_rd_dir):
                        json_files.extend(glob.glob(os.path.join(full_rd_dir, '*.json')))
                    json_files.extend(glob.glob(os.path.join(saved_data_dir, '*.json')))
                    if json_files:
                        saved_json_file = max(json_files, key=os.path.getctime)
                        with open(saved_json_file, 'r', encoding='utf-8') as f:
                            saved_data = json.load(f)
                            saved_segments = saved_data.get('segments')
                            if isinstance(saved_segments, list) and len(saved_segments) > 0:
                                segments = saved_segments
        except Exception:
            pass

    # Validate segments (final UI view segments are required either in request or saved_data)
    if not isinstance(segments, list) or len(segments) == 0:
        return jsonify({'error': 'segments is required and must be a non-empty list (provide in request or save segments first)'}), 400

    if not isinstance(reports, list) or len(reports) == 0:
        return jsonify({'error': 'No reports provided'}), 400

    combined_doc = Document()

    # Ensure each report uses the UI-provided segments as its table_of_contents
    for report in reports:
        try:
            if not isinstance(report, dict):
                report = {'title': str(report)}
            # Trust UI segments exactly as provided (do not modify)
            report['table_of_contents'] = segments
            print(f"Processing report: {json.dumps(report, indent=2)}")
            generate_docx_from_data(report, doc=combined_doc)
        except Exception as e:
            print(f"Error processing report {report.get('title', 'Unknown')}: {e}")
            combined_doc.add_paragraph(f"Error processing report: {report.get('title', 'Unknown')}: {str(e)}")
            return jsonify({'error': f"Failed to process report {report.get('title', 'Unknown')}: {str(e)}"}), 500

    # Save the combined document to a BytesIO object
    output = io.BytesIO()
    try:
        combined_doc.save(output)
        output.seek(0)
        return send_file(
            output,
            download_name='market_research_report.docx',
            as_attachment=True,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        print(f"Error sending DOCX file: {e}")
        return jsonify({'error': 'Failed to generate report'}), 500

@app.route('/api/generate-ai-segments', methods=['POST'])
@login_required(role=('researcher','admin'))
def generate_ai_segments_endpoint():
    """Endpoint to generate segments using AI based on title"""
    data = request.json
    title = data.get('title', '').strip()
    
    if not title:
        return jsonify({'error': 'Title is required'}), 400
    
    # Generate segments using AI
    result = generate_ai_segments(title)
    
    if 'error' in result:
        return jsonify(result), 500
    
    return jsonify(result)

@app.route('/api/analyze-ai-segments', methods=['POST'])
@login_required(role=('researcher','admin'))
def analyze_ai_segments_endpoint():
    """Endpoint to analyze scraped segmentation and generate consolidated segments"""
    data = request.json or {}
    title = (data.get('title') or '').strip()
    scraped_segments = data.get('scraped_segments') or []

    if not scraped_segments:
        return jsonify({'error': 'scraped_segments is required'}), 400

    result = analyze_ai_segments(title, scraped_segments)
    if 'error' in result:
        return jsonify(result), 500

    return jsonify(result)

@app.route('/api/generate-image', methods=['POST'])
def generate_image_endpoint():
    """Endpoint to generate image based on title"""
    try:
        data = request.json or {}
        title = (data.get('title') or '').strip()
        
        if not title:
            return jsonify({'error': 'Title is required'}), 400
        
        # Import here to avoid import issues
        import sys
        import os
        current_dir = os.path.dirname(os.path.abspath(__file__))
        if current_dir not in sys.path:
            sys.path.insert(0, current_dir)
        
        from image_gen import generate_market_image
        
        # Generate the image
        output_file = generate_market_image(title)
        
        # Read the generated image file
        with open(output_file, 'rb') as f:
            image_data = f.read()
        
        # Return image as base64
        import base64
        image_base64 = base64.b64encode(image_data).decode('utf-8')
        
        return jsonify({
            'success': True,
            'image': f'data:image/webp;base64,{image_base64}',
            'filename': output_file
        })
    except Exception as e:
        print(f"Error generating image: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/generate-companies', methods=['POST'])
@login_required(role=('researcher','admin'))
def generate_companies_endpoint():
    """Endpoint to validate and generate company profiles for a market"""
    data = request.json or {}
    market_name = (data.get('market_name') or '').strip()
    scraped_companies = data.get('scraped_companies') or []
    
    if not market_name:
        return jsonify({'error': 'Market name is required'}), 400
    
    result = validate_and_generate_companies(market_name, scraped_companies)
    
    if 'error' in result:
        return jsonify(result), 500
    
    return jsonify(result)

@app.route('/api/generate-excel', methods=['POST'])
@login_required(role=('researcher','admin'))
def generate_excel_endpoint():
    """Endpoint to generate Excel file from JSON data in dominating_region folder"""
    try:
        data = request.json or {}
        market_name = (data.get('market_name') or '').strip()
        
        if not market_name:
            return jsonify({'error': 'Market name is required'}), 400
        
        # Import excel_gen module
        import sys
        import os
        current_dir = os.path.dirname(os.path.abspath(__file__))
        multi_scraper_dir = os.path.join(current_dir, 'multi_scraper')
        
        if multi_scraper_dir not in sys.path:
            sys.path.insert(0, multi_scraper_dir)
        
        from excel_gen import generate_excel, clean_filename
        
        # Build path to dominating_region folder (same level as app.py)
        dominating_region_dir = os.path.join(current_dir, 'dominating_region')
        dominating_region_dir = os.path.abspath(dominating_region_dir)
        
        # Check if folder exists
        if not os.path.exists(dominating_region_dir):
            return jsonify({'error': f'dominating_region folder not found at {dominating_region_dir}'}), 404
        
        # Find JSON file matching market_name
        json_files = [f for f in os.listdir(dominating_region_dir) if f.endswith('.json')]
        
        if not json_files:
            return jsonify({'error': 'No JSON files found in dominating_region folder'}), 404
        
        # Try to find an exact match or use the first file
        json_path = None
        for json_file in json_files:
            if market_name.lower() in json_file.lower():
                json_path = os.path.join(dominating_region_dir, json_file)
                break
        
        # If no match found, use the most recent or first file
        if not json_path:
            json_path = os.path.join(dominating_region_dir, json_files[0])
        
        # Check if JSON file exists
        if not os.path.exists(json_path):
            return jsonify({'error': 'JSON file not found'}), 404
        
        # Generate Excel workbook
        wb, extracted_market_name = generate_excel(json_path)
        
        # Save to BytesIO for download
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        
        # Generate safe filename
        safe_name = clean_filename(extracted_market_name or market_name)
        excel_filename = f"{safe_name}.xlsx"
        
        return send_file(
            output,
            download_name=excel_filename,
            as_attachment=True,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
    
    except Exception as e:
        print(f"Error generating Excel: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Failed to generate Excel: {str(e)}'}), 500

@app.route('/api/save-data', methods=['POST'])
@login_required()
def save_data_endpoint():
    """Endpoint to save market research data as JSON to saved_data folder

    Now saves both 'segments' (final edited segments from UI) and 'ai_gen_seg'
    (original AI-generated segments, unchanged).
    """
    try:
        data = request.json or {}
        
        # Validate required fields
        title = (data.get('title') or '').strip()
        if not title:
            return jsonify({'error': 'Title is required'}), 400
        
        industry_classification = data.get('industryClassification', {})
        market_inputs = data.get('marketInputs', {})
        segments = data.get('segments', [])
        ai_gen_seg = data.get('ai_gen_seg', [])
        companies = data.get('companies', [])
        
        # Validate segments and companies as before
        if not isinstance(segments, list) or len(segments) == 0:
            return jsonify({'error': 'At least one segment is required'}), 400
        
        if not isinstance(companies, list) or len(companies) == 0:
            return jsonify({'error': 'At least one company is required'}), 400
        
        # Validate ai_gen_seg if provided: must be a list (can be empty)
        if ai_gen_seg is not None and not isinstance(ai_gen_seg, list):
            return jsonify({'error': 'ai_gen_seg (original AI segments) must be a list if provided'}), 400

        # Normalize ai_gen_seg into a flat numbered list of strings for saving.
        def flatten_ai_list(items):
            out = []
            def walk(arr, prefix=''):
                for idx, it in enumerate(arr):
                    num = f"{prefix}{idx+1}" if prefix == '' else f"{prefix}.{idx+1}"
                    if isinstance(it, str):
                        s = it.strip()
                        # If already numbered like '1. ...', keep it; else, format with computed num
                        if re.match(r'^\d+(?:\.\d+)*\.\s*', s):
                            out.append(s)
                        else:
                            out.append(f"{num}. {s}")
                    elif isinstance(it, dict):
                        name = str(it.get('name', '')).strip()
                        if name:
                            out.append(f"{num}. {name}")
                        subs = it.get('subsegments') or []
                        if isinstance(subs, list) and len(subs) > 0:
                            walk(subs, num)
                    else:
                        # ignore unsupported items
                        continue
            try:
                walk(items, '')
            except Exception:
                app.logger.exception('Error while flattening ai_gen_seg')
            return out

        if ai_gen_seg is None:
            ai_flat = []
        else:
            if isinstance(ai_gen_seg, list):
                # If all items are strings, use as-is (trimmed)
                if all(isinstance(x, str) for x in ai_gen_seg):
                    ai_flat = [x.strip() for x in ai_gen_seg]
                else:
                    # Mixed or hierarchical dicts -> flatten
                    ai_flat = flatten_ai_list(ai_gen_seg)
            else:
                ai_flat = []

        # Lightweight logging about prepared ai_gen_seg (flat list) for saving
        try:
            ai_len = len(ai_flat)
            app.logger.info(f"/api/save-data prepared ai_gen_seg for saving - type=flat_list, length={ai_len}")
            if ai_len > 0:
                sample = ai_flat[:3]
                app.logger.info(f"/api/save-data ai_gen_seg sample (first {len(sample)}): {sample}")
        except Exception:
            app.logger.exception('Error while logging ai_gen_seg metadata')
        
        # Create saved_data/full_rd_data folder if it doesn't exist
        current_dir = os.path.dirname(os.path.abspath(__file__))
        saved_data_dir = os.path.join(current_dir, 'saved_data')
        os.makedirs(saved_data_dir, exist_ok=True)

        full_rd_dir = os.path.join(saved_data_dir, 'full_rd_data')
        os.makedirs(full_rd_dir, exist_ok=True)

        # Generate filename from title (sanitize) — save using only market name (no timestamp)
        import re
        safe_title = re.sub(r'[<>:"/\\|?*]', '', title)[:50]  # Limit to 50 chars
        filename = f"{safe_title}.json"
        filepath = os.path.join(full_rd_dir, filename)

        # If a file with the same market name exists, back it up before overwriting
        try:
            if os.path.exists(filepath):
                bak_path = filepath + '.bak'
                # Overwrite any existing .bak with the latest backup
                if os.path.exists(bak_path):
                    os.remove(bak_path)
                os.replace(filepath, bak_path)
                app.logger.info(f"Existing file backed up to {bak_path}")
        except Exception:
            app.logger.exception('Failed to create backup of existing saved_data file')
        
        # Prepare the data structure, including original AI-generated segments
        save_data = {
            'title': title,
            'timestamp': data.get('timestamp', now_ist()),
            'industryClassification': industry_classification,
            'marketInputs': market_inputs,
            'segments': segments,
            'ai_gen_seg': ai_flat,  # Always save ai_gen_seg as a flat list of numbered strings
            'companies': companies,
            'createdBy': session.get('username', 'unknown'),
            'version': '1.0'
        }
        
        # Write to JSON file
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(save_data, f, indent=2, ensure_ascii=False)
        
        app.logger.info(f'Data saved to {filepath} by user {session.get("username")}')
        # Log ai_gen_seg saved summary
        try:
            ai_len = len(ai_gen_seg) if isinstance(ai_gen_seg, (list, tuple)) else 'non-list'
            app.logger.info(f"/api/save-data completed - ai_gen_seg saved, length={ai_len}")
        except Exception:
            app.logger.exception('Error while logging saved ai_gen_seg')

        # Insert or Update rd_submissions table (upsert by market_name)
        try:
            db = get_db()
            # Check if submission with same market_name exists
            existing = db.execute(
                'SELECT id FROM rd_submissions WHERE market_name = ?', (title,)
            ).fetchone()
            
            if existing:
                # UPDATE existing record
                db.execute('''
                    UPDATE rd_submissions SET
                        researcher_username = ?,
                        json_path = ?,
                        submitted_at = ?,
                        timestamp = ?,
                        sector = ?,
                        industry_group = ?,
                        industry = ?,
                        sub_industry = ?,
                        value_unit = ?,
                        cagr = ?,
                        market_size_2024 = ?,
                        market_size_2025 = ?,
                        projected_size_2033 = ?,
                        segments = ?,
                        ai_gen_seg = ?,
                        companies = ?,
                        created_by = ?,
                        version = ?
                    WHERE market_name = ?
                ''', (
                    session.get('username', 'unknown'),
                    filepath,
                    now_ist(),
                    save_data.get('timestamp'),
                    industry_classification.get('sector'),
                    industry_classification.get('industryGroup'),
                    industry_classification.get('industry'),
                    industry_classification.get('subIndustry'),
                    market_inputs.get('valueUnit'),
                    market_inputs.get('cagr2025_2033'),
                    market_inputs.get('marketSize2024'),
                    market_inputs.get('marketSize2025'),
                    market_inputs.get('projectedSize2033'),
                    json.dumps(segments),
                    json.dumps(ai_flat),
                    json.dumps(companies),
                    save_data.get('createdBy'),
                    save_data.get('version'),
                    title
                ))
                db.commit()
                app.logger.info(f'RD submission record UPDATED for {title}')
            else:
                # INSERT new record
                db.execute('''
                    INSERT INTO rd_submissions 
                    (market_name, researcher_username, json_path, submitted_at, timestamp,
                     sector, industry_group, industry, sub_industry,
                     value_unit, cagr, market_size_2024, market_size_2025, projected_size_2033,
                     segments, ai_gen_seg, companies, created_by, version)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (
                    title,
                    session.get('username', 'unknown'),
                    filepath,
                    now_ist(),
                    save_data.get('timestamp'),
                    industry_classification.get('sector'),
                    industry_classification.get('industryGroup'),
                    industry_classification.get('industry'),
                    industry_classification.get('subIndustry'),
                    market_inputs.get('valueUnit'),
                    market_inputs.get('cagr2025_2033'),
                    market_inputs.get('marketSize2024'),
                    market_inputs.get('marketSize2025'),
                    market_inputs.get('projectedSize2033'),
                    json.dumps(segments),
                    json.dumps(ai_flat),
                    json.dumps(companies),
                    save_data.get('createdBy'),
                    save_data.get('version')
                ))
                db.commit()
                app.logger.info(f'RD submission record CREATED for {title}')
        except Exception as e:
            app.logger.exception(f'Failed to insert/update rd_submission record: {e}')
            # Don't fail the whole operation if SQL insert fails

        return jsonify({
            'success': True,
            'message': 'Data saved successfully',
            'filename': filename,
            'filepath': filepath
        }), 200

    except Exception as e:
        print(f"Error saving data: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Failed to save data: {str(e)}'}), 500


@app.route('/api/update-submission', methods=['POST'])
@login_required()
def update_submission():
    """Update an existing RD submission from the edit modal"""
    try:
        data = request.json or {}
        submission_id = data.get('id')
        
        if not submission_id:
            return jsonify({'error': 'Submission ID is required'}), 400
        
        db = get_db()
        
        # Build update query dynamically based on provided fields
        updates = []
        params = []
        
        field_map = {
            'sector': 'sector',
            'industry_group': 'industry_group',
            'industry': 'industry',
            'sub_industry': 'sub_industry',
            'value_unit': 'value_unit',
            'cagr': 'cagr',
            'market_size_2024': 'market_size_2024',
            'market_size_2025': 'market_size_2025',
            'projected_size_2033': 'projected_size_2033',
        }
        
        for json_key, db_col in field_map.items():
            if json_key in data:
                updates.append(f'{db_col} = ?')
                params.append(data[json_key])
        
        # Handle JSON fields
        if 'segments' in data:
            updates.append('segments = ?')
            params.append(json.dumps(data['segments']) if isinstance(data['segments'], list) else data['segments'])
        
        if 'ai_gen_seg' in data:
            updates.append('ai_gen_seg = ?')
            params.append(json.dumps(data['ai_gen_seg']) if isinstance(data['ai_gen_seg'], list) else data['ai_gen_seg'])
        
        if 'companies' in data:
            updates.append('companies = ?')
            params.append(json.dumps(data['companies']) if isinstance(data['companies'], list) else data['companies'])
        
        if not updates:
            return jsonify({'error': 'No fields to update'}), 400
        
        # Add submission_id to params
        params.append(submission_id)
        
        query = f"UPDATE rd_submissions SET {', '.join(updates)} WHERE id = ?"
        db.execute(query, params)
        db.commit()
        
        # Also update the JSON file if it exists
        row = db.execute('SELECT json_path, market_name FROM rd_submissions WHERE id = ?', (submission_id,)).fetchone()
        if row and row['json_path'] and os.path.exists(row['json_path']):
            try:
                with open(row['json_path'], 'r', encoding='utf-8') as f:
                    file_data = json.load(f)
                
                # Update file data with new values
                if 'industryClassification' not in file_data:
                    file_data['industryClassification'] = {}
                if 'marketInputs' not in file_data:
                    file_data['marketInputs'] = {}
                
                if 'sector' in data:
                    file_data['industryClassification']['sector'] = data['sector']
                if 'industry_group' in data:
                    file_data['industryClassification']['industryGroup'] = data['industry_group']
                if 'industry' in data:
                    file_data['industryClassification']['industry'] = data['industry']
                if 'sub_industry' in data:
                    file_data['industryClassification']['subIndustry'] = data['sub_industry']
                if 'value_unit' in data:
                    file_data['marketInputs']['valueUnit'] = data['value_unit']
                if 'cagr' in data:
                    file_data['marketInputs']['cagr2025_2033'] = data['cagr']
                if 'market_size_2024' in data:
                    file_data['marketInputs']['marketSize2024'] = data['market_size_2024']
                if 'market_size_2025' in data:
                    file_data['marketInputs']['marketSize2025'] = data['market_size_2025']
                if 'projected_size_2033' in data:
                    file_data['marketInputs']['projectedSize2033'] = data['projected_size_2033']
                if 'segments' in data:
                    file_data['segments'] = data['segments']
                if 'ai_gen_seg' in data:
                    file_data['ai_gen_seg'] = data['ai_gen_seg']
                if 'companies' in data:
                    file_data['companies'] = data['companies']
                
                with open(row['json_path'], 'w', encoding='utf-8') as f:
                    json.dump(file_data, f, indent=2, ensure_ascii=False)
                
                app.logger.info(f'JSON file also updated for submission {submission_id}')
            except Exception as e:
                app.logger.exception(f'Failed to update JSON file: {e}')
        
        app.logger.info(f'Submission {submission_id} updated by {session.get("username")}')
        return jsonify({'success': True, 'message': 'Submission updated successfully'}), 200
        
    except Exception as e:
        app.logger.exception(f'Error updating submission: {e}')
        return jsonify({'error': f'Failed to update: {str(e)}'}), 500


@app.route('/api/mark-downloaded', methods=['POST'])
@login_required()
def mark_downloaded():
    """Mark an RD submission as downloaded when user downloads report/excel"""
    try:
        data = request.json or {}
        market_name = (data.get('market_name') or '').strip()
        
        if not market_name:
            return jsonify({'error': 'market_name is required'}), 400
        
        db = get_db()
        # Update the most recent submission for this market
        db.execute('''
            UPDATE rd_submissions 
            SET downloaded = 1, last_downloaded_at = ?
            WHERE market_name = ?
            AND id = (
                SELECT id FROM rd_submissions 
                WHERE market_name = ? 
                ORDER BY submitted_at DESC LIMIT 1
            )
        ''', (now_ist(), market_name, market_name))
        db.commit()
        
        app.logger.info(f'Marked {market_name} as downloaded by {session.get("username")}')
        return jsonify({'success': True}), 200
        
    except Exception as e:
        app.logger.exception(f'Failed to mark as downloaded: {e}')
        return jsonify({'error': str(e)}), 500


@app.route('/api/save-current-segments', methods=['POST'])
@login_required()
def save_current_segments():
    """Autosave current 'View Segments' without user clicking Save.

    Stores to 'saved_data/latest_segments.json' so other endpoints can use it
    when 'segments' is not provided in request payload.
    """

    try:
        data = request.json or {}
        segments = data.get('segments')
        if not isinstance(segments, list) or len(segments) == 0:
            return jsonify({'error': 'segments (non-empty list) is required'}), 400
        ai_gen_seg = data.get('ai_gen_seg', [])
        market_name = (data.get('market_name') or '').strip()
        
        # Accept company_data as string or companies as list, always save as string
        company_data = data.get('company_data', '')
        if not company_data and isinstance(data.get('companies'), list):
            company_data = '\n'.join([str(c) for c in data.get('companies') if c])
        company_data = company_data.strip()
        
        # Accept financial data fields (use None to detect if they were provided)
        value_2024_provided = 'value_2024' in data or 'value_2023' in data
        value_2024 = float(data.get('value_2024') or data.get('value_2023', 0)) if value_2024_provided else None
        currency = data.get('currency', '').strip() if 'currency' in data else None
        cagr = float(data.get('cagr', 0)) if 'cagr' in data else None
        
        # Accept classification and input data
        industry_classification = data.get('industryClassification')
        market_inputs = data.get('marketInputs')

        current_dir = os.path.dirname(os.path.abspath(__file__))
        saved_data_dir = os.path.join(current_dir, 'saved_data')
        os.makedirs(saved_data_dir, exist_ok=True)

        # Save autosave file into a dedicated short_rd_data folder, saved by market name
        short_rd_dir = os.path.join(saved_data_dir, 'short_rd_data')
        os.makedirs(short_rd_dir, exist_ok=True)

        # Use market name as filename; fallback to 'latest_segments' if market name missing
        import re
        safe_name = re.sub(r'[<>:"/\\|?*]', '', market_name)[:50] if market_name else 'latest_segments'
        filename = f"{safe_name}.json"
        latest_path = os.path.join(short_rd_dir, filename)

        # Backup existing short RD file if present
        try:
            if os.path.exists(latest_path):
                bak_path = latest_path + '.bak'
                if os.path.exists(bak_path):
                    os.remove(bak_path)
                os.replace(latest_path, bak_path)
                app.logger.info(f"Existing short RD file backed up to {bak_path}")
        except Exception:
            app.logger.exception('Failed to backup existing short RD file')

        # Load existing data if file exists to preserve all fields
        existing_data = {}
        if os.path.exists(latest_path):
            try:
                with open(latest_path, 'r', encoding='utf-8') as f:
                    existing_data = json.load(f)
            except Exception:
                pass
        
        # Merge with existing data, updating only provided fields
        # Use provided values if they exist, otherwise fall back to existing data
        payload = {
            'market_name': market_name,
            'segments': segments,
            'ai_gen_seg': ai_gen_seg,
            'timestamp': now_ist(),
            'savedBy': session.get('username', 'unknown'),
            'industryClassification': industry_classification if industry_classification is not None else existing_data.get('industryClassification', {}),
            'marketInputs': market_inputs if market_inputs is not None else existing_data.get('marketInputs', {}),
            'companies': company_data if company_data else existing_data.get('companies', ''),
            'value_2024': value_2024 if value_2024 is not None else existing_data.get('value_2024', 0.0),
            'currency': currency if currency else existing_data.get('currency', 'million'),
            'cagr': cagr if cagr is not None else existing_data.get('cagr', 0.0)
        }
        
        # Try to write with retry logic for Windows file locking issues
        import time
        max_retries = 3
        for attempt in range(max_retries):
            try:
                with open(latest_path, 'w', encoding='utf-8') as f:
                    json.dump(payload, f, indent=2, ensure_ascii=False)
                break  # Success, exit retry loop
            except PermissionError as pe:
                if attempt < max_retries - 1:
                    app.logger.warning(f"File locked, retrying autosave to {latest_path} (attempt {attempt + 1}/{max_retries})")
                    time.sleep(0.5)  # Wait 500ms before retry
                else:
                    # Final attempt failed, provide helpful error message
                    error_msg = f"Cannot autosave to {filename} - file may be open in another program. Please close the file."
                    app.logger.error(error_msg)
                    return jsonify({'error': error_msg, 'warning': True}), 200  # Return 200 but with warning

        app.logger.info(f'Autosaved latest segments to {latest_path}')
        return jsonify({'success': True, 'message': 'Latest segments autosaved', 'file_path': latest_path}), 200
    except Exception as e:
        print(f"Error autosaving segments: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Failed to autosave segments: {str(e)}'}), 500

@app.route('/api/generate-toc', methods=['POST'])
def generate_toc():
    """Generate Table of Contents document using toc.py"""
    try:
        from multi_scraper.toc import build_standard_toc
        
        data = request.json or {}
        
        market_name = (data.get('market_name') or 'Market').strip()
        segments_req = data.get('segments')
        ai_segments = data.get('ai_segments')  # Accept AI-generated segments
        headings = data.get('headings', [])
        levels = data.get('levels', [])
        segment_data = data.get('segment_data', '')
        company_data = data.get('company_data', '').strip()
        kmi_data = data.get('kmi_data', '').strip()  # Get KMI data
        
        if not market_name:
            return jsonify({'error': 'Market name is required'}), 400
        
        # If segments not supplied in request, try falling back to latest saved_data file
        if not isinstance(segments_req, list) or len(segments_req) == 0:
            try:
                current_dir = os.path.dirname(os.path.abspath(__file__))
                saved_data_dir = os.path.join(current_dir, 'saved_data')
                if os.path.exists(saved_data_dir):
                    json_files = glob.glob(os.path.join(saved_data_dir, '*.json'))
                    if json_files:
                        saved_json_file = max(json_files, key=os.path.getctime)
                        with open(saved_json_file, 'r', encoding='utf-8') as f:
                            saved_data = json.load(f)
                            saved_segments = saved_data.get('segments')
                            if isinstance(saved_segments, list) and len(saved_segments) > 0:
                                segments_req = saved_segments
            except Exception:
                pass

        # segments (final view segments) are mandatory and authoritative
        if not isinstance(segments_req, list) or len(segments_req) == 0:
            return jsonify({'error': 'segments is required and must be a non-empty list (provide in request or save segments first)'}), 400

        # Trust UI segments exactly as provided (do not modify numbering or text)
        toc_content = build_standard_toc(market_name, ai_segments_data=segments_req, company_names=company_data, kmi_data=kmi_data)
        
        # Create TOC document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt
        
        toc_temp_file_name = f"TOC_{market_name}_Market_SkyQuest.docx"
        toc_temp_file_path = os.path.join(tempfile.gettempdir(), toc_temp_file_name)
        
        # Create base document template path
        base_dir = os.path.dirname(os.path.abspath(__file__))
        doc_path = os.path.join(base_dir, 'multi_scraper', 'toc.docx')
        
        if not os.path.exists(doc_path):
            doc = Document()
            doc.save(doc_path)
        
        toc_doc = Document(doc_path)
        
        for heading, level in toc_content:
            paragraph = toc_doc.add_paragraph(heading)
            paragraph.style = 'List Paragraph'
            numbering = paragraph._element.get_or_add_pPr().get_or_add_numPr()
            numId = OxmlElement('w:numId')
            numId.set(qn('w:val'), '1')
            ilvl = OxmlElement('w:ilvl')
            ilvl.set(qn('w:val'), str(level))
            numbering.append(numId)
            numbering.append(ilvl)
            run = paragraph.runs[0]
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
            if level == 0:
                run.bold = True 
            paragraph.paragraph_format.line_spacing = 1.5
        
        toc_doc.save(toc_temp_file_path)
        
        return jsonify({
            'success': True,
            'file_path': toc_temp_file_path,
            'filename': toc_temp_file_name
        }), 200
    
    except Exception as e:
        print(f"Error generating TOC: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Failed to generate TOC: {str(e)}'}), 500

@app.route('/api/generate-short-rd', methods=['POST'])
def generate_short_rd():
    """Generate Short Research Document using segments from request body only"""
    try:
        data = request.json or {}
        
        market_name = (data.get('market_name') or 'Market').strip()
        headings = data.get('headings', [])
        levels = data.get('levels', [])
        segment_data = data.get('segment_data', '').strip()
        ai_segments = data.get('ai_segments')  # New: accept AI-generated segments
        segments_req = data.get('segments')  # Explicit segments from UI
        # Accept company_data as string or companies as list, always save as string
        company_data = data.get('company_data', '')
        if not company_data and isinstance(data.get('companies'), list):
            company_data = '\n'.join([str(c) for c in data.get('companies') if c])
        company_data = company_data.strip()
        # Accept both value_2023 and value_2024 for compatibility
        value_2024 = float(data.get('value_2024') or data.get('value_2023', 0))
        currency = data.get('currency', 'million').strip()
        cagr = float(data.get('cagr', 0))
        
        # Get Market Inputs (MANDATORY for Short RD)
        market_inputs = data.get('marketInputs', {})
        
        # ===== MANDATORY FIELD VALIDATION FOR SHORT RD =====
        validation_errors = []
        
        if not market_name:
            validation_errors.append('Market name')
        
        if not market_inputs or not isinstance(market_inputs, dict):
            validation_errors.append('Market Inputs')
        else:
            # Validate that market inputs has all required fields
            required_market_fields = ['unit', 'cagr', 'marketSize2024', 'projectedSize2033']
            missing_fields = []
            for field in required_market_fields:
                if not market_inputs.get(field):
                    missing_fields.append(field)
            
            if missing_fields:
                validation_errors.append(f'Market Inputs fields: {", ".join(missing_fields)}')
        
        if validation_errors:
            error_message = f'Required fields missing for Short RD generation: {", ".join(validation_errors)}'
            return jsonify({'error': error_message}), 400
        
        # --- Save all frontend data to short_rd_data JSON (before generating doc) ---
        current_dir = os.path.dirname(os.path.abspath(__file__))
        saved_data_dir = os.path.join(current_dir, 'saved_data')
        short_rd_dir = os.path.join(saved_data_dir, 'short_rd_data')
        os.makedirs(short_rd_dir, exist_ok=True)
        import re
        safe_name = re.sub(r'[<>:"/\\|?*]', '', market_name)[:50] if market_name else 'short_rd'
        json_filename = f"{safe_name}.json"
        json_path = os.path.join(short_rd_dir, json_filename)
        # Compose full frontend data for JSON
        # Load existing data if file exists to preserve all fields
        existing_data = {}
        if os.path.exists(json_path):
            try:
                with open(json_path, 'r', encoding='utf-8') as f:
                    existing_data = json.load(f)
            except Exception:
                pass
        
        # Merge with existing data, updating only provided fields
        short_rd_payload = {
            "market_name": market_name,
            "segments": data.get('segments', existing_data.get('segments', [])),
            "ai_gen_seg": data.get('ai_gen_seg', existing_data.get('ai_gen_seg', [])),
            "timestamp": data.get('timestamp', now_ist()),
            "savedBy": session.get('username', 'unknown'),
            "industryClassification": data.get('industryClassification', {}),
            "marketInputs": market_inputs if isinstance(market_inputs, dict) else existing_data.get('marketInputs', {}),
            "companies": company_data,
            "value_2024": value_2024,
            "currency": currency,
            "cagr": cagr
        }
        
        # Try to write with retry logic for Windows file locking issues
        import time
        max_retries = 3
        for attempt in range(max_retries):
            try:
                with open(json_path, 'w', encoding='utf-8') as f:
                    json.dump(short_rd_payload, f, indent=2, ensure_ascii=False)
                break  # Success, exit retry loop
            except PermissionError as pe:
                if attempt < max_retries - 1:
                    app.logger.warning(f"File locked, retrying write to {json_path} (attempt {attempt + 1}/{max_retries})")
                    time.sleep(0.5)  # Wait 500ms before retry
                else:
                    # Final attempt failed, provide helpful error message
                    error_msg = f"Cannot write to {json_filename} - file may be open in another program (text editor, Excel, etc.). Please close the file and try again."
                    app.logger.error(error_msg)
                    return jsonify({'error': error_msg}), 500

        # Build segment data for export_to_word using request body only
        toc_entries = []
        segments = []
        
        def extract_level_from_segment(segment_text):
            """Extract hierarchy level from segment numbering"""
            import re
            match = re.match(r'^(\d+(?:\.\d+)*)', segment_text)
            if match:
                numbers = match.group(1).split('.')
                return len(numbers) - 1  # 0-based level
            return 0
        
        def clean_segment_text(segment_text):
            """Remove numbering from segment text"""
            import re
            return re.sub(r'^\d+(?:\.\d+)*\.\s*', '', segment_text)
        
        # If segments not supplied in request, first check autosaved 'latest_segments.json', then full saved files
        if not isinstance(segments_req, list) or len(segments_req) == 0:
            try:
                current_dir = os.path.dirname(os.path.abspath(__file__))
                saved_data_dir = os.path.join(current_dir, 'saved_data')
                if os.path.exists(saved_data_dir):
                    # Check for autosaved short RD files (saved by market name) and pick most recent
                    short_rd_dir = os.path.join(saved_data_dir, 'short_rd_data')
                    latest_path_legacy = os.path.join(saved_data_dir, 'latest_segments.json')
                    if os.path.exists(short_rd_dir):
                        short_files = glob.glob(os.path.join(short_rd_dir, '*.json'))
                        if short_files:
                            latest_short = max(short_files, key=os.path.getctime)
                            try:
                                with open(latest_short, 'r', encoding='utf-8') as f:
                                    latest = json.load(f)
                                    latest_segments = latest.get('segments')
                                    if isinstance(latest_segments, list) and len(latest_segments) > 0:
                                        segments_req = latest_segments
                            except Exception:
                                app.logger.exception(f'Failed to read short RD autosave file: {latest_short}')
                    elif os.path.exists(latest_path_legacy):
                        with open(latest_path_legacy, 'r', encoding='utf-8') as f:
                            latest = json.load(f)
                            latest_segments = latest.get('segments')
                            if isinstance(latest_segments, list) and len(latest_segments) > 0:
                                segments_req = latest_segments
                    # Fallback to most recent saved JSON file (include full_rd_data folder)
                    if not (isinstance(segments_req, list) and len(segments_req) > 0):
                        full_rd_dir = os.path.join(saved_data_dir, 'full_rd_data')
                        json_files = []
                        if os.path.exists(full_rd_dir):
                            json_files.extend(glob.glob(os.path.join(full_rd_dir, '*.json')))
                        json_files.extend(glob.glob(os.path.join(saved_data_dir, '*.json')))
                        if json_files:
                            saved_json_file = max(json_files, key=os.path.getctime)
                            with open(saved_json_file, 'r', encoding='utf-8') as f:
                                saved_data = json.load(f)
                                saved_segments = saved_data.get('segments')
                                if isinstance(saved_segments, list) and len(saved_segments) > 0:
                                    segments_req = saved_segments
            except Exception:
                pass

        # segments (final view segments) are mandatory and authoritative
        if not isinstance(segments_req, list) or len(segments_req) == 0:
            return jsonify({'error': 'segments is required and must be a non-empty list (provide in request or save segments first)'}), 400

        # Build toc_entries from segments (trust UI strings exactly but parse level)
        for segment in segments_req:
            s = str(segment).strip()
            if not s:
                continue
            level = extract_level_from_segment(s)
            clean_text = clean_segment_text(s)
            if level == 0:
                toc_heading = f"Global {market_name} Size by {clean_text} & CAGR (2026-2033)"
                toc_entries.append((toc_heading, level))
                toc_entries.append(("Market Overview", 1))
                segments.append((toc_heading, level))
                segments.append(("Market Overview", 1))
            else:
                toc_entries.append((clean_text, level))
                segments.append((clean_text, level))
        
        # Generate Short RD
        rd_temp_file_name = f"RD_{market_name}_SkyQuest.docx"
        rd_temp_file_path = os.path.join(tempfile.gettempdir(), rd_temp_file_name)
        
        export_to_word(
            data=toc_entries,
            market_name=market_name,
            value_2024=value_2024,
            currency=currency,
            cagr=cagr,
            companies=company_data,
            output_path=rd_temp_file_path
        )
        
        return jsonify({
            'success': True,
            'file_path': rd_temp_file_path,
            'filename': rd_temp_file_name,
            'short_rd_json': json_path
        }), 200
    
    except Exception as e:
        print(f"Error generating Short RD: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Failed to generate Short RD: {str(e)}'}), 500

@app.route('/api/download-file')
def download_file():
    """Download generated file"""
    try:
        file_path = request.args.get('file_path')
        
        if not file_path:
            return "Error: File path is required.", 400
        
        # Prevent path traversal attacks - normalize and validate path
        file_path = os.path.abspath(file_path)
        temp_dir = os.path.abspath(tempfile.gettempdir())
        current_dir = os.path.abspath(os.path.dirname(__file__))
        
        # Only allow files from temp directory or current directory tree
        if not (file_path.startswith(temp_dir) or file_path.startswith(current_dir)):
            app.logger.warning(f'Path traversal attempt blocked: {file_path}')
            return "Error: Invalid file path.", 403
        
        if not os.path.exists(file_path):
            return "Error: The file does not exist.", 404
        
        file_name = os.path.basename(file_path)
        
        file_name = os.path.basename(file_path)
        
        return send_file(
            file_path,
            as_attachment=True,
            download_name=file_name, 
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    
    except Exception as e:
        return f"Error downloading file: {str(e)}", 500

@app.route('/api/submit-to-skyquest', methods=['POST'])
@login_required(role='admin')
def submit_to_skyquest():
    """Submit report to SkyQuest website API - Admin only
    
    Process: Get token -> Generate DOCX -> Generate Image -> Generate Excel -> Upload
    Each step is done sequentially with no retries.
    Note: DOCX takes ~10min, Image ~1-2min, Excel ~2min
    """
    import requests
    import shutil
    
    try:
        data = request.json or {}
        submission_id = data.get('submission_id')
        
        if not submission_id:
            return jsonify({'error': 'submission_id is required'}), 400
        
        # Get submission from database
        db = get_db()
        cur = db.execute('SELECT * FROM rd_submissions WHERE id = ?', (submission_id,))
        row = cur.fetchone()
        
        if not row:
            return jsonify({'error': 'Submission not found'}), 404
        
        submission = dict(row)
        market_name = submission.get('market_name', 'Unknown')
        
        # Create website_submission folder
        current_dir = os.path.dirname(os.path.abspath(__file__))
        safe_market_name = market_name.replace(' ', '_').replace('/', '-')
        submission_folder = os.path.join(current_dir, 'website_submission', safe_market_name)
        os.makedirs(submission_folder, exist_ok=True)
        
        print(f'[SkyQuest] Starting submission for: {market_name}')
        print(f'[SkyQuest] Output folder: {submission_folder}')
        
        # Step 1: Get SkyQuest token
        print('[SkyQuest] Step 1/5: Getting auth token...')
        login_url = 'https://www.skyquestt.com/api/login'
        
        # Get credentials from environment variables
        skyquest_email = os.getenv('SKYQUEST_EMAIL')
        skyquest_password = os.getenv('SKYQUEST_PASSWORD')
        
        if not skyquest_email or not skyquest_password:
            return jsonify({'error': 'SkyQuest credentials not configured. Set SKYQUEST_EMAIL and SKYQUEST_PASSWORD environment variables.'}), 500
        
        login_response = requests.post(login_url, data={
            'email': skyquest_email,
            'password': skyquest_password
        }, timeout=60)
        
        if login_response.status_code != 200:
            return jsonify({'error': f'Failed to get SkyQuest token: {login_response.text}'}), 500
        
        token_data = login_response.json()
        token = token_data.get('token')
        
        if not token:
            return jsonify({'error': 'No token received from SkyQuest'}), 500
        
        print(f'[SkyQuest] Token received, expires: {token_data.get("expires_in")}')
        
        # Step 2: Generate DOCX report (takes ~10 minutes)
        print('[SkyQuest] Step 2/5: Generating DOCX report (this may take 10+ minutes)...')
        
        segments = submission.get('segments', [])
        if isinstance(segments, str):
            try:
                segments = json.loads(segments)
            except:
                segments = []
        
        ai_gen_seg = submission.get('ai_gen_seg', [])
        if isinstance(ai_gen_seg, str):
            try:
                ai_gen_seg = json.loads(ai_gen_seg)
            except:
                ai_gen_seg = []
        
        companies = submission.get('companies', [])
        if isinstance(companies, str):
            try:
                companies = json.loads(companies)
            except:
                companies = []
        
        # Debug: Print all submission data
        print(f'[SkyQuest] Submission data:')
        print(f'[SkyQuest]   - sector: {submission.get("sector")}')
        print(f'[SkyQuest]   - industry_group: {submission.get("industry_group")}')
        print(f'[SkyQuest]   - industry: {submission.get("industry")}')
        print(f'[SkyQuest]   - sub_industry: {submission.get("sub_industry")}')
        print(f'[SkyQuest]   - cagr: {submission.get("cagr")}')
        print(f'[SkyQuest]   - value_unit: {submission.get("value_unit")}')
        print(f'[SkyQuest]   - segments count: {len(segments)}')
        print(f'[SkyQuest]   - companies count: {len(companies)}')
        
        # Build report data with all required fields
        report_data = {
            'title': market_name,
            'table_of_contents': segments,
            'companies': companies,
            'industry_classification': {
                'sector': submission.get('sector', ''),
                'industry_group': submission.get('industry_group', ''),
                'industry': submission.get('industry', ''),
                'sub_industry': submission.get('sub_industry', '')
            },
            'market_inputs': {
                'cagr': submission.get('cagr'),
                'unit': submission.get('value_unit', 'Million'),
                'value_2024': submission.get('market_size_2024'),
                'value_2025': submission.get('market_size_2025'),
                'value_2033': submission.get('projected_size_2033')
            }
        }
        
        # Check for existing DOCX in submission folder
        import time
        import glob as glob_module
        
        existing_docx = glob_module.glob(os.path.join(submission_folder, '*.docx'))
        if existing_docx:
            doc_path = existing_docx[0]
            print(f'[SkyQuest] Using existing DOCX: {doc_path}')
        else:
            doc = Document()
            generate_docx_from_data(report_data, doc=doc)
            doc_path = os.path.join(submission_folder, f'{safe_market_name}.docx')
            doc.save(doc_path)
            print(f'[SkyQuest] DOCX saved: {doc_path}')
        
        # Step 3: Generate Image (takes ~1-2 minutes) with retry
        print('[SkyQuest] Step 3/5: Checking for existing image or generating new...')
        
        # Check for existing image
        existing_images = glob_module.glob(os.path.join(submission_folder, '*.webp'))
        if not existing_images:
            existing_images = glob_module.glob(os.path.join(submission_folder, '*.png'))
        if not existing_images:
            existing_images = glob_module.glob(os.path.join(submission_folder, '*.jpg'))
        
        image_path = None
        if existing_images:
            image_path = existing_images[0]
            print(f'[SkyQuest] Using existing image: {image_path}')
        else:
            # Try up to 10 times with 30 sec delay
            max_retries = 10
            for attempt in range(1, max_retries + 1):
                print(f'[SkyQuest] Image generation attempt {attempt}/{max_retries}...')
                try:
                    from image_gen import generate_market_image
                    image_output = generate_market_image(market_name)
                    image_dest = os.path.join(submission_folder, os.path.basename(image_output))
                    if os.path.exists(image_output):
                        shutil.move(image_output, image_dest)
                    image_path = image_dest
                    print(f'[SkyQuest] Image saved: {image_path}')
                    break
                except Exception as e:
                    print(f'[SkyQuest] Image attempt {attempt} failed: {e}')
                    if attempt < max_retries:
                        print(f'[SkyQuest] Waiting 30 seconds before retry...')
                        time.sleep(30)
                    else:
                        print('[SkyQuest] All image attempts failed, continuing without image...')
                        image_path = None
        
        # Step 4: Generate Excel if data exists (takes ~2 minutes)
        print('[SkyQuest] Step 4/5: Checking for existing Excel or generating new...')
        
        # Check for existing Excel
        existing_excel = glob_module.glob(os.path.join(submission_folder, '*.xlsx'))
        
        excel_path = None
        if existing_excel:
            excel_path = existing_excel[0]
            print(f'[SkyQuest] Using existing Excel: {excel_path}')
        else:
            try:
                import sys
                multi_scraper_dir = os.path.join(current_dir, 'multi_scraper')
                if multi_scraper_dir not in sys.path:
                    sys.path.insert(0, multi_scraper_dir)
                from excel_gen import generate_excel, generate_template_excel, clean_filename
                
                # Try to find matching JSON in dominating_region folder
                dominating_region_dir = os.path.join(current_dir, 'dominating_region')
                matching_json = None
                
                if os.path.exists(dominating_region_dir):
                    json_files = [f for f in os.listdir(dominating_region_dir) if f.endswith('.json')]
                    
                    # Try multiple matching strategies
                    # 1. Exact match (case insensitive)
                    for jf in json_files:
                        if market_name.lower() in jf.lower():
                            matching_json = os.path.join(dominating_region_dir, jf)
                            break
                    
                    # 2. If no match, try removing " Market" suffix and match beginning
                    if not matching_json:
                        market_name_base = market_name.lower().replace(' market', '').strip()
                        for jf in json_files:
                            jf_lower = jf.lower()
                            # Check if filename starts with the market name base
                            if jf_lower.startswith(market_name_base):
                                matching_json = os.path.join(dominating_region_dir, jf)
                                print(f'[SkyQuest] Matched by prefix: {jf}')
                                break
                    
                    # 3. If still no match, try fuzzy match (first 30 chars)
                    if not matching_json and len(market_name) > 30:
                        market_prefix = market_name.lower()[:30]
                        for jf in json_files:
                            if market_prefix in jf.lower():
                                matching_json = os.path.join(dominating_region_dir, jf)
                                print(f'[SkyQuest] Matched by fuzzy prefix: {jf}')
                                break
                
                if matching_json:
                    # Generate Excel from dominating_region data
                    wb, extracted_name = generate_excel(matching_json)
                    excel_path = os.path.join(submission_folder, f'{clean_filename(extracted_name or market_name)}.xlsx')
                    wb.save(excel_path)
                    print(f'[SkyQuest] Excel generated from data: {excel_path}')
                else:
                    # Generate template Excel if no data available
                    print('[SkyQuest] No dominating_region data found, generating template Excel...')
                    value_unit = submission.get('value_unit', 'Million')
                    wb = generate_template_excel(market_name, value_unit)
                    excel_path = os.path.join(submission_folder, f'{clean_filename(market_name)}.xlsx')
                    wb.save(excel_path)
                    print(f'[SkyQuest] Template Excel generated: {excel_path}')
                    
            except Exception as e:
                print(f'[SkyQuest] Excel generation failed: {e}')
                import traceback
                traceback.print_exc()
                # Generate template Excel as fallback
                try:
                    from excel_gen import generate_template_excel, clean_filename
                    value_unit = submission.get('value_unit', 'Million')
                    wb = generate_template_excel(market_name, value_unit)
                    excel_path = os.path.join(submission_folder, f'{clean_filename(market_name)}.xlsx')
                    wb.save(excel_path)
                    print(f'[SkyQuest] Fallback template Excel generated: {excel_path}')
                except Exception as fallback_error:
                    print(f'[SkyQuest] Fallback Excel generation also failed: {fallback_error}')
                    excel_path = None
        
        # Step 5: Upload to SkyQuest
        print('[SkyQuest] Step 5/5: Uploading to SkyQuest API...')
        upload_url = 'https://www.skyquestt.com/api/importRDFile'
        headers = {'Authorization': f'Bearer {token}'}
        
        files = {}
        file_handles = []
        
        if os.path.exists(doc_path):
            fh = open(doc_path, 'rb')
            files['report_rd'] = fh
            file_handles.append(fh)
        
        if image_path and os.path.exists(image_path):
            fh = open(image_path, 'rb')
            files['report_image'] = fh
            file_handles.append(fh)
        
        if excel_path and os.path.exists(excel_path):
            fh = open(excel_path, 'rb')
            files['report_graph'] = fh
            file_handles.append(fh)
        
        # Validate required files before upload
        if 'report_rd' not in files:
            return jsonify({'error': 'DOCX report is required but could not be generated'}), 400
        
        if 'report_graph' not in files:
            return jsonify({'error': 'Excel file is required but could not be generated. Check logs for details.'}), 400
        
        print(f'[SkyQuest] Files to upload:')
        print(f'[SkyQuest]   - report_rd: {doc_path}')
        print(f'[SkyQuest]   - report_image: {image_path if image_path else "None (optional)"}')
        print(f'[SkyQuest]   - report_graph: {excel_path}')
        
        try:
            upload_response = requests.post(upload_url, headers=headers, files=files, timeout=300)
            
            # Close all file handles
            for fh in file_handles:
                fh.close()
            
            print(f'[SkyQuest] Response status: {upload_response.status_code}')
            print(f'[SkyQuest] Response body: {upload_response.text}')
            
            if upload_response.status_code == 200:
                print('[SkyQuest] Upload successful!')
                
                # Update submission status in database
                db.execute('''
                    UPDATE rd_submissions 
                    SET downloaded = 1, last_downloaded_at = ? 
                    WHERE id = ?
                ''', (datetime.now(timezone.utc).isoformat(), submission_id))
                db.commit()
                
                return jsonify({
                    'success': True,
                    'message': f'Successfully submitted {market_name} to SkyQuest',
                    'files': {
                        'doc': doc_path if os.path.exists(doc_path) else None,
                        'image': image_path,
                        'excel': excel_path
                    },
                    'skyquest_response': upload_response.text
                })
            else:
                print(f'[SkyQuest] Upload failed: {upload_response.status_code}')
                return jsonify({
                    'error': f'SkyQuest upload failed with status {upload_response.status_code}',
                    'details': upload_response.text
                }), 500
                
        except Exception as e:
            # Close file handles on error
            for fh in file_handles:
                try:
                    fh.close()
                except:
                    pass
            raise e
        
    except Exception as e:
        print(f'[SkyQuest] Error: {str(e)}')
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Failed to submit to SkyQuest: {str(e)}'}), 500

if __name__ == '__main__':
    port = int(os.getenv('PORT', '5001'))
    # Use stat reloader instead of watchdog to avoid Windows socket errors
    app.run(host='0.0.0.0', port=port, debug=True, use_reloader=True, reloader_type='stat')
