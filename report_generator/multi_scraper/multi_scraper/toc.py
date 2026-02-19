from flask import Flask, render_template, request, send_file, redirect, url_for
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os
from docx.shared import RGBColor
import json
import glob

def build_standard_toc(market_name, ai_segments_data=None, company_names=None, kmi_data=None):
    """
    Build complete TOC structure with:
    0. Key Market Insights (from KMI data)
    1. Market Segments (from saved JSON)
    2. Regions
    """
    toc = []
    
    if kmi_data:
        if isinstance(kmi_data, list):
            for kmi in kmi_data:
                if isinstance(kmi, tuple):
                    toc.append(kmi)
                else:
                    toc.append((str(kmi), 1))
        elif isinstance(kmi_data, str):
            for kmi in kmi_data.strip().split('\n'):
                if kmi.strip():
                    toc.append((kmi.strip(), 1))
    
    segment_names = []
    segments_list = None
    
    if ai_segments_data:
        if isinstance(ai_segments_data, dict) and 'segments' in ai_segments_data:
            segments_list = ai_segments_data['segments']
        elif isinstance(ai_segments_data, list):
            segments_list = ai_segments_data
    
    if segments_list:
        def extract_level_from_segment(segment_text):
            import re
            match = re.match(r'^(\d+(?:\.\d+)*)', segment_text)
            if match:
                numbers = match.group(1).split('.')
                return len(numbers) - 1  # 0-based level
            return 0
        
        def clean_segment_text(segment_text):
            import re
            return re.sub(r'^\d+(?:\.\d+)*\.\s*', '', segment_text)
        
        if segments_list and isinstance(segments_list[0], str):
            current_segment_group = None
            current_segment_group_level = -1
            
            for segment in segments_list:
                if isinstance(segment, str) and segment.strip():
                    level = extract_level_from_segment(segment)
                    clean_text = clean_segment_text(segment)
                    
                    if level == 0:
                        segment_names.append(clean_text)
                        toc.append((f"Global {market_name} Size by {clean_text} & CAGR (2026-2033)", 0))
                        toc.append(("Market Overview", 1))
                        current_segment_group = clean_text
                        current_segment_group_level = 0
                    else:
                        toc.append((clean_text, level))
    
    toc.append((f"Global {market_name} Size & CAGR (2026-2033)", 0))
    
    segment_text = ", ".join(segment_names) if segment_names else "Product Type, Application, Chemistry, End User, Market Structure"
    
    regions = [
        ("North America", ["US", "Canada"]),
        ("Europe", ["Germany", "Spain", "France", "UK", "Italy", "Rest of Europe"]),
        ("Asia Pacific", ["China", "India", "Japan", "South Korea", "Rest of Asia-Pacific"]),
        ("Latin America", ["Mexico", "Brazil", "Rest of Latin America"]),
        ("Middle East & Africa", ["GCC Countries", "South Africa", "Rest of Middle East & Africa"]),
    ]
    
    for region, countries in regions:
        toc.append((f"{region} ({segment_text})", 1))
        for country in countries:
            toc.append((country, 2))
    
    toc.append(("Competitive Intelligence", 0))
    toc.append(("Top 5 Player Comparison", 1))
    toc.append(("Market Positioning of Key Players, 2025", 1))
    toc.append(("Strategies Adopted by Key Market Players", 1))
    toc.append(("Recent Developments in the Market", 1))
    toc.append(("Company Market Share Analysis, 2025", 1))
    toc.append(("Company Profiles of All Key Players", 1))
    toc.append(("Company Details", 2))
    toc.append(("Product Portfolio Analysis", 2))
    toc.append(("Company's Segmental Share Analysis", 2))
    toc.append(("Revenue Y-O-Y Comparison (2023-2025)", 2))
    
    toc.append(("Key Company Profiles", 0))
    
    if company_names:
        if isinstance(company_names, str):
            company_list = [c.strip() for c in company_names.split('\n') if c.strip()]
        else:
            company_list = company_names if isinstance(company_names, list) else []
        
        for company in company_list[:50]:
            toc.append((company, 1))
            toc.append(("Company Overview", 2))
            toc.append(("Business Segment Overview", 2))
            toc.append(("Financial Updates", 2))
            toc.append(("Key Developments", 2))
    
    toc.append(("Conclusion & Recommendations", 0))
    
    return toc

def transform_market_data(data, market_name):
    segments = {}
    current_segment = None
    current_level = 0

    for item, level in data:
        if level not in [0, 1, 2]:
            continue
        if level == 0:
            if f"Global {market_name} Size by" in item:
                segment_name = item.replace(f"Global {market_name} Size by ", "").split(" & CAGR")[0]
                current_segment = segment_name
                segments[current_segment] = []
            continue

        if current_segment and level > 0:
            if item != "Market Overview":
                if level == 1:
                    segments[current_segment].append(item)
                elif level == 2:
                    if segments[current_segment] and isinstance(segments[current_segment][-1], list):
                        segments[current_segment][-1].append(item)
                    else:
                        segments[current_segment].append([item])
    
    if not segments:
        segments["General"] = ["Market Overview", "Analysis"]
    
    formatted_output = []
    for segment, sub_segments in segments.items():
        formatted_subs = []
        for sub in sub_segments:
            if isinstance(sub, list):
                if formatted_subs:
                    formatted_subs[-1] += f" ({', '.join(sub)})"
            else:
                formatted_subs.append(sub)

        if formatted_subs:
            segment_line = f"Segment {segment}: Sub-Segments {', '.join(formatted_subs)}"
            formatted_output.append(segment_line)
    
    return formatted_output, segments

def generate_segmental_analysis(segments_data, market_name):
    text = f"Global {market_name} is segmented by "
    segment_names = []
    segment_details = []

    for segment, sub_segments in segments_data.items():
        segment_names.append(segment)
        sub_details = []
        for i, sub in enumerate(sub_segments):
            if isinstance(sub, list):
                continue
            else:
                sub_details.append(sub)

        if not sub_details:
            segment_details.append(
                f"Based on {segment}, no specific sub-segments were identified."
            )
        elif len(sub_details) > 1:
            joined_sub_details = ", ".join(sub_details[:-1]) + " and " + sub_details[-1]
        else:
            joined_sub_details = sub_details[0]

        if sub_details:
            segment_details.append(
                f"Based on {segment}, the market is segmented into {joined_sub_details}."
            )

    if len(segment_names) > 1:
        text += ", ".join(segment_names) + " and region. "
    elif len(segment_names) == 1:
        text += segment_names[0] + " and region. "
    else:
        text += "various factors and region. "

    text += " ".join(segment_details)
    text = (
        text
        + " Based on region, the market is segmented into North America, Europe, Asia Pacific, Latin America and Middle East & Africa. "
    )
    return text

def title_h1(segments_data, market_name):
    segments_text_list = []
    
    for segment, sub_segments in segments_data.items():
        top_level_subsegments = []  
        
        for sub in sub_segments:
            if isinstance(sub, str):  
                top_level_subsegments.append(sub)
        if top_level_subsegments:
            subsegments_text = f" ({', '.join(top_level_subsegments[:2])})"
        else:
            subsegments_text = ""

        segments_text_list.append("By " + segment + subsegments_text)

    region_text_full = "By Region"
    region_text_short = "By Region"
    
    if len(segments_text_list) >= 1:
        segments_text_list.append(region_text_full)
        title_text = f"{market_name} Size, Share, Growth Analysis, {', '.join(segments_text_list)} - Industry Forecast 2026-2033"
        word_count = len(title_text.split())
    else:
        title_text = f"{market_name} Size, Share, Growth Analysis, By Region - Industry Forecast 2026-2033"
        word_count = len(title_text.split())
        return title_text
    
    if word_count > 35:
        segments_text_list[-1] = region_text_short
        segments_text = ", ".join(segments_text_list)
        title_text = f"{market_name} Size, Share, Growth Analysis, {segments_text} - Industry Forecast 2026-2033"
        word_count = len(title_text.split())
    
    if word_count > 35 and len(segments_text_list) > 1:
        for i in range(len(segments_text_list) - 2, -1, -1):
            if "(" in segments_text_list[i]:
                segment_name = segments_text_list[i].split("(")[0].strip()
                segments_text_list[i] = segment_name
                break
        
        segments_text = ", ".join(segments_text_list)
        title_text = f"{market_name} Size, Share, Growth Analysis, {segments_text} - Industry Forecast 2026-2033"
        word_count = len(title_text.split())
    
    if word_count > 35 and len(segments_text_list) > 1:
        simplified_list = []
        for i, seg in enumerate(segments_text_list[:-1]):
            if i == len(segments_text_list) - 2:
                simplified_list.append(seg)
            else:
                segment_name = seg.split("(")[0].strip()
                simplified_list.append(segment_name)
        simplified_list.append(segments_text_list[-1])
        segments_text = ", ".join(simplified_list)
        title_text = f"{market_name} Size, Share, Growth Analysis, {segments_text} - Industry Forecast 2026-2033"

    return title_text

def export_to_word(data, market_name, value_2024, currency, cagr, companies, output_path="Market_Report.docx"):

    value_2025 = value_2024 * (1 + cagr / 100) ** 1
    value_2032 = value_2024 * (1 + cagr / 100) ** 9
    value_2025 = round(value_2025, 2)
    value_2032 = round(value_2032, 2)

    doc = Document()
    formatted_output, segments = transform_market_data(data,market_name)

    def set_poppins_style(paragraph, size=12, bold=False, color=RGBColor(0, 0, 0)):
        run = paragraph.add_run()
        run.font.name = "Poppins"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Poppins")
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = color  
        return run

    title = doc.add_heading(level=1)
    title_run = set_poppins_style(title, size=16, bold=True, color=RGBColor(0, 0, 0))
    title_run.text = "Report Name"
    text_paragraph = doc.add_paragraph()
    text_run = set_poppins_style(text_paragraph, size=12, color=RGBColor(0, 0, 0))
    text_run.text = f"{market_name}"

    upcoming = doc.add_heading(level=1)
    upcoming_run = set_poppins_style(upcoming, size=16, bold=True, color=RGBColor(0, 0, 0))
    upcoming_run.text = "Upcoming"
    text_paragraph = doc.add_paragraph()
    text_run = set_poppins_style(text_paragraph, size=12, color=RGBColor(0, 0, 0))
    text_run.text = "No"

    segments_heading = doc.add_heading(level=1)
    segments_heading_run = set_poppins_style(segments_heading, size=16, bold=True, color=RGBColor(0, 0, 0))
    segments_heading_run.text = "Segments"

    for seg_name, sub_segments in segments.items():
        segment_heading = doc.add_heading(level=2)
        segment_run = set_poppins_style(segment_heading, size=16, bold=True, color=RGBColor(0, 0, 0))
        segment_run.text = "Segment"

        segment_name_paragraph = doc.add_paragraph()
        segment_name_run = set_poppins_style(segment_name_paragraph, size=12, color=RGBColor(0, 0, 0))
        segment_name_run.text = seg_name

        level1_items = []
        level2_children = {}
        current_l1 = None
        for sub in sub_segments:
            if isinstance(sub, str):
                level1_items.append(sub)
                current_l1 = sub
                level2_children[sub] = []
            elif isinstance(sub, list) and current_l1:
                level2_children[current_l1].extend(sub)

        sub_segment_label = doc.add_heading(level=2)
        sub_segment_label_run = set_poppins_style(
            sub_segment_label, size=16, bold=False, color=RGBColor(0, 0, 0)
        )
        sub_segment_label_run.text = "Sub-Segments"

        sub_segment_paragraph = doc.add_paragraph()
        sub_segment_run = set_poppins_style(sub_segment_paragraph, size=12, color=RGBColor(0, 0, 0))
        sub_segment_run.text = ", ".join(level1_items)

        for l1_name, l2_items in level2_children.items():
            if l2_items:
                sub_sub_label = doc.add_heading(level=2)
                sub_sub_label_run = set_poppins_style(
                    sub_sub_label, size=16, bold=False, color=RGBColor(0, 0, 0)
                )
                sub_sub_label_run.text = f"Sub-Segments-For-{l1_name}"

                sub_sub_paragraph = doc.add_paragraph()
                sub_sub_run = set_poppins_style(sub_sub_paragraph, size=12, color=RGBColor(0, 0, 0))
                sub_sub_run.text = ", ".join(l2_items)

    market_heading = doc.add_heading(level=1)
    market_heading_run = set_poppins_style(market_heading, size=16, bold=True, color=RGBColor(0, 0, 0))
    market_heading_run.text = "Market Insights"

    text_paragraph = doc.add_paragraph()
    text_run = set_poppins_style(text_paragraph, size=12, color=RGBColor(0, 0, 0))
    text_run.text = f"Global {market_name} size was valued at USD {value_2024} {currency} in 2024 and is poised to grow from USD {value_2025} {currency} in 2025 to USD {value_2032} {currency} by 2033, growing at a CAGR of {cagr}% during the forecast period (2026-2033)."

    market_heading_1 = doc.add_heading(level=1)
    market_heading_run_1 = set_poppins_style(market_heading_1, size=16, bold=True, color=RGBColor(0, 0, 0))
    market_heading_run_1.text = "Segmental Analysis"

    text_paragraph = doc.add_paragraph()
    text_run = set_poppins_style(text_paragraph, size=12, color=RGBColor(0, 0, 0))
    text_run.text = generate_segmental_analysis(segments, market_name)

    top_players_heading = doc.add_paragraph()
    top_players_run = set_poppins_style(top_players_heading, size=12, bold=True, color=RGBColor(0, 0, 0))
    top_players_run.text = "Top Player's Company Profiles"
    
    for c1 in companies.splitlines():
        if not c1.strip():
            continue
        company_paragraph = doc.add_paragraph(style="Normal")
        company_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pPr = company_paragraph._element.get_or_add_pPr()
        numPr = pPr.get_or_add_numPr()
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), '0')
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), '1')
        numPr.append(ilvl)
        numPr.append(numId)
        company_run = set_poppins_style(company_paragraph, size=12, color=RGBColor(0, 0, 0))
        company_run.text = c1.strip()

    H1_text = doc.add_heading(level=1)
    H1_run = set_poppins_style(H1_text, size=16, bold=True, color=RGBColor(0, 0, 0))
    H1_run.text = "H1 Title"

    text_paragraph = doc.add_paragraph()
    title_text = title_h1(segments, market_name)
    text_run = set_poppins_style(
        text_paragraph,
        size=12,
        color=RGBColor(255, 0, 0) if len(title_text.split()) > 35 else RGBColor(0, 0, 0)
    )
    text_run.text = title_text
    doc.save(output_path)
    return output_path

app = Flask(__name__)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
doc_path = os.path.join(BASE_DIR, 'toc.docx')
rd_path = os.path.join(BASE_DIR, 'rd.docx')

if not os.path.exists(doc_path):
    doc = Document()
    doc.save(doc_path)

def get_level(i):
    return i.split(" ", 1)[0].count(".")

def clean(name):
    a = name.split(" ", 1)[1]
    if "(Page No." in a:
        a = a.split(" (Page No.", 1)[0].strip()
    return a.strip()

def get_level1(i):
    try:
        i = i.replace("\t", " ")
        return i.split(" ", 1)[0].count(".")
    except:
        return i.split(" ", 1)[0].count(".")

def clean1(name):
    try:
        name = name.replace("\t", " ")
        a = name.split(" ", 1)[1]
    except:
        a = name.split(" ", 1)[1]
    if "(Page No." in a:
        a = a.split(" (Page No.", 1)[0].strip()
    return a.strip()

def add_bullet_point_text(doc, text, level):
    paragraph = doc.add_paragraph(text)
    paragraph.style = 'List Paragraph'
    numbering = paragraph._element.get_or_add_pPr().get_or_add_numPr()
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), '1')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level))
    numbering.append(numId)
    numbering.append(ilvl)
    run = paragraph.runs[0]
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    if level == 0:
        run.bold = True 
    paragraph.paragraph_format.line_spacing = 1.5

@app.route("/", methods=["GET", "POST"])
def index():
    data = {
        "kmi": [],
        "toc_entries": [],
        "segments": [],
        "regions": [],
        "companies": [],
        "conclusion":[("Conclusion & Recommendations",0)]
    }
    raw_segments = []
    market_details = {
    "market_name": request.form.get("market_name", "").strip(),
    "value_2024": float(request.form.get("value_2024") or  0),
    "currency": request.form.get("currency", "million").strip(),
    "cagr": float(request.form.get("cagr") or 0),
    }

    if request.method == "POST":
        market_name = request.form.get("market_name", "").strip()
        if not market_name:
            return render_template("index.html", error="Market name is required!")

        kmi_data = request.form.get("kmi_data", "").strip()
        if kmi_data:
            data["kmi"].extend([(km.title().strip(), 1) for km in kmi_data.splitlines() if km.strip()])
        print(data["kmi"])
        
        segment_data = request.form.get("segment_data", "").strip()
        
        if not segment_data:
            headings = request.form.getlist("headings[]")
            levels = request.form.getlist("levels[]")
            for heading, level in zip(headings, levels):
                level = int(level)
                if level == 0:
                    raw_segments.append(heading.title())
                    toc_heading = f"Global {market_name} Market Size by {heading.title()} & CAGR (2026-2033)"
                    data["toc_entries"].append((toc_heading, level))
                    data["toc_entries"].append(("Market Overview", 1))
                else:
                    data["toc_entries"].append((heading.title(), level))
        else:
            for seg in segment_data.splitlines():
                seg_level = get_level1(seg) - 1
                cleaned = clean1(seg)
                if seg_level == 0:
                    raw_segments.append(cleaned)
                    cleaned_1 = f"Global {market_name} Market Size by {cleaned} & CAGR (2026-2033)"
                    data["toc_entries"].append((cleaned_1, seg_level))
                    data["toc_entries"].append(("Market Overview", 1))
                else:
                    data["toc_entries"].append((cleaned, seg_level))

        company_data = request.form.get("company_data", "").strip()
        if company_data:
            data["companies"].extend([
                ("Competitive Intelligence", 0),
                ("Top 5 Player Comparison", 1),
                ("Market Positioning of Key Players, 2025", 1),
                ("Strategies Adopted by Key Market Players", 1),
                ("Recent Developments in the Market", 1),
                ("Company Market Share Analysis, 2025", 1),
                ("Company Profiles of All Key Players", 1),
                ("Company Details", 2),
                ("Product Portfolio Analysis", 2),
                ("Company's Segmental Share Analysis", 2),
                ("Revenue Y-O-Y Comparison (2023-2025)", 2),
                ("Key Company Profiles", 0),
            ])

            for company in company_data.splitlines():
                company_name = company.strip()
                if company_name:
                    data["companies"].append((company_name, 1))  
                    data["companies"].extend([ 
                        ("Company Overview", 2),
                        ("Business Segment Overview", 2),
                        ("Financial Updates", 2),
                        ("Key Developments", 2),
                    ])


        regions = [
            ("North America", ["US", "Canada"]),
            ("Europe", ["Germany", "Spain", "France", "UK", "Italy", "Rest of Europe"]),
            ("Asia Pacific", ["China", "India", "Japan", "South Korea", "Rest of Asia-Pacific"]),
            ("Latin America", ["Mexico","Brazil", "Rest of Latin America"]),
            ("Middle East & Africa", ["GCC Countries", "South Africa", "Rest of Middle East & Africa"]),
        ]

        data["regions"].append((f"Global {market_name} Market Size & CAGR (2026-2033)", 0))
        segment_text = ", ".join(raw_segments) if raw_segments else "No segments available"
        print(data["segments"])
        for region, subregions in regions:
            data["regions"].append((f"{region} ({segment_text})", 1))
            data["regions"].extend([(subregion, 2) for subregion in subregions])
        toc_content = data["kmi"] + data["toc_entries"] + data["regions"] + data["companies"]+data["conclusion"]

        toc_temp_file_name = f"TOC_{market_name}_SkyQuest.docx"
        rd_temp_file_name = f"RD_{market_name}_SkyQuest.docx"

        toc_temp_file_path = os.path.join(tempfile.gettempdir(), toc_temp_file_name)
        rd_temp_file_path = os.path.join(tempfile.gettempdir(), rd_temp_file_name)

        toc_doc = Document(doc_path)
        for heading, level in toc_content:
            add_bullet_point_text(toc_doc, heading, level)
        toc_doc.save(toc_temp_file_path)

        export_to_word(
            data=data["toc_entries"] + data["segments"],
            market_name=market_name,
            value_2024=market_details["value_2024"],
            currency=market_details["currency"],
            cagr=market_details["cagr"],
            companies=company_data,
            output_path=rd_temp_file_path
        )

        return render_template(
            "index.html",
            file_ready=True,
            toc_file_path=toc_temp_file_path,
            rd_file_path=rd_temp_file_path
        )

    return render_template("index.html", file_ready=False)

@app.route("/download")
def download():
    file_path = request.args.get("file_path")
    
    if not file_path or not os.path.exists(file_path):
        return "Error: The file does not exist. Please generate the document first.", 404
    file_name = os.path.basename(file_path)
    
    return send_file(
        file_path,
        as_attachment=True,
        download_name=file_name, 
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5001)