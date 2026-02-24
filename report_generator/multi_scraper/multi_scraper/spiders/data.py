import scrapy
import json
import os
import re
import random
from urllib.parse import urlparse
from w3lib.html import remove_tags
from parsel import Selector
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from openai import OpenAI
from itertools import cycle
from dotenv import load_dotenv
from concurrent.futures import ThreadPoolExecutor, as_completed

# Load environment variables
load_dotenv()

TOC_NUM_RE = re.compile(r"^\s*\d+(?:\.\d+)*[.)]?\s+")
# Initialize OpenAI client
client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

# Regional Hierarchy Configuration
GPT_REGION_DOMINANCE = ["North America", "Europe", "Asia Pacific"]
REGION_DOMINANCE = GPT_REGION_DOMINANCE + ["Latin America", "Middle East & Africa"]
EUROPE_COUNTRIES = ["Germany", "United Kingdom", "France"]

# Global market inputs storage
CURRENT_MARKET_INPUTS = {}

# Region-to-countries mapping (static)
REGION_COUNTRIES_MAP = {
    "North America": ["United States", "Canada"],
    "Europe": ["Germany", "United Kingdom", "France"],
    "Asia Pacific": ["Japan", "South Korea"],
}

def build_regions_mapping_prompt(market_name):
    """Build prompt for GPT to determine regional dominance for a market"""
    return f"""
You are a market research analyst. For the {market_name}, determine the regional market dominance ranking.

Available regions (in order of priority): {', '.join(GPT_REGION_DOMINANCE)}

Task 1: Rank the top 3 regions by market dominance for {market_name}
Return ONLY a JSON object in this exact format (no markdown, no explanation):
{{
  "first": "region name",
  "second": "region name", 
  "third": "region name"
}}

IMPORTANT: You MUST choose 3 different regions from: {', '.join(GPT_REGION_DOMINANCE)}
Do not repeat regions.
"""

def build_europe_countries_prompt(market_name):
    """Build prompt for GPT to determine country status in Europe"""
    return f"""
You are a market research analyst. For the {market_name} in Europe, classify each country's market status.

Countries to classify: {', '.join(EUROPE_COUNTRIES)}

For each country, determine if it is: "dominant", "fastest_growing", or "emerging" in the {market_name}.

Return ONLY a JSON object in this exact format (no markdown, no explanation):
{{
  "Germany": "status",
  "United Kingdom": "status",
  "France": "status"
}}

Where status is one of: "dominant", "fastest_growing", "emerging"
"""

def extract_segments_from_toc(table_of_contents):
    """Extract segments and sub-segments from table of contents.
    
    Parses numbered format like:
    "1. Segment Name"
    "1.1. Sub-segment Name"
    "1.2. Sub-segment Name"
    "2. Another Segment"
    
    Returns: {
        "Segment Name": ["Sub-segment 1", "Sub-segment 2"],
        "Another Segment": ["Sub-segment 1"]
    }
    
    Note: Segments and sub-segments are converted to title-case.
    """
    segments = {}
    current_segment = None
    
    if not isinstance(table_of_contents, list):
        return segments
    
    for line in table_of_contents:
        if not isinstance(line, str):
            continue
        
        # Remove numbering prefix like "1. ", "1.1. ", etc.
        match = re.match(r'^\s*(\d+(?:\.\d+)*)[.)]?\s+(.+)$', line.strip())
        if not match:
            continue
        
        numbering = match.group(1)
        text = match.group(2).strip()
        
        # Convert to title-case
        text = text.title()
        
        # Determine hierarchy level by counting dots
        depth = numbering.count('.')
        
        if depth == 0:
            # This is a top-level segment
            current_segment = text
            if current_segment not in segments:
                segments[current_segment] = []
        elif depth == 1 and current_segment is not None:
            # This is a sub-segment, add to current segment
            segments[current_segment].append(text)
    
    # Remove segments with no sub-segments (they're just headers without actual subdivisions)
    segments = {k: v for k, v in segments.items() if v}
    
    return segments

def save_dominating_regions_json(market_name, regions_data, europe_data, market_inputs=None, segments_data=None):
    """Save regional ranking and Europe classification to JSON file"""
    # Create dominating_region folder if it doesn't exist
    folder_path = "dominating_region"
    os.makedirs(folder_path, exist_ok=True)
    
    # Extract market inputs or use defaults
    if market_inputs is None:
        market_inputs = {}
    
    unit = market_inputs.get("unit", "").title() if market_inputs.get("unit", "") else ""
    cagr = market_inputs.get("cagr", "")
    value_2024 = market_inputs.get("value_2024", "")
    value_2033 = market_inputs.get("value_2033", "")
    
    # Prepare data structure
    output_data = {
        "market_name": market_name,
        "unit": unit,
        "cagr": cagr,
        "value_2024": value_2024,
        "value_2033": value_2033,
        "REGIONAL_RANKING": regions_data,
        "EUROPE_COUNTRY_CLASSIFICATION": europe_data
    }
    
    # Add SEGMENTS if provided
    if segments_data:
        output_data["SEGMENTS"] = segments_data
    
    # Create filename from market name
    cleaned_filename = re.sub(r'[<>:"/\\|?*]', '', market_name).strip()[:50]
    file_path = os.path.join(folder_path, f"{cleaned_filename}_dominating_regions.json")
    
    # Save to JSON file
    with open(file_path, "w", encoding="utf-8") as f:
        json.dump(output_data, f, indent=2, ensure_ascii=False)
    
    print(f"\n✅ Saved Dominating Regions JSON: {file_path}\n")
    return file_path


def get_regions_mapping_from_gpt(market_name, market_inputs=None, segments_data=None):
    """Dynamically build REGIONS_MAPPING based on GPT analysis of market"""
    global CURRENT_MARKET_INPUTS
    
    # Set market inputs if provided
    if market_inputs:
        CURRENT_MARKET_INPUTS = market_inputs
    
    try:
        # Get regional ranking from GPT
        prompt_regions = build_regions_mapping_prompt(market_name)
        response_regions = client.chat.completions.create(
            model="gpt-5-mini",
            messages=[{"role": "user", "content": prompt_regions}],
        )
        regions_data = json.loads(response_regions.choices[0].message.content.strip())
        
        # Print first GPT response
        print("\n" + "="*70)
        print("🤖 GPT RESPONSE #1: REGIONAL RANKING")
        print("="*70)
        print(json.dumps(regions_data, indent=2))
        print("="*70 + "\n")
        
        # Get Europe country classification from GPT
        prompt_europe = build_europe_countries_prompt(market_name)
        response_europe = client.chat.completions.create(
            model="gpt-5-mini",
            messages=[{"role": "user", "content": prompt_europe}],
        )
        europe_data = json.loads(response_europe.choices[0].message.content.strip())
        
        # Print second GPT response
        print("\n" + "="*70)
        print("🤖 GPT RESPONSE #2: EUROPE COUNTRY CLASSIFICATION")
        print("="*70)
        print(json.dumps(europe_data, indent=2))
        print("="*70 + "\n")
        
        # Save the GPT responses to JSON with market metrics from global variable
        save_dominating_regions_json(
            market_name, 
            regions_data, 
            europe_data,
            market_inputs=CURRENT_MARKET_INPUTS,
            segments_data=segments_data
        )
        
        # Build REGIONS_MAPPING dynamically
        regions_mapping = {}
         
        # First region (dominant)
        first_region = regions_data.get("first", "North America")
        regions_mapping["dominant"] = {
            "region": first_region,
            "countries": REGION_COUNTRIES_MAP.get(first_region, [])
        }
        
        # Second region (second)
        second_region = regions_data.get("second", "Europe")
        if second_region == "Europe":
            # Add country classification for Europe
            countries_with_types = [
                {"name": country, "type": europe_data.get(country, "emerging")}
                for country in EUROPE_COUNTRIES
            ]
            regions_mapping["second"] = {
                "region": second_region,
                "countries": countries_with_types
            }
        else:
            regions_mapping["second"] = {
                "region": second_region,
                "countries": REGION_COUNTRIES_MAP.get(second_region, [])
            }
        
        # Third region (third)
        third_region = regions_data.get("third", "Asia Pacific")
        regions_mapping["third"] = {
            "region": third_region,
            "countries": REGION_COUNTRIES_MAP.get(third_region, [])
        }
        
        return regions_mapping
        
    except Exception as e:
        print(f"Error building REGIONS_MAPPING from GPT: {e}")
        # Fallback to default mapping
        return {
            "dominant": {
                "region": "North America",
                "countries": ["United States", "Canada"]
            },
            "second": {
                "region": "Europe",
                "countries": [
                    {"name": "Germany", "type": "dominant"},
                    {"name": "United Kingdom", "type": "fastest_growing"},
                    {"name": "France", "type": "emerging"}
                ]
            },
            "third": {
                "region": "Asia Pacific",
                "countries": ["Japan", "South Korea"]
            }
        }

# Initialize default REGIONS_MAPPING (will be overridden dynamically per market)
REGIONS_MAPPING = {
    "dominant": {
        "region": "North America",
        "countries": ["United States", "Canada"]
    },
    "second": {
        "region": "Europe",
        "countries": [
            {"name": "Germany", "type": "dominant"},
            {"name": "United Kingdom", "type": "fastest_growing"},
            {"name": "France", "type": "emerging"}
        ]
    },
    "third": {
        "region": "Asia Pacific",
        "countries": ["Japan", "South Korea"]
    }
}

# Cascading dropdown data structure for industry classification
INDUSTRY_CLASSIFICATION = {
    "sectors": [
        {
            "name": "Materials",
            "id": "materials",
            "industry_groups": [
                {
                    "name": "Diversified Materials",
                    "id": "diversified_materials",
                    "industries": [
                        {
                            "name": "Metals & Mining",
                            "id": "metals_mining",
                            "sub_industries": [
                                "Silver",
                                "Diversified Metals & Mining",
                                "Gold",
                                "Copper",
                                "Steel",
                                "Aluminum",
                                "Precious Metals & Minerals"
                            ]
                        },
                        {
                            "name": "Chemicals",
                            "id": "chemicals",
                            "sub_industries": [
                                "Commodity Chemicals",
                                "Fertilizers & Agricultural Chemicals",
                                "Specialty Chemicals",
                                "Industrial Gases",
                                "Diversified Chemicals"
                            ]
                        },
                        {
                            "name": "Paper & Forest Products",
                            "id": "paper_forest",
                            "sub_industries": [
                                "Paper Products",
                                "Forest Products"
                            ]
                        },
                        {
                            "name": "Construction Materials",
                            "id": "construction_materials",
                            "sub_industries": [
                                "Construction Materials"
                            ]
                        },
                        {
                            "name": "Containers & Packaging",
                            "id": "containers_packaging",
                            "sub_industries": [
                                "Paper Packaging",
                                "Metal & Glass Containers"
                            ]
                        }
                    ]
                }
            ]
        },
        {
            "name": "Utilities",
            "id": "utilities",
            "industry_groups": [
                {
                    "name": "Diversified Utilities Services",
                    "id": "diversified_utilities",
                    "industries": [
                        {
                            "name": "Electric Utilities",
                            "id": "electric_utilities",
                            "sub_industries": [
                                "Electric Utilities"
                            ]
                        },
                        {
                            "name": "Gas Utilities",
                            "id": "gas_utilities",
                            "sub_industries": [
                                "Gas Utilities"
                            ]
                        },
                        {
                            "name": "Multi-Utilities",
                            "id": "multi_utilities",
                            "sub_industries": [
                                "Multi-Utilities"
                            ]
                        },
                        {
                            "name": "Water Utilities",
                            "id": "water_utilities",
                            "sub_industries": [
                                "Water Utilities"
                            ]
                        },
                        {
                            "name": "Independent Power and Renewable Electricity Producers",
                            "id": "independent_power_renewable",
                            "sub_industries": [
                                "Independent Power Producers & Energy Traders",
                                "Renewable Electricity"
                            ]
                        }
                    ]
                }
            ]
        },
        {
            "name": "Real Estate",
            "id": "real_estate",
            "industry_groups": [
                {
                    "name": "Real Estate Services",
                    "id": "real_estate_services",
                    "industries": [
                        {
                            "name": "Equity Real Estate Investment Trusts (REITs)",
                            "id": "equity_reits",
                            "sub_industries": [
                                "Diversified REITs",
                                "Industrial REITs",
                                "Office REITs",
                                "Health Care REITs",
                                "Retail REITs",
                                "Specialized REITs",
                                "Hotel & Resort REITs",
                                "Residential REITs"
                            ]
                        },
                        {
                            "name": "Real Estate Management & Development",
                            "id": "real_estate_mgmt_dev",
                            "sub_industries": [
                                "Real Estate Operating Companies",
                                "Real Estate Development",
                                "Real Estate Services",
                                "Diversified Real Estate Activities"
                            ]
                        }
                    ]
                }
            ]
        },
        {
            "name": "Consumer Staples",
            "id": "consumer_staples",
            "industry_groups": [
                {
                    "name": "Food & Staples Retailing",
                    "id": "food_staples_retailing",
                    "industries": [
                        {
                            "name": "Food & Staples Retailing",
                            "id": "Food & Staples Retailing",
                            "sub_industries": [
                                "Drug Retail",
                                "Food Retail",
                                "Hypermarkets & Super Centers",
                                "Food Distributors"
                            ]
                        }
                    ]
                },
                {
                    "name": "Food, Beverage & Tobacco",
                    "id": "food_beverage_tobacco",
                    "industries": [
                        {
                            "name": "Food Products",
                            "id": "food_products",
                            "sub_industries": [
                                "Agricultural products",
                                "Packaged Foods & Meats",
                            ]
                        },
                        {
                            "name": "Beverages",
                            "id": "beverages",
                            "sub_industries": [
                                "Brewers",
                                "Soft Drinks",
                                "Distillers & Vintners"
                            ]
                        },
                        {
                            "name": "Tobacco",
                            "id": "tobacco",
                            "sub_industries": [
                                "Tobacco Products"
                            ]
                        }
                    ]
                },
                {
                    "name": "Household & Personal Products",
                    "id": "household_personal_products",
                    "industries": [
                        {
                            "name": "Household Products",
                            "id": "household_products",
                            "sub_industries": [
                                "Household"
                            ]
                        },
                        {
                            "name": "Personal Products",
                            "id": "personal_products",
                            "sub_industries": [
                                "Personal Products"
                            ]
                        }
                    ]
                }
            ]
        },
        {
            "name": "Health Care",
            "id": "health_care",
            "industry_groups": [
                {
                    "name": "Health Care Equipment & Services",
                    "id": "health_care_equipment_services",
                    "industries": [
                        {
                            "name": "Health Care Equipment & Supplies",
                            "id": "health_care_equipment_supplies",
                            "sub_industries": [
                                "Health Care Equipment",
                                "Health Care Supplies"
                            ]
                        },
                        {
                            "name": "Health Care Providers & Services",
                            "id": "health_care_providers_services",
                            "sub_industries": [
                                "Health Care Distributors",
                                "Health Care Facilities",
                                "Managed Health Care",
                                "Health Care Services"
                            ]
                        },
                        {
                            "name": "Health Care Technology",
                            "id": "health_care_technology",
                            "sub_industries": [
                                "Health Care Technology"
                            ]
                        }
                    ]
                },
                {
                    "name": "Pharmaceuticals, Biotechnology & Life Sciences",
                    "id": "pharma_biotech_life_sciences",
                    "industries": [
                        {
                            "name": "Biotechnology",
                            "id": "biotechnology",
                            "sub_industries": [
                                "Biotechnology"
                            ]
                        },
                        {
                            "name": "Pharmaceuticals",
                            "id": "pharmaceuticals",
                            "sub_industries": [
                                "Pharmaceuticals"
                            ]
                        },
                        {
                            "name": "Life Sciences Tools & Services",
                            "id": "life_sciences_tools_services",
                            "sub_industries": [
                                "Life Sciences Tools & Services"
                            ]
                        }
                    ]
                }
            ]
        },
        {
            "name": "Financials",
            "id": "financials",
            "industry_groups": [
                {
                    "name": "Insurance",
                    "id": "insurance",
                    "industries": [
                        {
                            "name": "Insurance Services",
                            "id": "insurance_services",
                            "sub_industries": [
                                "Reinsurance",
                                "Insurance Brokers",
                                "Multi-line Insurance",
                                "Life & Health Insurance",
                                "Property & Casualty Insurance"
                            ]
                        }
                    ]
                },
                {
                    "name": "Diversified Financials",
                    "id": "diversified_financials",
                    "industries": [
                        {
                            "name": "Consumer Finance",
                            "id": "consumer_finance",
                            "sub_industries": [
                                "Consumer Finance"
                            ]
                        },
                        {
                            "name": "Diversified Financial Services",
                            "id": "diversified_financial_services",
                            "sub_industries": [
                                "Multi-Sector Holdings",
                                "Other Diversified Financial Services",
                                "Specialized Finance"
                            ]
                        },
                        {
                            "name": "Capital Markets",
                            "id": "capital_markets",
                            "sub_industries": [
                                "Financial Exchanges & Data",
                                "Investment Banking & Brokerage",
                                "Diversified Capital Markets",
                                "Asset Management & Custody Banks"
                            ]
                        },
                        {
                            "name": "Mortgage Real Estate Investment Trusts (REITs)",
                            "id": "mortgage_reits",
                            "sub_industries": [
                                "Mortgage REITs"
                            ]
                        }
                    ]
                },
                {
                    "name": "Banks",
                    "id": "banks",
                    "industries": [
                        {
                            "name": "Banking Services",
                            "id": "banking_services",
                            "sub_industries": [
                                "Regional Banks",
                                "Diversified Banks"
                            ]
                        },
                        {
                            "name": "Thrifts & Mortgage Finance",
                            "id": "thrifts_mortgage_finance",
                            "sub_industries": [
                                "Thrifts & Mortgage Finance"
                            ]
                        }
                    ]
                }
            ]
        },
        {
            "name": "Industrials",
            "id": "industrials",
            "industry_groups": [
                {
                    "name": "Capital Goods",
                    "id": "capital_goods",
                    "industries": [
                        {
                            "name": "Building Products",
                            "id": "building_products",
                            "sub_industries": [
                                "Building Products"
                            ]
                        },
                        {
                            "name": "Industrial Conglomerates",
                            "id": "industrial_conglomerates",
                            "sub_industries": [
                                "Industrial Conglomerates"
                            ]
                        },
                        {
                            "name": "Machinery",
                            "id": "machinery",
                            "sub_industries": [
                                "Industrial Machinery",
                                "Construction Machinery & Heavy Trucks",
                                "Agricultural & Farm Machinery"
                            ]
                        },
                        {
                            "name": "Aerospace & Defense",
                            "id": "aerospace_defense",
                            "sub_industries": [
                                "Aerospace & Defense"
                            ]
                        },
                        {
                            "name": "Construction & Engineering",
                            "id": "construction_engineering",
                            "sub_industries": [
                                "Construction & Engineering"
                            ]
                        },
                        {
                            "name": "Electrical Equipment",
                            "id": "electrical_equipment",
                            "sub_industries": [
                                "Electrical Components & Equipment",
                                "Heavy Electrical Equipment"
                            ]
                        },
                        {
                            "name": "Trading Companies & Distributors",
                            "id": "trading_companies_distributors",
                            "sub_industries": [
                                "Trading Companies & Distributors"
                            ]
                        }
                    ]
                },
                {
                    "name": "Transportation",
                    "id": "transportation",
                    "industries": [
                        {
                            "name": "Airlines",
                            "id": "airlines",
                            "sub_industries": [
                                "Airlines"
                            ]
                        },
                        {
                            "name": "Air Freight & Logistics",
                            "id": "air_freight_logistics",
                            "sub_industries": [
                                "Air Freight & Logistics"
                            ]
                        },
                        {
                            "name": "Marine",
                            "id": "marine",
                            "sub_industries": [
                                "Marine"
                            ]
                        },
                        {
                            "name": "Road & Rail",
                            "id": "road_rail",
                            "sub_industries": [
                                "Railroads",
                                "Trucking"
                            ]
                        },
                        {
                            "name": "Transportation Infrastructure",
                            "id": "transportation_infrastructure",
                            "sub_industries": [
                                "Marine Ports & Services",
                                "Highways & Railtracks",
                                "Airport Services"
                            ]
                        }
                    ]
                }
            ]
        },
        {
            "name": "Information Technology",
            "id": "information_technology",
            "industry_groups": [
                {
                    "name": "Semiconductors & Semiconductor Equipment",
                    "id": "semiconductors_semiconductor_equipment",
                    "industries": [
                        {
                            "name": "Semiconductor Services & Semiconductor Equipment",
                            "id": "semiconductor_services_equipment",
                            "sub_industries": [
                                "Semiconductor Equipment",
                                "Semiconductors"
                            ]
                        }
                    ]
                },
                {
                    "name": "Technology Hardware & Equipment",
                    "id": "technology_hardware_equipment",
                    "industries": [
                        {
                            "name": "Communications Equipment",
                            "id": "communications_equipment",
                            "sub_industries": [
                                "Communications Equipment"
                            ]
                        },
                        {
                            "name": "Technology Hardware, Storage & Peripherals",
                            "id": "technology_hardware_storage_peripherals",
                            "sub_industries": [
                                "Technology Hardware, Storage & Peripherals"
                            ]
                        },
                        {
                            "name": "Electronic Equipment, Instruments & Components",
                            "id": "electronic_equipment_instruments_components",
                            "sub_industries": [
                                "Electronic Components",
                                "Electronic Equipment & Instruments",
                                "Technology Distributors",
                                "Electronic Manufacturing Services"
                            ]
                        }
                    ]
                },
                {
                    "name": "Software & Services",
                    "id": "software_services",
                    "industries": [
                        {
                            "name": "IT Services",
                            "id": "it_services",
                            "sub_industries": [
                                "Data Processing & Outsourced Services",
                                "Internet Services & Infrastructure",
                                "IT Consulting & Other Services"
                            ]
                        },
                        {
                            "name": "Software",
                            "id": "software",
                            "sub_industries": [
                                "Systems Software",
                                "Application Software",
                                "Home Entertainment Software"
                            ]
                        },
                        {
                            "name": "Internet Software & Services",
                            "id": "internet_software_services",
                            "sub_industries": [
                                "Internet Software & Services"
                            ]
                        }
                    ]
                }
            ]
        },
        {
            "name": "Communication Services",
            "id": "communication_services",
            "industry_groups": [
                {
                    "name": "Media & Entertainment",
                    "id": "media_entertainment",
                    "industries": [
                        {
                            "name": "Media",
                            "id": "media",
                            "sub_industries": [
                                "Cable & Satellite",
                                "Broadcasting",
                                "Publishing",
                                "Advertising"
                            ]
                        },
                        {
                            "name": "Entertainment",
                            "id": "entertainment",
                            "sub_industries": [
                                "Movies & Entertainment",
                                "Interactive Home Entertainment"
                            ]
                        },
                        {
                            "name": "Interactive Media & Services",
                            "id": "interactive_media_services",
                            "sub_industries": [
                                "Interactive Media & Services"
                            ]
                        }
                    ]
                },
                {
                    "name": "Telecommunication Services",
                    "id": "telecommunication_services",
                    "industries": [
                        {
                            "name": "Diversified Telecommunication Services",
                            "id": "diversified_telecommunication_services",
                            "sub_industries": [
                                "Alternative Carriers",
                                "Integrated Telecommunication Services"
                            ]
                        },
                        {
                            "name": "Wireless Telecommunication Services",
                            "id": "wireless_telecommunication_services",
                            "sub_industries": [
                                "Wireless Telecommunication Services"
                            ]
                        }
                    ]
                }
            ]
        },
        {
            "name": "Energy",
            "id": "energy",
            "industry_groups": [
                {
                    "name": "Energy Fuel",
                    "id": "energy_fuel",
                    "industries": [
                        {
                            "name": "Oil, Gas & Consumable Fuels",
                            "id": "oil_gas_consumable_fuels",
                            "sub_industries": [
                                "Oil & Gas Refining & Marketing",
                                "Integrated Oil & Gas",
                                "Oil & Gas Storage & Transportation",
                                "Coal & Consumable Fuels",
                                "Oil & Gas Exploration & Production"
                            ]
                        }
                    ]
                },
                {
                    "name": "Energy Service",
                    "id": "energy_service",
                    "industries": [
                        {
                            "name": "Energy Equipment & Services",
                            "id": "energy_equipment_services",
                            "sub_industries": [
                                "Oil & Gas Equipment & Services",
                                "Oil & Gas Drilling"
                            ]
                        }
                    ]
                }
            ]
        },
        {
            "name": "Consumer Discretionary",
            "id": "consumer_discretionary",
            "industry_groups": [
                {
                    "name": "Consumer Durables & Apparel",
                    "id": "consumer_durables_apparel",
                    "industries": [
                        {
                            "name": "Household Durables",
                            "id": "household_durables",
                            "sub_industries": [
                                "Consumer Electronics",
                                "Housewares & Specialties",
                                "Homebuilding",
                                "Home Furnishings",
                                "Household Appliances"
                            ]
                        },
                        {
                            "name": "Leisure Products",
                            "id": "leisure_products",
                            "sub_industries": [
                                "Leisure Products"
                            ]
                        },
                        {
                            "name": "Textiles, Apparel & Luxury Goods",
                            "id": "textiles_apparel_luxury_goods",
                            "sub_industries": [
                                "Apparel, Accessories & Luxury Goods",
                                "Footwear",
                                "Textiles"
                            ]
                        }
                    ]
                },
                {
                    "name": "Consumer Services",
                    "id": "consumer_services",
                    "industries": [
                        {
                            "name": "Hotels, Restaurants & Leisure",
                            "id": "hotels_restaurants_leisure",
                            "sub_industries": [
                                "Casinos & Gaming",
                                "Hotels, Resorts & Cruise Lines",
                                "Leisure Facilities",
                                "Restaurants"
                            ]
                        },
                        {
                            "name": "Diversified Consumer Services",
                            "id": "diversified_consumer_services",
                            "sub_industries": [
                                "Education Services",
                                "Specialized Consumer Services"
                            ]
                        }
                    ]
                },
                {
                    "name": "Automobiles & Components",
                    "id": "automobiles_components",
                    "industries": [
                        {
                            "name": "Auto Components",
                            "id": "auto_components",
                            "sub_industries": [
                                "Tires & Rubber",
                                "Auto Parts & Equipment"
                            ]
                        },
                        {
                            "name": "Automobiles",
                            "id": "automobiles",
                            "sub_industries": [
                                "Motorcycle Manufacturers",
                                "Automobile Manufacturers"
                            ]
                        }
                    ]
                },
                {
                    "name": "Retailing",
                    "id": "retailing",
                    "industries": [
                        {
                            "name": "Multiline Retail",
                            "id": "multiline_retail",
                            "sub_industries": [
                                "Department Stores",
                                "General Merchandise Stores"
                            ]
                        },
                        {
                            "name": "Specialty Retail",
                            "id": "specialty_retail",
                            "sub_industries": [
                                "Apparel Retail",
                                "Homefurnishing Retail",
                                "Automotive Retail",
                                "Specialty Stores",
                                "Computer & Electronics Retail",
                                "Home Improvement Retail"
                            ]
                        },
                        {
                            "name": "Distributors",
                            "id": "distributors",
                            "sub_industries": [
                                "Distributors"
                            ]
                        },
                        {
                            "name": "Internet & Direct Marketing Retail",
                            "id": "internet_direct_marketing_retail",
                            "sub_industries": [
                                "Internet & Direct Marketing Retail"
                            ]
                        }
                    ]
                },
                {
                    "name": "Media",
                    "id": "media_consumer_discretionary",
                    "industries": [
                        {
                            "name": "Diversified Media Services",
                            "id": "diversified_media_services",
                            "sub_industries": [
                                "Advertising",
                                "Movies & Entertainment",
                                "Publishing",
                                "Cable & Satellite",
                                "Broadcasting"
                            ]
                        }
                    ]
                }
            ]
        }
    ]
}

class MarketResearchSpider(scrapy.Spider):
    name = "data"
    SKYQUEST_TOC_AUTOMAP = {
        "browserHtml": True,
        "actions": [
            {
                "action": "click",
                "selector": {
                    "type": "css",
                    "state": "visible",
                    "value": 'a[href="#tab_default_3"]'
                },
                "delay": 0,
                "button": "left",
                "onError": "return"
            }
        ],
    }
 
    def __init__(self, urls=None, *args, **kwargs):
        super(MarketResearchSpider, self).__init__(*args, **kwargs)
     
        # Get URLs from command line argument
        if urls:
            self.start_urls = [url.strip() for url in urls.split(',') if url.strip()]
        else:
            self.start_urls = []
     
        # If no URLs provided, ask for input
        if not self.start_urls:
            input_urls = input("Enter URLs separated by comma: ")
            self.start_urls = [url.strip() for url in input_urls.split(',') if url.strip()]
     
        if not self.start_urls:
            self.logger.warning("🚫 No URLs provided. Exiting.")
            self.crawler.engine.close_spider(self, reason="no_urls")
     
        self.output_dir = "scraped_json"
        os.makedirs(self.output_dir, exist_ok=True)
     
        self.ws_re = re.compile(r"\s+")
        self.num_dot_fix_re = re.compile(r"^((?:\d+\.)*\d+)(?!\.)\s+")
        self.no_docx = kwargs.pop('no_docx', False)

    def sanitize_filename(self, text):
        return re.sub(r'[^\w\-]', '_', text).lower()
    
    def extract_domain_from_url(self, url):
        """Extract domain name from URL for unique file naming"""
        try:
            from urllib.parse import urlparse
            parsed = urlparse(url)
            domain = parsed.netloc.replace('www.', '').split('.')[0]
            return domain.lower()
        except Exception as e:
            self.logger.warning(f"Could not extract domain from {url}: {e}")
            return "unknown"

 
    def start_requests(self):
        if not self.start_urls:
            self.logger.warning("🚫 No URLs provided. Exiting.")
            return
     
        for url in self.start_urls:
            if "grandviewresearch.com" in url:
                # Add "/segmentation" to the URL if it's not already there
                if not url.endswith("/segmentation"):
                    segmentation_url = f"{url}/segmentation"
                else:
                    segmentation_url = url
                meta = {'original_url': url.replace('/segmentation', '')}
                yield scrapy.Request(segmentation_url, callback=self.parse_grandview, meta=meta)
            elif "marketsandmarkets.com" in url:
                yield scrapy.Request(url, callback=self.parse_markets)
            elif "snsinsider.com" in url:
                yield scrapy.Request(url, callback=self.parse_sns)
            elif "mordorintelligence.com" in url:
                yield scrapy.Request(url, callback=self.parse_mordor, meta={"zyte_api_automap": True})
            elif "fortunebusinessinsights.com" in url:## modified ##
                yield scrapy.Request(
                    url,
                    callback=self.parse_fortune_main,
                    meta={
                        "zyte_api_automap": True,
                        "zyte_api_render": {"browserHtml": True}
                    }
                )
            elif "futuremarketinsights.com" in url: ## modified ##
                yield scrapy.Request(
                    url, 
                    callback=self.parse_future,
                    meta={
                        "zyte_api_automap": {"browserHtml": True}
                    }
                )
            elif "alliedmarketresearch.com" in url:
                yield scrapy.Request(url, callback=self.parse_allied, meta={"zyte_api": {"browserHtml": True}})
            elif "skyquestt.com" in url:
                yield scrapy.Request(
                    url,
                    callback=self.parse_skyquestt,
                    meta={
                        "zyte_api_automap": self.SKYQUEST_TOC_AUTOMAP
                    },
                    dont_filter=True
                )
            elif "marketresearchfuture.com" in url: ## Modified ##
                yield scrapy.Request(
                    url,
                    callback=self.parse_mrf,
                    meta={
                        "zyte_api": {
                            "browserHtml": True
                        }
                    }
                )          
# deb 1st edit start                      
            elif "technavio.com" in url:
                yield scrapy.Request(
                    url,
                    callback=self.parse_technavio,
                    meta={"zyte_api_automap": True}
                )
            elif "straitsresearch.com" in url:
                if not url.endswith("/segmentation"):
                    segmentation_url = url.rstrip("/") + "/segmentation"
                else:
                    segmentation_url = url

                yield scrapy.Request(
                    segmentation_url,
                    callback=self.parse_straits,
                    meta={"main_url": url.replace("/segmentation", "")}
                ) 
#deb 1st edit stop
            elif "precedenceresearch.com" in url:   # Modified #rd
                yield scrapy.Request(
                    url=url,
                    callback=self.parse_toc
                )
            elif "verifiedmarketresearch.com" in url:
                yield scrapy.Request(
                    url,
                    callback=self.parse_verified,
                    meta={
                        "zyte_api_automap": self.VERIFIED_TOC_AUTOMAP
                    },
                    dont_filter=True
                )  
            elif "verifiedmarketreports.com" in url:
                yield scrapy.Request(
                    url,
                    callback=self.parse_verified_reports,
                    meta={
                        "zyte_api": {
                            "browserHtml": True
                        }
                    }
                )
            elif "gminsights.com" in url:
                yield scrapy.Request(
                    url,
                    callback=self.parse_global_market,
                    meta={
                        "zyte_api": {
                            "browserHtml": True
                        }
                    }
                )

            else:
                self.logger.warning(f"Unsupported domain in URL: {url}")
            
 
    # ==================== GRANDVIEW RESEARCH ====================
    def parse_grandview(self, response):
        summary = response.css('div.report_summary.full.non-indexable ul')
     
        # Extract company profiles from the main report page (not segmentation page)
        company_profiles = []
        original_url = response.meta.get('original_url', response.url.replace('/segmentation', ''))
     
        # Only fetch company profiles if we have a different URL than the current one
        if original_url != response.url:
            yield scrapy.Request(original_url,
                               callback=self.parse_grandview_company_profiles,
                               meta={'response_data': response, 'original_url': original_url})
            return
     
        # If we're already on the main page, process normally
        self.process_grandview_data(response, summary, company_profiles)
 
    def parse_grandview_company_profiles(self, response):
        """Parse company profiles from the main report page"""
        response_data = response.meta['response_data']
        original_url = response.meta['original_url']
     
        # Extract company profiles
        company_profiles = self.extract_grandview_company_profiles(response)
     
        # Process the final data with company profiles
        self.process_grandview_data(response_data, response_data.css('div.report_summary.full.non-indexable ul'), company_profiles)
 
    def process_grandview_data(self, response, summary, company_profiles):
        """Process and save the final data"""
        ## Modified Starts Here ##
        if summary:
            def parse_and_format(ul, prefix=[]):
                lines = []
                for idx, li in enumerate(ul.css(':scope > li'), 1):
                    # ---- extract name ----
                    strong_texts = li.css('strong::text, strong *::text').getall()
                    if strong_texts:
                        name = ' '.join(s.strip() for s in strong_texts if s.strip())
                    else:
                        direct_texts = li.xpath('text()').getall()
                        name = ' '.join(t.strip() for t in direct_texts if t.strip())
                    if not name:
                        continue

                    current_prefix = prefix + [str(idx)]
                    number = ".".join(current_prefix)
                    cleaned_name = self.clean_toc_entry(name)
                    if self.is_regional_section(cleaned_name):
                        continue
                    depth = len(current_prefix) ## modification ##
                    # Skip sub / sub-sub with >6 words
                    if depth >= 2 and len(cleaned_name.split()) > 6:
                        continue
                    # Add segment & sub-segment
                    lines.append(f"{number}. {cleaned_name}")
                    # ---- handle children ----
                    for child_ul in li.css(':scope > ul'):

                        # CASE: sub-sub segments (depth == 2)
                        if len(current_prefix) == 2:
                            sub_sub_lis = child_ul.css(':scope > li')
                            # append ONLY if count > 2
                            if len(sub_sub_lis) >= 2:
                                child_lines = parse_and_format(child_ul, current_prefix)
                                lines.extend([
                                    line for line in child_lines
                                    if not self.is_regional_line(line)
                                ])
                        # CASE: segment → sub-segment (depth < 2)
                        elif len(current_prefix) < 2:
                            child_lines = parse_and_format(child_ul, current_prefix)
                            lines.extend([
                                line for line in child_lines
                                if not self.is_regional_line(line)
                            ])

                return lines
            numbered_lines = parse_and_format(summary[0])
            # Filter out any remaining regional lines that might have slipped through
            filtered_lines = [line for line in numbered_lines if not self.is_regional_line(line)]
            # Clean the title
            raw_title = response.css('title::text').get()
            cleaned_title = self.clean_title(raw_title)
            self.market_name = cleaned_title.replace(" Market", "").strip()
            # Clean all TOC entries
            cleaned_lines = [self.clean_toc_line(line) for line in filtered_lines]
            
            # Extract segments from cleaned_lines
            segments = extract_segments_from_toc(cleaned_lines)
            
            # Prepare the data with both table_of_contents and company_profiles
            data = {
                'title': cleaned_title,
                'url': response.meta.get('original_url', response.url),
                'table_of_contents': cleaned_lines
            }
            
            # Add segments if available
            if segments:
                data['segments'] = segments
            
            # Add company profiles if available
            if company_profiles:
                data['company_profiles'] = company_profiles
            self.save_to_json(data, response.meta.get('original_url', response.url))
        else:
            self.logger.error("Could not find the segmentation <ul>!")
            self.logger.info(f"Page title: {response.css('title::text').get()}")
            # Save error response for debugging
            self.save_to_json({
                'title': response.css('title::text').get(),
                'url': response.meta.get('original_url', response.url),
                'table_of_contents': ["Could not find the segmentation <ul>!"],
                'error': "TOC structure not found"
            }, response.meta.get('original_url', response.url))
 
    def extract_grandview_company_profiles(self, response):
        """Extract company names from the company profiles section"""
        company_profiles = []
        selectors = [
            'div.report_external_com_sec',
            'div[class*="company"]',
            'div[class*="profile"]',
            'div[class*="key"]',
            'div[class*="player"]'
        ]

        for selector in selectors:
            company_section = response.css(selector)
            if company_section:
                # Try to find direct li text first
                companies = company_section.css('ul li::text').getall()
                # If li text empty, try li with p inside
                if not companies:
                    for ul in company_section.css('ul'):
                        li_texts = ul.css('li p::text, li::text').getall()
                        companies.extend(li_texts)

                # If still empty, fallback to all text filtering
                if not companies:
                    all_text = company_section.css('*::text').getall()
                    companies = [text for text in all_text if self.is_likely_company_name(text)]
                # Clean and append
                for company in companies:
                    cleaned_company = company.strip()
                    if (self.is_likely_company_name(cleaned_company) and
                        cleaned_company not in company_profiles):
                        company_profiles.append(cleaned_company)

                if company_profiles:
                    break

        # ===========================
        # 2. Heading + sibling <ul> approach
        # ===========================
        keywords = [
            "key companies",
            "key company",
            "key players",
            "key player",
            "major players",
            "major player"
        ]

        for heading in response.css("h2, h3"):
            heading_text = " ".join(heading.css("*::text").getall()).strip().lower()
            if any(keyword in heading_text for keyword in keywords):
                # Get next sibling <ul> immediately after heading
                ul = heading.xpath("following-sibling::ul[1]")
                if ul:
                    companies = ul.css("li p::text, li::text").getall()
                    for company in companies:
                        cleaned = company.strip()
                        if cleaned and cleaned not in company_profiles:
                            company_profiles.append(cleaned)

                    if company_profiles:
                        break
        return company_profiles
 
    # ==================== MARKETSANDMARKETS ====================
    ## Modified Starts Here ##
    def parse_markets_accordion(self, response):
        toc_sections = []
        for item in response.css("div.accordion-item"):
            main_title = item.css("div.TOCcustHead div:nth-child(1)::text").get()
            main_title = main_title.strip() if main_title else None
            sub_sections = []
            for bullet_item in item.css("ul.toc_list li"):
                # Extract the main title (bulletsHead)
                main_title_elem = bullet_item.css("div.bulletsHead::text").get()
                if main_title_elem:
                    main_title_text = main_title_elem.strip()
                 
                    # Extract all sub-sub-sections (bullets)
                    sub_sub_sections = []
                    for bullet in bullet_item.css("div.bullets::text"):
                        bullet_text = bullet.get().strip()
                        if bullet_text:
                            sub_sub_sections.append(bullet_text)
                 
                    sub_sections.append({
                        "title": main_title_text,
                        "sub_sub_sections": sub_sub_sections
                    })
                else:
                    # Handle regular list items without bulletsHead
                    lines = [
                        remove_tags(t).strip()
                        for t in bullet_item.css("*::text").getall()
                        if remove_tags(t).strip()
                    ]
                    if lines:
                        sub_sections.append({
                            "title": lines[0],
                            "sub_sub_sections": []
                        })
            toc_sections.append({
                "main_title": main_title,
                "sub_sections": sub_sections,
            })
        formatted_toc = []
        chapter_counter = 0
        for chapter in toc_sections:
            if not chapter["main_title"]:
                continue
            # remove numbering
            chapter_title = remove_tags(re.sub(r"^\d+(\.\d+)*\s*", "", chapter["main_title"])).strip()
            # 🚫 skip if main segment is only explanatory text in ()
            if re.fullmatch(r"\(.*\)", chapter_title):
                continue
            stop_keywords = ["BY REGION", "BY COUNTRY", "BY GEOGRAPHICAL"]
            if "BY" not in chapter_title.upper() or any(keyword in chapter_title.upper() for keyword in stop_keywords):
                continue
            # normalize
            match = re.search(r"BY.*", chapter_title, re.IGNORECASE)
            if not match:
                continue
            chapter_title = match.group(0).split("BY", 1)[1].strip().upper()
            # renumber chapters
            chapter_counter += 1
            formatted_toc.append(f"{chapter_counter}. {chapter_title}")
            # Start sub-section numbering from 1
            section_counter = 1
            for sub in chapter.get("sub_sections", []):
                title = remove_tags(re.sub(r"^\d+(\.\d+)*\s*", "", sub.get("title", "").strip()))
                if (
                    not title
                    or len(title.split()) > 6
                    or re.search(r"\b(primary insights|key primary insights)\b", title.lower())
                    or "introduction" in title.lower()
                ):
                    continue

                formatted_toc.append(f"{chapter_counter}.{section_counter}. {title}")
                # Add sub-sub-sections only if there are 2 or more
                sub_sub_sections = sub.get("sub_sub_sections", [])
                if len(sub_sub_sections) >= 2:
                    sub_sub_counter = 1
                    for sub_sub in sub_sub_sections:
                        if (
                            not sub_sub
                            or len(sub_sub.split()) > 6
                            or re.search(r"\b(primary insights|key primary insights)\b", sub_sub.lower())
                            or "introduction" in sub_sub.lower()
                        ):
                            continue

                        formatted_toc.append(f"{chapter_counter}.{section_counter}.{sub_sub_counter}. {sub_sub}")
                        sub_sub_counter += 1
                section_counter += 1

        # Extract only "Company Profiles" section
        company_profiles = []
        company_section = response.xpath(
            "//div[contains(@class,'TOCcustHead')][.//text()[contains(., 'COMPANY PROFILES')]]"
            "/following-sibling::div[contains(@class,'accordion-item-body')][1]"
        )

        inside_key_players = False
        for li in company_section.css("ul.toc_list > li"):
            head = li.css("div.bulletsHead::text").get()
            head_clean = head.strip().upper() if head else ""
            # Detect KEY / MAJOR PLAYERS
            if "KEY PLAYERS" in head_clean or "MAJOR PLAYERS" in head_clean:
                inside_key_players = True
            # Stop when another section starts (e.g. OTHER PLAYERS)
            elif head_clean and inside_key_players:
                break
            # Collect companies (same li or following li)
            if inside_key_players:
                for company in li.css("div.bullets::text").getall():
                    company = company.strip()
                    if company and company.isupper():  # uppercase filter if desired
                        company_profiles.append(company)

        return formatted_toc, company_profiles
    
    def parse_markets_table(self, response):
        toc_sections = []
        for item in response.css("div.tblTOC div.clsTR"):

            main_title = item.css("div.txthead::text").get()
            main_title = main_title.strip() if main_title else None

            # IGNORE titles fully wrapped in parentheses
            if main_title and re.fullmatch(r"\(.*\)", main_title):
                main_title = None

            sub_sections = []
            sub_title = item.css("div.txtsubhead::text").get()
            sub_title = sub_title.strip() if sub_title and sub_title.strip() != "\xa0" else None
            sub_sections = [sub_title] if sub_title else []

            sub_sub_title = item.css("div:nth-child(4)::text").get()
            sub_sub_title = sub_sub_title.strip() if sub_sub_title else None

            toc_sections.append({
                "main_title": main_title,
                "sub_sections": sub_sections,
                "sub_sub_title": sub_sub_title,
            })

        formatted_toc = []
        chapter_counter = 0
        section_counter = 0
        sub_sub_buffer = []
        current_subsection_key = None
        for chapter in toc_sections:
            main_title = chapter["main_title"]

            # allow empty main_title for sub-sub sections like 9.2.1
            if not main_title and not chapter["sub_sections"]:
                continue
            # CASE 1: MAIN CHAPTER TITLE (only when main_title exists)
            if main_title and not re.fullmatch(r"\d+(\.\d+)?", main_title):
                # 🔧 FIX: flush sub-sub sections before starting new chapter
                if sub_sub_buffer:
                    if len(sub_sub_buffer) >= 2:
                        formatted_toc.extend(sub_sub_buffer)
                    sub_sub_buffer = []
                    current_subsection_key = None

                chapter_title = remove_tags(main_title).strip()
                chapter_title_upper = chapter_title.upper()
                stop_keywords = ["BY REGION", "BY COUNTRY", "BY GEOGRAPHICAL"]

                if any(keyword in chapter_title_upper for keyword in stop_keywords):
                    break

                elif "BY" not in chapter_title_upper:
                    continue

                match = re.search(r"BY.*", chapter_title, re.IGNORECASE)
                if not match:
                    continue

                chapter_name = (match.group(0).split("BY", 1)[1].strip().upper())
                chapter_counter += 1
                section_counter = 0
                formatted_toc.append(f"{chapter_counter}. {chapter_name}")
                continue
            # CASE 2: SUBSECTION (7.2, 7.10)
            if re.fullmatch(r"\d+\.\d+", main_title):
                # Flush previous sub-sub sections
                if sub_sub_buffer:
                    if len(sub_sub_buffer) >= 2:
                        formatted_toc.extend(sub_sub_buffer)
                    sub_sub_buffer = []
                    current_subsection_key = None

                if chapter_counter == 0:
                    continue

                if not chapter["sub_sections"]:
                    continue

                sub_title = chapter["sub_sections"][0]
                if not sub_title:
                    continue

                # ✅ SKIP if title ends with a year (e.g., ", 2024")
                if re.search(r",\s*\d{4}$", sub_title):
                    continue

                if len(sub_title.split()) > 6 or re.search(r"\b(primary insights|key primary insights)\b", sub_title.lower()) or "introduction" in sub_title.lower():
                    continue

                section_counter += 1
                entry = f"{chapter_counter}.{section_counter}. {sub_title}"
                formatted_toc.append(entry)
                continue
            # CASE 3: SUB-SUB-SECTION (7.2.1)
            if chapter["sub_sections"]:
                sub_sub_no = chapter["sub_sections"][0]

                if (
                    re.fullmatch(r"\d+\.\d+\.\d+", sub_sub_no)
                    and chapter_counter > 0
                    and section_counter > 0
                ):
                    title = chapter.get("sub_sub_title")
                    title = title.strip() if title else None
                    if not title or len(title.split()) > 6:
                        continue
                    if title:
                        parent_key = ".".join(sub_sub_no.split(".")[:2])
                        if current_subsection_key != parent_key:
                            if len(sub_sub_buffer) >= 2:
                                formatted_toc.extend(sub_sub_buffer)
                            sub_sub_buffer = []
                            current_subsection_key = parent_key

                        entry = f"{chapter_counter}.{section_counter}.{len(sub_sub_buffer)+1}. {title}"
                        sub_sub_buffer.append(entry)

        if len(sub_sub_buffer) >= 2:
            formatted_toc.extend(sub_sub_buffer)

        # Extract company profiles
        company_profiles = []
        inside_company_profiles = False
        inside_key_players = False
        parent_chapter_number = None

        rows = response.xpath("//div[contains(@class,'clsTR')]")
        for row in rows:
            div_texts = row.xpath("./div/text()").getall()
            div_texts_clean = [t.strip() if t else "" for t in div_texts]

            head_texts = row.xpath(".//div[contains(@class,'txthead')]//text()").getall()
            head_clean = " ".join([t.strip() for t in head_texts if t.strip()])

            subhead = row.xpath(".//div[contains(@class,'txtsubhead')]/text()").get()
            subhead_clean = subhead.strip() if subhead else ""

            candidate = div_texts_clean[3] if len(div_texts_clean) >= 4 else None
            candidate = candidate.strip() if candidate and candidate != "\xa0" else None

            # Detect COMPANY PROFILES
            if "COMPANY PROFILES" in head_clean.upper():
                inside_company_profiles = True
                inside_key_players = False
                parent_chapter_match = re.match(r"(\d+)", head_clean)
                parent_chapter_number = parent_chapter_match.group(1) if parent_chapter_match else None
                continue

            if not inside_company_profiles:
                continue

            # Detect KEY PLAYERS, MAJOR PLAYERS
            if (
                re.fullmatch(r"\d+\.\d+", head_clean)
                and any(k in subhead_clean.upper() for k in ["KEY PLAYERS", "MAJOR PLAYERS"])
            ):
                inside_key_players = True
                continue

            # Stop at OTHER PLAYERS
            if (
                inside_key_players
                and re.fullmatch(r"\d+\.\d+", head_clean)
                and not any(k in subhead_clean.upper() for k in ["KEY PLAYERS", "MAJOR PLAYERS"])
            ):
                break

            # Extract company names
            if inside_key_players and parent_chapter_number:
                if re.match(rf"^{parent_chapter_number}\.\d+\.\d+$", subhead_clean):
                    if candidate and not re.match(r'^\d+(\.\d+)*$', candidate):
                        company_profiles.append(candidate)
        return formatted_toc, company_profiles
    
    def parse_segment_toc(self, response):
        table_of_contents = []
        chapter_counter = 0
        for p in response.css("div.tab-of-content p"):
            strong_text = p.css("strong::text").get()
            if not strong_text:
                continue
            strong_text = re.sub(r"\s+", " ", strong_text).strip()
            if " BY " not in strong_text.upper():
                continue
            if re.search(r"BY\s+(REGION|COUNTRY|GEOGRAPHY)", strong_text, re.I):
                break
            m = re.search(r"\bBY\s+(.+?)(?:\(|$)", strong_text, re.I)
            if not m:
                continue
            chapter_counter += 1
            section_counter = 0
            chapter_name = m.group(1).strip()
            table_of_contents.append(f"{chapter_counter}. {chapter_name}")
            lines = p.xpath(".//text()").getall()
            lines = [re.sub(r"\s+", " ", l).strip() for l in lines if l.strip()]
            lvl3_buffer = []
            current_section_no = None
            for line in lines:
                if line == strong_text:
                    continue
                if "INTRODUCTION" in line.upper() or "TABLE" in line.upper():
                    continue
                m2 = re.match(r"\d+\.(\d+)\s+(.+)", line)
                if m2 and not re.match(r"\d+\.\d+\.\d+", line):
                    title = m2.group(2).strip()
                    if len(title.split()) > 6:
                        continue
                    if len(lvl3_buffer) >= 2:
                        table_of_contents.extend(lvl3_buffer)
                    lvl3_buffer = []
                    section_counter += 1
                    current_section_no = f"{chapter_counter}.{section_counter}"
                    table_of_contents.append(f"{current_section_no}. {title}")
                    continue
 
                m3 = re.match(r"\d+\.(\d+)\.(\d+)\s+(.+)", line)
                if m3 and current_section_no:
                    title = m3.group(3).strip()
                    if len(title.split()) > 6:
                        continue
                    lvl3_index = len(lvl3_buffer) + 1
                    lvl3_buffer.append(f"{current_section_no}.{lvl3_index}. {title}")
 
            if lvl3_buffer:
                table_of_contents.extend(lvl3_buffer)
        # -------- COMPANY PROFILES --------
        company_profiles = []
        for p in response.css("div.tab-of-content p"):
            strong_text = p.css("strong::text").get()
            if not strong_text:
                continue
            if "COMPANY PROFILES" not in strong_text.upper():
                continue
            lines = p.xpath(".//text()").getall()
            lines = [re.sub(r"\s+", " ", l).strip() for l in lines if l.strip()]
            inside_key_players = False
            for line in lines:
                if re.search(r"\bKEY PLAYERS\b", line, re.I):
                    inside_key_players = True
                    continue
                if re.search(r"\bOTHER PLAYERS\b", line, re.I):
                    break
                if not inside_key_players:
                    continue
                m = re.match(r"\d+\.\d+\.\d+\s+(.+)", line)
                if not m:
                    continue
                company_name = m.group(1).strip()
                if "TABLE" in company_name.upper() or "FIGURE" in company_name.upper():
                    continue
                company_profiles.append(company_name)
        # -------- FALLBACK if no companies found in main pass --------
        if not company_profiles:
            for p in response.css("div.tab-of-content p"):
                strong_text = p.css("strong::text").get()
                if not strong_text:
                    continue
                if "COMPANY PROFILES" not in strong_text.upper():
                    continue
                lines = p.xpath(".//text()").getall()
                lines = [re.sub(r"\s+", " ", l).strip() for l in lines if l.strip()]
                for line in lines:
                    m = re.match(r"\d+(\.\d+)+\s+(.+)", line)
                    if m:
                        company_name = m.group(2).strip()
                        if "INTRODUCTION" in company_name.upper() or \
                        "TABLE" in company_name.upper() or \
                        "FIGURE" in company_name.upper():
                            continue
                        company_profiles.append(company_name)
                break  
        if not table_of_contents and not company_profiles:
            table_of_contents, company_profiles = self.parse_segment_toc_div(response)
        return table_of_contents, company_profiles
    
    def parse_segment_toc_div(self, response):
        table_of_contents = []
        chapter_counter = 0
        stop_parsing = False 
        div_nodes = response.css("div.tab-content div")
        i = 0
        while i < len(div_nodes) and not stop_parsing:
            node = div_nodes[i]
            strong_text = node.css("strong::text").get()
            if not strong_text:
                i += 1
                continue
            strong_text = re.sub(r"\s+", " ", strong_text).strip()
            # Skip if not a "BY" heading
            if " BY " not in strong_text.upper():
                i += 1
                continue
            # Stop parsing if heading indicates Regional Analysis, Country/Region breakdowns, or post-market sections
            if re.search(r"REGIONAL ANALYSIS|BY\s+(REGION|COUNTRY|GEOGRAPHY)|PESTLE|COMPETITIVE|COMPANY|AUTHOR", strong_text, re.I):
                break
            m = re.search(r"\bBY\s+(.+?)(?:\(|$)", strong_text, re.I)
            if not m:
                i += 1
                continue
            chapter_counter += 1
            section_counter = 0
            chapter_name = re.sub(r"\(Page No\..*?\)", "", m.group(1).strip())
            table_of_contents.append(f"{chapter_counter}. {chapter_name}")
            # -------- COLLECT FOLLOWING DIVS AS LEVEL 2 / LEVEL 3 --------
            lvl3_buffer = []
            current_section_no = None
            j = i + 1
            while j < len(div_nodes):
                next_node = div_nodes[j]
                next_strong = next_node.css("strong::text").get()
                if next_strong and " BY " in next_strong.upper():
                    break

                lines = next_node.xpath(".//text()").getall()
                lines = [re.sub(r"\s+", " ", l).strip() for l in lines if l.strip()]
                for line in lines:
                    # Stop parsing if line indicates Regional Analysis / post-market
                    if re.search(r"REGIONAL ANALYSIS|BY\s+(REGION|COUNTRY|GEOGRAPHY)|PESTLE|COMPETITIVE|COMPANY|AUTHOR", line, re.I):
                        stop_parsing = True
                        break

                    if line == strong_text or "INTRODUCTION" in line.upper() or "TABLE" in line.upper():
                        continue
                    m2 = re.match(r"\d+\.(\d+)\s+(.+)", line)
                    if m2 and not re.match(r"\d+\.\d+\.\d+", line):
                        if len(lvl3_buffer) >= 2:
                            table_of_contents.extend(lvl3_buffer)
                        lvl3_buffer = []
                        current_section_no = None
                        title = m2.group(2).strip()
                        # Skip extremely long titles (e.g., > 8 words)
                        if len(title.split()) > 6:
                            continue
                        section_counter += 1
                        current_section_no = f"{chapter_counter}.{section_counter}"
                        table_of_contents.append(f"{current_section_no}. {title}")
                        continue
                    m3 = re.match(r"\d+\.(\d+)\.(\d+)\s+(.+)", line)
                    if m3 and current_section_no:
                        title = m3.group(3).strip()
                        # Skip extremely long titles (e.g., > 8 words)
                        if len(title.split()) > 6:
                            continue
                        lvl3_index = len(lvl3_buffer) + 1
                        lvl3_buffer.append(f"{current_section_no}.{lvl3_index}. {title}")
                if stop_parsing:
                    break
                j += 1
            if len(lvl3_buffer) >= 2:
                table_of_contents.extend(lvl3_buffer)
            if stop_parsing:
                break
            i = j
        # -------- COMPANY PROFILES EXTRACTION ----------
        company_profiles = []
        found_company_section = False
        collecting_key_players = False
        for node in div_nodes:
            strong_text = node.css("strong::text").get()
            if strong_text and re.search(r"COMPANY PROFILES", strong_text, re.I):
                found_company_section = True
                continue  # move to next divs
            if found_company_section:
                lines = node.xpath(".//text()").getall()
                lines = [re.sub(r"\s+", " ", l).strip() for l in lines if l.strip()]
                for line in lines:
                    # Start collecting after KEY PLAYERS heading
                    if re.search(r"KEY PLAYERS", line, re.I):
                        collecting_key_players = True
                        continue
                    # Stop at OTHER PLAYERS or any post-market sections
                    if re.search(r"OTHER PLAYERS|DISCUSSION|KNOWLEDGE|AVAILABLE|RELATED REPORTS", line, re.I):
                        collecting_key_players = False
                        found_company_section = False
                        break
                    # Only collect if we are in KEY PLAYERS section
                    if collecting_key_players:
                        # Match company names like 16.2.1 ENERSYS
                        m = re.match(r"\d+\.\d+\.\d+\s+(.+)", line)
                        if m:
                            company_name = m.group(1).strip()
                            company_profiles.append(company_name)    
        return table_of_contents, company_profiles
    
    def parse_toc_from_accordion_xpath(self, response):
        toc = []
        accordion = response.xpath(
            '//div[contains(@class,"accordion")][.//strong[contains(.,"TABLE OF CONTENTS")]]'
        )
        if not accordion:
            return {"table_of_contents": toc}
        segment_headers = accordion.xpath(
            './/div[strong and contains(translate(., "abcdefghijklmnopqrstuvwxyz", "ABCDEFGHIJKLMNOPQRSTUVWXYZ"), " BY ")]'
        )
        segment_index = 0
        for header in segment_headers:
            header_text = " ".join(header.xpath(".//text()").getall())
            header_text = re.sub(r"\s+", " ", header_text).strip()
            header_upper = header_text.upper()
            if re.search(r"\b(BY REGION|KEY COUNTRY|BY COUNTRY)\b", header_upper):
                break
 
            segment_index += 1
            segment_name = re.sub(r"\(.*?\)", "", header_text)
            segment_name = re.sub(r".*BY", "", segment_name, flags=re.I).strip()
            toc.append(f"{segment_index}. {segment_name.title()}")
 
            # ---- parse ONLY following siblings ----
            sub_index = 0
            siblings = header.xpath("following-sibling::div")
            for sib in siblings:
                # Stop at next segment
                if sib.xpath(".//strong"):
                    break
 
                sib_text = " ".join(sib.xpath(".//text()").getall())
                sib_text = re.sub(r"\s+", " ", sib_text).strip()
 
                if not sib_text:
                    continue
 
                match = re.match(r"\d+(\.\d+)+\s+(.*)", sib_text)
                if not match:
                    continue
 
                sub_name = match.group(2).strip()
                if "INTRODUCTION" in sub_name.upper():
                    continue
 
                if len(sub_name.split()) > 6:
                    continue
 
                sub_index += 1
                toc.append(f"{segment_index}.{sub_index}. {sub_name.title()}")
 
        # Find COMPANY PROFILES header dynamically
        company_profiles = []
        company_header = accordion.xpath(
            './/div[strong and contains(translate(., "abcdefghijklmnopqrstuvwxyz", "ABCDEFGHIJKLMNOPQRSTUVWXYZ"), "COMPANY PROFILES")]'
        )
 
        if company_header:
            header_text = " ".join(company_header.xpath(".//text()").getall())
            header_text = re.sub(r"\s+", " ", header_text).strip()
 
            chapter_match = re.match(r"(\d+)\s+COMPANY PROFILES", header_text.upper())
 
            if chapter_match:
                chapter_no = chapter_match.group(1)
                siblings = company_header.xpath("following-sibling::div")
                in_key_players = False
                for div in siblings:
                    raw = " ".join(div.xpath(".//text()").getall())
                    text = re.sub(r"\s+", " ", raw.replace("\xa0", " ")).strip()
 
                    if not text:
                        continue
 
                    upper_text = text.upper()
                    if re.match(rf"{int(chapter_no)+1}\s+[A-Z]", upper_text):
                        break
                    # Detect KEY PLAYERS start
                    if "KEY PLAYERS" in upper_text:
                        in_key_players = True
                        continue
                    # Stop when OTHER PLAYERS starts
                    if "OTHER PLAYERS" in upper_text:
                        break
                    # Only parse when inside KEY PLAYERS
                    if not in_key_players:
                        continue
                    match = re.match(
                        rf"{chapter_no}\.\d+\.\d+\s+([A-Z0-9&.,()\- ]+)$",
                        text
                    )
                    if not match:
                        continue
 
                    company = match.group(1).strip()
                    # Exclude section labels
                    if re.search(
                        r"(BUSINESS OVERVIEW|PRODUCTS|RECENT DEVELOPMENTS|STRATEGY|WEAKNESSES)",
                        company.upper()
                    ):
                        continue
                    company_profiles.append(company.title())
 
        # Deduplicate companies (preserve order)
        seen = set()
        clean_companies = []
        for c in company_profiles:
            if c not in seen:
                seen.add(c)
                clean_companies.append(c)
 
        return {
            "table_of_contents": toc,
            "company_profiles": clean_companies
        }
 
    def parse_markets(self, response):
        accordion_exists = bool(response.css("div.accordion-item div.TOCcustHead"))
        table_exists = bool(response.css("div.tblTOC div.clsTR"))
        tab_content_exists = bool(response.css("div.tab-content p"))
        accordion_xpath_exists = bool(response.xpath('//div[contains(@class,"accordion")][.//div[contains(normalize-space(.),"TABLE OF CONTENTS")]]'))
        if accordion_exists:
            formatted_toc, company_profiles = self.parse_markets_accordion(response)
        elif table_exists:
            formatted_toc, company_profiles = self.parse_markets_table(response)
        elif tab_content_exists:
            formatted_toc, company_profiles = self.parse_segment_toc(response)
        elif accordion_xpath_exists:
            toc_data = self.parse_toc_from_accordion_xpath(response)
            formatted_toc = toc_data.get("table_of_contents", [])
            company_profiles = toc_data.get("company_profiles", [])    
        else:
            self.logger.warning("Unknown TOC structure")
            formatted_toc, company_profiles = [], []
 
        self.save_to_json({
            "title": self.clean_title(response.css("title::text").get()),
            "url": response.url,
            "table_of_contents": formatted_toc,
            "company_profiles": company_profiles
        }, response.url)
    ## Modified Ends Here ##
 
    # ==================== SNS INSIDER ====================
    def parse_sns(self, response):
        # Main page info
        title = self.clean_title(response.css('title::text').get() or "Untitled Report")
        main_url = response.url
        company_profiles = self.extract_sns_company_profiles(response)
        segmentation_url = f"{main_url.rstrip('/')}/segmentation"

        # Request segmentation page
        yield scrapy.Request(
            segmentation_url,
            callback=self.parse_sns_segmentation,
            meta={
                "title": title,
                "main_url": main_url,
                "company_profiles": company_profiles
            }
        )

    def parse_sns_segmentation(self, response):
        # Get main page data from meta
        title = response.meta["title"]
        main_url = response.meta["main_url"]
        company_profiles = response.meta["company_profiles"]
        segmentation_data = self.extract_sns_segmentation_data(response)

        self.save_to_json({
            "title": title,
            "url": main_url,
            "table_of_contents": segmentation_data,
            "company_profiles": company_profiles
        }, main_url)

    def extract_sns_segmentation_data(self, response):
        """Extract segmentation data from div.tab-content in flat TOC format, stopping at Regional Coverage."""
        toc = []
        segment_counter = 0
        for tab_content in response.css("div.tab-content"):
            p_tags = tab_content.css("p")
            for p in p_tags:
                p_text = p.xpath("string()").get()
                if not p_text:
                    continue
                p_text = p_text.strip()
                # Stop completely if we reach Regional Coverage
                if "Regional Coverage" in p_text:
                    return toc
                # Only consider main segments starting with "By"
                if not p_text.lower().startswith("by "):
                    continue
                # Increment main segment counter
                segment_counter += 1
                # Remove "By" from title
                segment_title = re.sub(r'^\s*by\s+', '', p_text, flags=re.IGNORECASE).strip()
                toc.append(f"{segment_counter}. {segment_title.upper()}")
                # Collect subsegments from following siblings
                sub_counter = 0
                siblings = p.xpath("following-sibling::*")
                for sibling in siblings:
                    # Stop completely if Regional Coverage appears
                    sib_text = sibling.xpath("string()").get()
                    if sib_text and "Regional Coverage" in sib_text:
                        return toc
 
                    tag_name = sibling.root.tag
                    if tag_name == "ul":
                        lis = sibling.css("li *::text, li::text").getall()
                        for li in lis:
                            li_text = li.strip()
                            if li_text:
                                sub_counter += 1
                                toc.append(f"{segment_counter}.{sub_counter}. {li_text}")
                    elif tag_name == "p":
                        if sib_text and sib_text.lower().startswith("by "):
                            break
        return toc
   
    def extract_sns_company_profiles(self, response):
        """Extract company profiles from various HTML structures, robust to variations in headings."""
        company_names = []
        # Loop through all relevant sections
        for tab_div in response.css("div.tab-content"):
            # Combine all heading and paragraph texts in this div
            heading_texts = tab_div.css("h2::text, h2 *::text, p::text, p *::text").getall()
            heading_texts = [t.strip() for t in heading_texts if t and t.strip()]
            # Check if any heading text mentions "Companies" or "Key Players" (ignore colon and case)
            if not any("companies are" in t.lower() or "key players" in t.lower() for t in heading_texts):
                continue
            # Grab all text inside li under ul (including nested tags)
            raw_items = tab_div.css("ul li strong::text, ul li strong a::text, ul li::text").getall()
            for text in raw_items:
                text = text.strip()
                if not text or re.fullmatch(r"[^\w]+", text):
                    continue
                # Filtering rules
                if len(text.split()) > 6 or len(text) > 60:
                    continue
                if re.search(r"\d|%|:|CAGR", text):
                    continue
                if text.startswith("By ") or text.startswith("In "):
                    continue
                company_names.append(text)
        company_names = sorted(set(company_names))
        return company_names
    
    # ==================== MORDOR INTELLIGENCE ====================
#start here
    def parse_mordor(self, response):
        toc_selector = response.css("#table-of-content")
        if not toc_selector:
            self.logger.warning(f"No TOC found on {response.url}")
            self.save_to_json({
                'title': self.clean_title(response.css('title::text').get()),
                'url': response.url,
                'table_of_contents': [],
                'company_profiles': []
            }, response.url)
            return
        # Extract all text content from TOC
        toc_text = toc_selector.xpath('normalize-space(.)').get()
     
        # Find all segmentation patterns that start with "By" including sub-segments
        segmentation_data = self.extract_mordor_segmentation_data(toc_text)
     
        # Normalize the numbering and extract text after "By"
        normalized_data = self.normalize_mordor_numbering(segmentation_data)
     
        # Extract company profiles
        company_profiles = self.extract_mordor_company_profiles(response)
     
        self.save_to_json({
            'title': self.clean_title(response.css('title::text').get()),
            'url': response.url,
            'table_of_contents': normalized_data,
            'company_profiles': company_profiles
        }, response.url)
    
    def extract_mordor_segmentation_data(self, toc_text):
        segmentation_lines = []
        toc_text = re.split(r'\d+\.\d+\s+By\s+Geography', toc_text)[0]
        main_pattern = r'(\d+\.\d+\s+By\s+[A-Za-z\s&/\-]+)'
        main_matches = re.findall(main_pattern, toc_text)
        for main_match in main_matches:
            main_header = re.sub(r'\s+', ' ', main_match.strip())

            base_number_match = re.search(r'^(\d+\.\d+)', main_header)
            if not base_number_match:
                continue

            base_number = base_number_match.group(1)
            segmentation_lines.append(main_header)

            sub_pattern = rf'({re.escape(base_number)}\.\d+(?:\.\d+)*\s+[A-Za-z\s&/\-]+)'
            sub_matches = re.findall(sub_pattern, toc_text)

            for sub_match in sub_matches:
                segmentation_lines.append(" " + re.sub(r'\s+', ' ', sub_match.strip()))

        return segmentation_lines
 
    def normalize_mordor_numbering(self, segmentation_data):
        """Normalize the numbering system to create proper hierarchical structure with sequential numbering"""
        normalized_data = []
        main_counter = 0
        sub_counter = 0
        sub_sub_counter = 0
        last_sub_item = None  # Track the last level-1 item number (e.g., "5.1.1")
        last_sub_sub_item = None  # Track the last level-2 item number (e.g., "5.1.1.1")
        
        for line in segmentation_data:
            # Check if it's a main header (no leading spaces)
            if not line.startswith(' '):
                main_counter += 1
                sub_counter = 0
                sub_sub_counter = 0
                last_sub_item = None
                last_sub_sub_item = None
                
                # Extract the text after "By" for main headers
                if "By " in line:
                    # Extract everything after "By "
                    text_part = line.split("By ", 1)[1].strip()
                    # Remove any numbering prefix that might be left
                    text_part = re.sub(r'^\d+\.\d+\s*', '', text_part).strip()
                else:
                    # If "By" is not found, use the original text after numbering
                    text_part = re.sub(r'^\d+\.\d+\s+', '', line).strip()
                
                # Create new numbering for main header
                new_line = f"{main_counter}. {text_part}"
                normalized_data.append(new_line)
            else:
                # It's a sub-item - need to parse the old numbering to understand hierarchy
                # Remove leading spaces for parsing
                clean_line = line.strip()
                
                # Extract old numbering pattern (e.g., "5.1.1", "5.1.1.1", etc.)
                old_number_match = re.match(r'^(\d+\.\d+(?:\.\d+)*)\s+(.+)', clean_line)
                
                if old_number_match:
                    old_number = old_number_match.group(1)
                    text_part = old_number_match.group(2).strip()
                    
                    # Count the dots in the old number to determine level
                    dot_count = old_number.count('.')
                    parts = old_number.split('.')
                    
                    if dot_count == 2:
                        # Level 1 sub-item (e.g., "5.1.1") - becomes 1.1, 1.2, 1.3, etc.
                        current_sub_item = f"{parts[0]}.{parts[1]}.{parts[2]}"
                        
                        # If this is a new level-1 item number, increment sub_counter
                        if current_sub_item != last_sub_item:
                            sub_counter += 1
                            sub_sub_counter = 0  # Reset sub-sub counter for new parent
                            last_sub_item = current_sub_item
                        
                        new_line = f" {main_counter}.{sub_counter}. {text_part}"
                    elif dot_count == 3:
                        # Level 2 sub-item (e.g., "5.1.1.1") - becomes 1.1.1, 1.1.2, etc.
                        current_sub_sub_item = old_number
                        
                        # If this is a new level-2 item, increment sub_sub_counter
                        if current_sub_sub_item != last_sub_sub_item:
                            sub_sub_counter += 1
                            last_sub_sub_item = current_sub_sub_item
                        
                        new_line = f" {main_counter}.{sub_counter}.{sub_sub_counter}. {text_part}"
                    else:
                        # For deeper or other nesting, just clean the text
                        new_number = f"{main_counter}." + ".".join(parts[2:])
                        new_line = f" {new_number}. {text_part}"
                    
                    normalized_data.append(new_line)
                else:
                    # Fallback: just clean and add with spacing
                    text_part = re.sub(r'^\d+(?:\.\d+)*\s*', '', clean_line).strip()
                    sub_counter += 1
                    normalized_data.append(f" {main_counter}.{sub_counter}. {text_part}")
        
        return normalized_data
 
    def extract_mordor_company_profiles(self, response):
        """Extract company profiles from the table of contents"""
        company_profiles = []
     
        # Look for the Company Profiles section with different patterns
        company_patterns = [
            "Company Profiles",
            "Company Profile",
            "Company Share Analysis",
            "Key Company Profiles",
            "Major Company Profiles"
        ]
     
        company_section = None
        for pattern in company_patterns:
            company_section = response.css(f"#table-of-content li:contains('{pattern}')")
            if company_section:
                break
     
        if company_section:
            # Find the specific section with company names (usually toc-level-3 items)
            # Look for the immediate ul containing company items
            company_ul = company_section.xpath("./following-sibling::ul[1]")
         
            if not company_ul:
                # Try another pattern - look for ul within the same parent
                company_ul = company_section.xpath("../following-sibling::ul[1]")
         
            if not company_ul:
                # Try to find any ul that contains toc-level-3 items after the company section
                company_ul = company_section.xpath("ancestor::li/following-sibling::li//ul")
         
            if company_ul:
                # Extract all company names from list items with specific class patterns
                company_items = company_ul.css("li.toc-level-3")
             
                # If no toc-level-3 items found, try to find any list items with company-like numbering
                if not company_items:
                    company_items = company_ul.css("li")
             
                for item in company_items:
                    # Get the text content and clean it up
                    company_text = item.xpath("normalize-space(.)").get()
                 
                    # Remove numbering prefix (e.g., "6.4.1 ")
                    company_name = re.sub(r'^\d+\.\d+\.\d+\s+', '', company_text)
                 
                    # Remove any trailing special characters or extra spaces
                    company_name = re.sub(r'[^\w\s&().,\-]', '', company_name).strip()
                 
                    # Filter out non-company names using a more specific pattern
                    if (self.is_valid_company_name(company_name) and
                        len(company_name) > 2 # Filter out very short names
                    ):
                        company_profiles.append(company_name)
     
        # If we still don't have company profiles, try a more direct approach
        if not company_profiles:
            # Look for all toc-level-3 items that might be companies
            all_toc_items = response.css("#table-of-content li.toc-level-3")
            for item in all_toc_items:
                company_text = item.xpath("normalize-space(.)").get()
                company_name = re.sub(r'^\d+\.\d+\.\d+\s+', '', company_text)
                company_name = re.sub(r'[^\w\s&().,\-]', '', company_name).strip()
             
                if (self.is_valid_company_name(company_name) and
                    len(company_name) > 2
                ):
                    company_profiles.append(company_name)
     
        # Remove duplicates and sort alphabetically
        company_profiles = sorted(list(set(company_profiles)))
        return company_profiles
# end here 
    # ==================== FORTUNE BUSINESS INSIGHTS ====================
    ## Modified Starts Here ##
    def extract_company_profiles(self, response):
            companies = []
            company_sections = response.xpath(
                '//div[@id="summary"]//h3['
                '('
                    'contains(translate(., "ABCDEFGHIJKLMNOPQRSTUVWXYZ", "abcdefghijklmnopqrstuvwxyz"), "list of key") and '
                    'contains(translate(., "ABCDEFGHIJKLMNOPQRSTUVWXYZ", "abcdefghijklmnopqrstuvwxyz"), "companies profiled")'
                ') or '
                'contains(translate(., "ABCDEFGHIJKLMNOPQRSTUVWXYZ", "abcdefghijklmnopqrstuvwxyz"), "long list of companies studied") or '
                'contains(translate(., "ABCDEFGHIJKLMNOPQRSTUVWXYZ", "abcdefghijklmnopqrstuvwxyz"), "key players covered")'
                ']/following-sibling::ul[1] | '
            
                '//div[@id="summary"]//h2['
                    'contains(translate(., "ABCDEFGHIJKLMNOPQRSTUVWXYZ", "abcdefghijklmnopqrstuvwxyz"), "list of") and '
                    'contains(translate(., "ABCDEFGHIJKLMNOPQRSTUVWXYZ", "abcdefghijklmnopqrstuvwxyz"), "companies")'
                ']/following-sibling::ul[1]'
            )
            # Fallback 1: any ul with country in parentheses
            if not company_sections:
                company_sections = response.xpath('//ul[.//li[contains(text(), "(") and contains(text(), ")")]]')
            # Fallback 2: ul with 3–20 li items
            if not company_sections:
                company_sections = response.xpath('//ul[count(li) > 3 and count(li) < 20]')
    
            # Words to exclude that are clearly not company names
            invalid_keywords = [
                "toc", "segmentation", "methodology", "request", "sample",
                "copy", "buy", "faqs", "testimonials", "terms", "privacy",
                "policy", "careers", "order", "how to", "contact", "press release",
                "for instance", "note", "announcement", "announced", "report", "study"
            ]
    
            for section in company_sections:
                for item in section.xpath('./li'):
                    # Combine all text nodes in this li, including <a> links
                    text = ''.join(item.xpath('.//text()').getall()).strip()
                    if not text:
                        continue
    
                    # Remove country info in parentheses
                    company_name = re.sub(r'\s*\([^)]*\).*$', '', text).strip()
                    company_name = company_name.replace('&amp;', '&')
                    company_name = re.sub(r'\s+', ' ', company_name)
    
                    # Filter valid company names
                    if (
                        len(company_name) > 1 and  # allow short names like LG
                        len(company_name.split()) <= 10 and  # allow long names
                        not any(word in company_name.lower() for word in invalid_keywords) and
                        company_name not in companies
                    ):
                        companies.append(company_name)
            return companies

    def parse_fortune_main(self, response):
        title = self.clean_title(response.css('title::text').get())
        main_url = response.url
        company_profiles = self.extract_company_profiles(response)

        # Build segmentation URL
        if "/industry-reports/" in main_url:
            segmentation_url = main_url.replace(
                "/industry-reports/",
                "/industry-reports/segmentation/"
            )
        else:
            segmentation_url = main_url.replace(
                "fortunebusinessinsights.com/",
                "fortunebusinessinsights.com/segmentation/"
            )

        # Only request the segmentation page first
        yield scrapy.Request(
            segmentation_url,
            callback=self.parse_fortune_segmentation,
            meta={
                "title": title,
                "main_url": main_url,
                "company_profiles": company_profiles,
                "is_main_page": False,  # Block summary fallback
                "zyte_api_automap": True,
                "zyte_api_render": {
                    "browserHtml": True,
                    "waitFor": "div.tab-content"
                }
            }
        )

    def parse_fortune_segmentation(self, response):
        structured_points = []
        main_url = response.meta["main_url"]
        title = response.meta["title"]
        company_profiles = response.meta.get("company_profiles", [])
        rows = response.xpath('//div[@id="industrycoverage"]//table//tr')
        main_count = 0
        start_processing = False  # <-- Flag to start after "Segmentation"
        for row in rows:
            # Decide where "By ..." lives
            if row.xpath('./td[@rowspan]'):
                left_td = row.xpath('./td[2]')
            elif start_processing and not row.xpath('./td[2]'):
                left_td = row.xpath('./td[1]')
            else:
                left_td = row.xpath('./td[1]')
            # Extract ONLY the "By ..." heading (ignore <li>)
            left_text = ' '.join(
                t.strip()
                for t in left_td.xpath('.//p//strong//text()').getall()
                if t.strip()
            )
            # Fallback if <strong> is missing
            if not left_text:
                left_text = ' '.join(
                    t.strip()
                    for t in left_td.xpath('./p[1]//text()').getall()
                    if t.strip()
                )
            # Detect start of segmentation (ALWAYS read from td[1])
            seg_text = ' '.join(
                t.strip() for t in row.xpath('./td[1]//text()').getall() if t.strip()
            )
            if not start_processing:
                if "segmentation" in seg_text.lower():
                    start_processing = True
                else:
                    continue

            # Skip rows not starting with "By"
            if not left_text.lower().startswith("by"):
                if not start_processing:
                    continue

                content_td = row.xpath('./td[2]')
                by_heads = content_td.xpath(
                    './/p/strong[starts-with(normalize-space(.), "By")]/text()'
                ).getall()
                if not by_heads:
                    continue

                for head in by_heads:
                    lower_head = head.lower()

                    # Skip region / geography / country
                    if any(k in lower_head for k in ("by region", "by geography", "by country")):
                        continue
                    main_count += 1
                    sub = sub_sub = sub_sub_sub = 0
                    segment = re.sub(r'^by\s+', '', head, flags=re.IGNORECASE).strip()
                    structured_points.append(f"{main_count}. {segment}")
                    ul = content_td.xpath(
                        f'.//p[strong[text()="{head}"]]/following-sibling::ul[1]'
                    )

                    for li in ul.xpath('.//li'):
                        text = ' '.join(li.xpath('.//text()').getall()).replace('\xa0', ' ').strip()
                        if text:
                            sub += 1
                            structured_points.append(f"{main_count}.{sub}. {text}")
                continue
            # Skip region / geography / country
            lower_left = left_text.lower()
            if any(k in lower_left for k in ("by region", "by geography", "by country")):
                continue
            main_count += 1
            sub = sub_sub = sub_sub_sub = 0
            segment = re.sub(r'^by\s+', '', left_text, flags=re.IGNORECASE).strip()
            structured_points.append(f"{main_count}. {segment}")
            content_td = row.xpath('./td[2]') if row.xpath('./td[2]') else row.xpath('./td[1]')
            found_symbol_items = False
            for p in content_td.xpath('.//p'):
                text = ''.join(p.xpath('.//text()').getall())
                text = text.replace('\xa0', ' ').strip()
                if not text:
                    continue
                # LEVEL 2 — ·
                if text.startswith('·'):
                    sub += 1
                    sub_sub = sub_sub_sub = 0
                    structured_points.append(
                        f"{main_count}.{sub}. {text.lstrip('·').strip()}"
                    )
                # LEVEL 3 — o
                elif text.startswith('o'):
                    clean_text = text.lstrip('o').strip()
                    if len(clean_text.split()) <= 6:
                        sub_sub += 1
                        sub_sub_sub = 0
                        structured_points.append(
                            f"{main_count}.{sub}.{sub_sub}. {clean_text}"
                        )
                # LEVEL 4 — §
                elif text.startswith('§'):
                    clean_text = text.lstrip('§').strip()
                    if len(clean_text.split()) <= 6:
                        sub_sub_sub += 1
                        structured_points.append(
                            f"{main_count}.{sub}.{sub_sub}.{sub_sub_sub}. {clean_text}"
                        )
            # --- <ul><li> fallback ---
            if not found_symbol_items:
                for li in content_td.xpath('.//li'):
                    text = ' '.join(li.xpath('.//text()').getall()).strip()
                    if text:
                        sub += 1
                        structured_points.append(f"{main_count}.{sub}. {text}")

        # Trigger main page fallback if segmentation page TOC is empty
        if not structured_points and not response.meta.get("is_main_page", False):
            # Fetch the main page as fallback
            yield scrapy.Request(
                response.meta["main_url"],
                callback=self.parse_fortune_segmentation,
                meta={
                    "title": response.meta["title"],
                    "main_url": response.meta["main_url"],
                    "company_profiles": response.meta.get("company_profiles", []),
                    "is_main_page": True,   # Allow summary H2 fallback
                    "zyte_api_automap": True,
                    "zyte_api_render": {
                        "browserHtml": True,
                        "waitFor": "div#summary"
                    }
                },
                dont_filter=True
            )
            return
        # ================= SUMMARY-H2 ANCHORED FALLBACK =================
        if not structured_points and response.meta.get("is_main_page") is True:
            # Locate "Segmentation" H2 first
            seg_h2 = response.xpath(
                '//div[@id="summary"]//h2[strong[normalize-space(text())="Segmentation"]]'
            )
            if seg_h2:
                # Fetch the first table AFTER the Segmentation heading
                table = seg_h2.xpath('following-sibling::table[1]')

                if table:
                    header_tds = table.xpath('.//tr[1]/td')
                    value_tds = table.xpath('.//tr[2]/td')

                    if len(header_tds) == len(value_tds):

                        for idx, header_td in enumerate(header_tds):
                            header_text = ' '.join(
                                t.strip()
                                for t in header_td.xpath('.//text()').getall()
                                if t.strip()
                            )

                            lower_header = header_text.lower()
                            if any(k in lower_header for k in ("region", "geography", "country")):
                                continue

                            if not lower_header.startswith("by"):
                                continue
                            main_count += 1
                            sub = 0
                            segment = re.sub(
                                r'^by\s+', '', header_text, flags=re.IGNORECASE
                            ).strip()
                            structured_points.append(f"{main_count}. {segment}")
                            value_td = value_tds[idx]
                            for li in value_td.xpath('.//li'):
                                text = ' '.join(
                                    li.xpath('.//text()').getall()
                                ).replace('\xa0', ' ').strip()

                                if text:
                                    sub += 1
                                    structured_points.append(
                                        f"{main_count}.{sub}. {text}"
                                    )

        result_data = {
            "title": title,
            "url": main_url,
            "table_of_contents": structured_points
        }
        if company_profiles:
            result_data["company_profiles"] = company_profiles
        self.save_to_json(result_data, main_url)
    ## Modified Ends Here ##
    # ==================== FUTURE MARKET INSIGHTS ====================
    def parse_future(self, response):
        # Extract market segmentation data and format as table_of_contents
        table_of_contents = self.extract_future_segmentation_as_toc(response)
        # Extract company profiles
        company_profiles = self.extract_future_company_profiles(response)
        result_data = {
            'title': self.clean_title(response.css('title::text').get()),
            'url': response.url,
            'table_of_contents': table_of_contents,
            'company_profiles': company_profiles
        }
        self.save_to_json(result_data, response.url)
    ## Modification Starts Here ##
    def extract_future_segmentation_as_toc(self, response):
        """Extract ONLY market segmentation data and format as numbered TOC"""
        toc_lines = []
        content_div = response.css('div.report_content_div div.tab_content')
        added_segments = set()
        # ---------- STEP 1: FIND SEGMENTATION H2 ----------
        segmentation_h2 = None
        for h2 in content_div.css('h2'):
            text = h2.xpath('normalize-space(.)').get()
            if not text:
                continue
            text_lower = text.lower()
            # skip H2 if it contains "segmental analysis"
            if 'segmental analysis' in text_lower:
                continue
            if (
                any(word in text_lower for word in ['segment', 'segments', 'key segments'])
                and '?' not in text_lower
            ):
                segmentation_h2 = h2
                break

        if not segmentation_h2:
            return toc_lines
        # ---------- STEP 2: COLLECT NODES UNTIL NEXT H2 ----------
        nodes = []
        for sib in segmentation_h2.xpath('following-sibling::*'):
            if sib.root.tag == 'h2':
                break
            nodes.append(sib)
        # ---------- STEP 3: PARSE SEGMENTS ----------
        main_num = 1
        current_category = None
        for node in nodes:
            tag = node.root.tag
            # ---- SEGMENT TITLE ----
            if tag == 'h3':
                raw_title = node.xpath('normalize-space(.)').get()
                if not raw_title:
                    current_category = None
                    continue
                title = raw_title.strip().rstrip(':').strip()
                if title.lower().startswith('by '):
                    title = title[3:].strip()
                title_clean = title.lower()
                if title_clean in {'region', 'regions', 'country', 'countries'}:
                    current_category = None
                    continue
                if title_clean in added_segments:
                    current_category = None
                    continue
                current_category = title
            # ---- SUB SEGMENTS with nested UL and minimum 2 sub-sub-items ----
            elif tag == 'ul' and current_category:
                sub_num = 1
                sub_lines = []
                first_level_has_valid_items = False
                for li in node.xpath('./li'):
                    li_text = li.xpath('normalize-space(text())').get()
                    if not li_text:
                        continue
                    li_text = li_text.strip()
                    # check if this li has a nested ul
                    nested_ul = li.xpath('./ul')
                    if nested_ul:
                        nested_items = []
                        for nested_li in nested_ul.xpath('./li'):
                            nested_text = nested_li.xpath('normalize-space(.)').get()
                            if nested_text:
                                nested_text = nested_text.strip()
                                nested_items.append(nested_text)
                        # only add if >=2 nested items
                        if len(nested_items) >= 2:
                            sub_lines.append(f"{main_num}.{sub_num}. {li_text}")
                            sub_sub_num = 1
                            for nested_text in nested_items:
                                sub_lines.append(f"{main_num}.{sub_num}.{sub_sub_num}. {nested_text}")
                                sub_sub_num += 1
                            first_level_has_valid_items = True
                    else:
                        # no nested ul, treat li normally as sub-segment
                        sub_lines.append(f"{main_num}.{sub_num}. {li_text}")
                        first_level_has_valid_items = True
                    sub_num += 1
                # only add category if there is at least one valid sub-item
                if first_level_has_valid_items:
                    toc_lines.append(f"{main_num}. {current_category}")
                    added_segments.add(current_category.lower())
                    toc_lines.extend(sub_lines)
                    main_num += 1
                current_category = None
        return toc_lines

    def extract_future_company_profiles(self, response):
        """Extract company profiles from h2 + ul blocks inside report_content_div"""
        company_profiles = []
        content_div = response.css('div.report_content_div div.tab_content')
        # ---- FIND COMPANY PROFILE H2 ----
        company_h2 = None
        for h2 in content_div.css('h2'):
            h2_text = h2.xpath('normalize-space(.)').get()
            if not h2_text:
                continue
            h2_lower = h2_text.lower()
            if any(keyword in h2_lower for keyword in [
                'key industry participants',
                'key players',
                'key companies',
                'company profiles',
                'market players'
            ]):
                company_h2 = h2
                break
        # No company section
        if company_h2:
            # ---- PARSE UNTIL NEXT H2 ----
            for sib in company_h2.xpath('following-sibling::*'):
                if sib.root.tag == 'h2':
                    break
                if sib.root.tag == 'ul':
                    for li in sib.css('li'):
                        text = li.xpath('normalize-space(.)').get()
                        if text:
                            company_profiles.append(text)
        # ---- FALLBACK: TABLE EXTRACTION IF NO UL FOUND ----
        if not company_profiles:
            # Find table rows containing "Key Companies Profiled"
            rows = response.xpath("//tbody/tr")
            for tr in rows:
                td_label = tr.xpath("normalize-space(td[1])").get()
                td_value = tr.xpath("normalize-space(td[2])").get()
                if td_label and "Key Companies Profiled" in td_label and td_value:
                    # Split by comma
                    for company in td_value.split(","):
                        company = company.strip()
                        if company:
                            company_profiles.append(company)
                    break  
        # ---- DEDUPLICATE (PRESERVE ORDER) ----
        company_profiles = list(dict.fromkeys(company_profiles))
        return company_profiles
    ## Modification Ends Here ##
    # ==================== ALLIED MARKET RESEARCH ====================
    def parse_allied(self, response):
        report_id = response.css('input#report_id::attr(value)').get()
        if not report_id:
            script_content = response.xpath('//script[contains(text(), "report_id")]/text()').get()
            if script_content and 'report_id' in script_content:
                report_id = script_content.split('report_id:')[1].split(',')[0].strip().strip("'\"")
        if report_id:
            toc_url = f"https://www.alliedmarketresearch.com/get-report-toc-rev/{report_id}"
            yield scrapy.Request(
                toc_url,
                callback=self.parse_allied_toc,
                meta={
                    "page_url": response.url,
                    "title": response.css('title::text').get()
                }
            )
        else:
            # Extract just the market name from the full title
            full_title = response.css('title::text').get()
            market_name = self.extract_market_name(full_title)
         
            self.save_to_json({
                'title': market_name,
                'url': response.url,
                'table_of_contents': ["❌ Report ID not found. TOC cannot be extracted."],
                'company_profiles': ["❌ Report ID not found. Company profiles cannot be extracted."]
            }, response.url)
 
    def extract_market_name(self, full_title):
        """Extract just the market name from the full title"""
        if not full_title:
            return "Unknown Market"
     
        # Remove everything after "Market" or "Market Size"
        market_match = re.search(r'^(.*?Market)\b', full_title)
        if market_match:
            return market_match.group(1).strip()
     
        # If "Market" is not found, return the first part of the title
        return full_title.split('|')[0].split('-')[0].strip()
 
## Modification starts here##
    def parse_allied_toc(self, response):
        flat_toc_list = []
        company_profiles_list = []
     
        # Variables for segmentation extraction
        chapter_number = 0
        start_collecting = False
        stop_collecting = False
     
        # Variables for company profiles extraction
        in_company_profiles = False
        for chapter_card in response.css('#acordTabOCnt > .card'):
            chapter_title = chapter_card.css('.card-header .btn-link .fw-700::text').get()
            if chapter_title:
                chapter_title = re.sub(r'^\d+(\.\d+)*\.\s*', '', chapter_title.strip())
             
                # Check if we're in the company profiles section
                if "company profile" in chapter_title.lower() or "company profiles" in chapter_title.lower():
                    in_company_profiles = True
                    stop_collecting = True # Stop collecting segmentation sections
                elif in_company_profiles:
                    # We've passed the company profiles section
                    break
             
                # Segmentation extraction logic
                if not start_collecting and not in_company_profiles:
                    if "market overview" in chapter_title.lower() or "market landscape" in chapter_title.lower():
                        start_collecting = True
                        continue # Skip market overview/landscape chapter itself
                    else:
                        continue
                # Check if this is the "BY REGION" chapter and stop collecting
                if not in_company_profiles and " BY REGION" in chapter_title.upper():
                    stop_collecting = True
                    continue
                # Extract segmentation sections - handle both "BY" format and other formats
                if not in_company_profiles and not stop_collecting:
                    # Extract the segment type - handle different formats
                    if " BY " in chapter_title.upper():
                        # Format: "CHAPTER X : MARKET NAME, BY SEGMENT_TYPE"
                        segment_type = chapter_title.upper().split(" BY ")[-1].strip()
                        chapter_number += 1
                        flat_toc_list.append(f"{chapter_number}. {segment_type}")
                    elif "," in chapter_title and ":" in chapter_title:
                        # Format: "CHAPTER X : MARKET NAME, SEGMENT_TYPE"
                        # Extract the part after the last comma
                        segment_part = chapter_title.split(",")[-1].strip()
                        chapter_number += 1
                        flat_toc_list.append(f"{chapter_number}. {segment_part}")
                    else:
                        # Generic fallback - just use the chapter title
                        chapter_number += 1
                        flat_toc_list.append(f"{chapter_number}. {chapter_title}")
            # Process segmentation sections - only main sections (no sub-sections)
            if not in_company_profiles and not stop_collecting:
                card_body = chapter_card.css('.card-body')
                if card_body:
                    # Get all h3 elements which represent main sections
                    main_sections = card_body.css('h3, p span span, p span, p')
                    section_number = 1
                    seen_subsegments = set()
                    temp_subsegments = []
                    # First, collect all valid sub-segments for this chapter
                    for section in main_sections:
                        text = ' '.join(section.css('::text').getall()).strip()
                        if not text:
                            continue

                        if re.match(r'^\d+\.\d+\.\d+\.', text):
                            continue

                        match = re.match(r'^\d+\.\d+\.\s*(.+)', text)
                        if not match:
                            continue

                        clean_title = match.group(1).strip()

                        if re.search(
                            r'market size|forecast|key market|overview|opportunit|by region',
                            clean_title,
                            re.IGNORECASE
                        ):
                            continue

                        # Skip if more than 6 words
                        if len(clean_title.split()) > 6:
                            continue

                        key = (chapter_number, clean_title.lower())
                        if key in seen_subsegments:
                            continue

                        seen_subsegments.add(key)
                        temp_subsegments.append(clean_title)
                    # Only add sub-segments if count >= 2
                    if len(temp_subsegments) >= 2:
                        for sub in temp_subsegments:
                            flat_toc_list.append(f"{chapter_number}.{section_number}. {sub}")
                            section_number += 1

            # Process company profiles - extract only company names without numbering
            if in_company_profiles:
                card_body = chapter_card.css('.card-body')
                if card_body:
                    # Extract company names from h3 elements
                    company_h3s = card_body.css('p, p span span, p span, h3')
                    for item in company_h3s:
                        text = ' '.join(item.css('::text').getall()).strip()
                        if not text:
                            continue

                        # Remove 3rd-level or non-numbered
                        match = re.match(r'^\d+\.\d+\.\s*(.+)', text)
                        if not match:
                            continue

                        company_name = match.group(1).strip()

                        # Remove junk headings
                        if re.search(
                            r'overview|definition|research|dynamics|analysis|description',
                            company_name,
                            re.IGNORECASE
                        ):
                            continue

                        # Clean suffixes
                        company_name = re.sub(r'\s*\(.*?\)', '', company_name)
                        company_name = re.sub(
                            r'\s*(Inc|Ltd|LLC|Corp|Co|Limited|Pvt)\.?$',
                            '',
                            company_name,
                            flags=re.IGNORECASE
                        ).strip()

                        if company_name and company_name not in company_profiles_list:
                            company_profiles_list.append(company_name)
        # Extract just the market name from the full title
        full_title = response.meta["title"]
        market_name = self.extract_market_name(full_title)
        self.save_to_json({
            'title': market_name,
            'url': response.meta["page_url"],
            'table_of_contents': flat_toc_list if flat_toc_list else ["No segmentation sections found after Market Overview."],
            'company_profiles': company_profiles_list if company_profiles_list else ["No company profiles found in TOC."]
        }, response.meta["page_url"])
    ## Modification ends here##
# ==================== SkyQuest Technology ====================
    def parse_skyquestt(self, response):
        segments_and_companies = self.extract_segments_and_companies(response)
        result_data = {
            "title": self.clean_title(response.css('title::text').get()),
            "url": response.url,
            "table_of_contents": segments_and_companies.get("table_of_contents", []),
            "company_profiles": segments_and_companies.get("company_profiles", [])
        }
        self.save_to_json(result_data, response.url)
 
    def extract_segments_and_companies(self, response):
        table_of_contents = []
        company_profiles = []
        base = 'div.accordion-body div.special-toc-class > ul > li'
        segment_counter = 0
        matched_items = response.css(base)
        for item in matched_items:
            # Heading (direct <strong> child)
            heading = item.xpath('./strong/text() | ./b/text()').get()
            if not heading:
                continue
            heading_clean = heading.strip()
            if "by region" in heading_clean.lower():
                continue
            # 1️⃣ SEGMENTATION (BY ...)
            if " by " in heading_clean.lower():
                segment_counter += 1
                segment_name = heading_clean.split(' by ', 1)[1]
                segment_name = segment_name.split('&')[0].strip()
                table_of_contents.append(f"{segment_counter}. {segment_name}")
 
                # Try direct <ul> first
                sub_list = item.xpath('./ul/li')
                # Fallback: check next sibling <li> for <ul>
                if not sub_list:
                    sub_list = item.xpath('./following-sibling::li[1]/ul/li')
 
                sub_counter = 0
                for sub_li in sub_list:
                    sub_title = sub_li.xpath('./text()').get()
                    if not sub_title:
                        continue
 
                    sub_title = sub_title.strip()
                    if sub_title.lower() == "market overview":
                        continue
 
                    sub_counter += 1
                    table_of_contents.append(f"{segment_counter}.{sub_counter}. {sub_title}")
 
                    # SUB-SUB SEGMENTS
                    sub_sub_list = sub_li.xpath('./ul/li')
                    # Filter valid sub-sub-segments first
                    valid_sub_sub = [
                        s for s in sub_sub_list
                        if s.xpath('./text()').get() and s.xpath('./text()').get().strip().lower() != "market overview"
                    ]
                    # Only proceed if 2 or more valid sub-sub-segments
                    if len(valid_sub_sub) >= 2:
                        sub_sub_counter = 0
                        for sub_sub_li in valid_sub_sub:
                            sub_sub_title = sub_sub_li.xpath('./text()').get().strip()
                            sub_sub_counter += 1
                            table_of_contents.append(
                                f"{segment_counter}.{sub_counter}.{sub_sub_counter}. {sub_sub_title}"
                            )
           
            # 2️⃣ KEY COMPANY PROFILES
            elif re.search(r'^key\s+company\s+profiles$', heading_clean, re.I):
                # Try direct <ul/li> first
                company_list = item.xpath('./ul/li')
                # Fallback: check next sibling <li> for <ul>
                if not company_list:
                    company_list = item.xpath('./following-sibling::li[1]/ul/li')
 
                for company_li in company_list:
                    company_name = company_li.xpath('./text()').get()
                    if not company_name:
                        continue
                    company_name = company_name.strip()
                    if len(company_name.split()) > 10:
                        continue
                    company_profiles.append(company_name)
        return {
            "table_of_contents": table_of_contents,
            "company_profiles": company_profiles
        }
#===================== MRF =====================
    def parse_mrf(self, response):
        raw_title = response.css("title::text").get()
        title = self.clean_title(raw_title)
        # Fallback ONLY if title is missing
        if not title:
            h1_fallback = response.xpath(
                "//div[contains(@class,'rd-title-cont')]"
                "//h1[contains(@class,'report-title')]/text()"
            ).get()
 
            if h1_fallback:
                title = self.clean_title(h1_fallback.strip().title())
        segments_and_companies = self.parse_mrf_market(response)
        result_data = {
            "title": title,
            "url": response.url,
            "table_of_contents": segments_and_companies.get("table_of_contents", []),
            "company_profiles": segments_and_companies.get("company_profiles", [])
        }
        self.save_to_json(result_data, response.url)
 
    def parse_mrf_market(self, response):
        table_of_contents = []
        company_profiles = []
        section = response.xpath(
            "//div[contains(@class,'section-heading-two')]"
            "[./h2[normalize-space()='Market Segmentation']]"
            "/following-sibling::div[contains(@class,'section-content')][1]"
        )
        # MARKET SEGMENTATION
        segment_counter = 1
        cards = section.xpath(".//div[contains(@class,'inner-section-cont')]")
        for card in cards:
            heading = card.xpath(
                ".//h3[contains(@class,'sec-heading-cont')]/i/text()"
            ).get()
            if not heading:
                continue
            match = re.search(
                r".*(?:market|industry|by)\s+(.+?)\s+outlook\s*$",
                heading,
                re.IGNORECASE
            )
            if not match:
                continue
            segment_title = match.group(1).strip().upper()
            table_of_contents.append(
                f"{segment_counter}. {segment_title}"
            )
            sub_segments = card.xpath(".//ul/li/text()").getall()
            sub_counter = 1
            for sub in sub_segments:
                # split by newlines and clean
                parts = [p.strip() for p in sub.split("\n") if p.strip()]
 
                if not parts:
                    continue
 
                # FIRST part → normal sub-segment (e.g. BAGS)
                parent = parts[0].upper()
                table_of_contents.append(
                    f"{segment_counter}.{sub_counter}. {parent}"
                )
                # REMAINING parts → sub-sub-segments
                if len(parts) > 1:
                    sub_sub_counter = 1
                    for child in parts[1:]:
                        table_of_contents.append(
                            f"{segment_counter}.{sub_counter}.{sub_sub_counter}. {child.upper()}"
                        )
                        sub_sub_counter += 1
                sub_counter += 1
            segment_counter += 1
 
        # COMPANY PROFILES
        players_block = response.xpath(
            "//div[contains(@class,'section-content')]"
            "//div[contains(@class,'sec-cont-sub-heading')]"
            "[./h3[normalize-space()='Major Players']]"
            "/following-sibling::div[contains(@class,'section-description')][1]"
        )
        if players_block:
            raw_text = " ".join(players_block.xpath(".//text()").getall())
            raw_text = re.sub(r"\s+", " ", raw_text).strip()
 
            # CASE 1: Companies have country codes → split by '),'
            if re.search(r"\([A-Z]{2}\)", raw_text):
                companies = [c.strip() for c in raw_text.split("),") if c.strip()]
 
                for company in companies:
                    if not company.endswith(")"):
                        company += ")"
                    company = re.sub(r"\s*\(", " (", company.strip())
                    company_profiles.append(company)
 
            # CASE 2: No country codes → split by commas and 'and'
            else:
                raw_text = re.sub(r"\s+and\s+", ", ", raw_text, flags=re.IGNORECASE)
                companies = [c.strip() for c in raw_text.split(",") if c.strip()]
 
                for company in companies:
                    company_profiles.append(company)
        return {
            "table_of_contents": table_of_contents,
            "company_profiles": company_profiles
        }
    
# deb start from here technavio

# ==================== TECHNAVIO ====================
    def parse_technavio(self, response):
        container = response.css("div.market-description")

        if not container:
            self.logger.warning("Technavio market-description not found")
            self.save_to_json({
                "title": self.clean_title(response.css("title::text").get()),
                "url": response.url,
                "table_of_contents": [],
                "company_profiles": []
            }, response.url)
            return

        # ----------- TITLE -----------
        title = container.css("h2::text").get()
        title = self.clean_title(title if title else response.css("title::text").get())

        # ----------- SEGMENTATION (FINAL — STRICT TECHNAVIO) -----------
        toc = []
        main_counter = 0

        # Find ONLY the correct segmentation UL block
        seg_ul = container.xpath(
            ".//h2[contains(translate(., 'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'segmented')]"
            "/following::ul[1][.//li/ul]"
        )

        # Fallback for OLD layout
        if not seg_ul:
            seg_ul = container.xpath(
                ".//*[self::h2 or self::h3]"
                "[contains(translate(., 'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'segment')]"
                "/following-sibling::ul[1]"
            )

        for ul in seg_ul:

            for li in ul.xpath("./li"):

                main_text = li.xpath("normalize-space(text())").get()
                if not main_text:
                    continue

                # Skip Geography block completely
                if re.search(r'^geography', main_text, re.I):
                    continue

                main_counter += 1
                toc.append(f"{main_counter}. {main_text}")

                sub_counter = 0

                # Only immediate sub-segments (prevents country extraction)
                for sub in li.xpath("./ul/li[not(./ul)]"):

                    sub_text = sub.xpath("normalize-space(text())").get()
                    if not sub_text:
                        continue

                    sub_counter += 1
                    toc.append(f"{main_counter}.{sub_counter}. {sub_text}")

        # ----------- COMPANY PROFILES (FINAL — ONLY REAL COMPANIES) -----------
        companies = []

        container = response.css("div.market-description")

        for heading in container.css("h2, h3"):

            heading_text = " ".join(heading.css("::text").getall()).lower()

            if not any(k in heading_text for k in [
                "vendor", "company", "player", "leading", "competitive", "key companies"
            ]):
                continue

            ul = heading.xpath("following::ul[1]")
            if not ul:
                continue

            for li in ul.css("li"):

                name = li.xpath("normalize-space(.)").get()
                if not name:
                    continue

                name = name.strip()

                # Remove sentence / paragraph fragments
                if re.search(r'\b(is|are|was|were|rise|rising|growth|market|forecast|driven|increasing)\b', name, re.I):
                    continue

                # Skip geography / segmentation junk
                if re.search(r'\b(region|country|market|analysis|snapshot|methodology|table of contents)\b', name, re.I):
                    continue

                # Skip long text (not company)
                if len(name.split()) > 8:
                    continue

                # ✔ Accept only company-like names
                if re.search(r'\b(inc|ltd|corp|plc|llc|group|holdings|co\.|sa|ag|gmbh)\b', name, re.I):
                    companies.append(name)
                    continue

                # ✔ Accept short brand names (EssilorLuxottica, KOIA, Crussh etc.)
                if 1 <= len(name.split()) <= 4 and not re.search(r'\d{4}|%', name):
                    companies.append(name)

        # ----------- FALLBACK (if heading missing) -----------
        if len(companies) < 5:
            for ul in container.css("ul"):

                items = [
                    li.xpath("normalize-space(.)").get().strip()
                    for li in ul.css("li")
                    if li.xpath("normalize-space(.)").get()
                ]


                valid = []
                for name in items:

                    if len(name.split()) > 8:
                        continue

                    if re.search(r'\b(is|are|was|were|growth|market|forecast|region|country|analysis)\b', name, re.I):
                        continue

                    if re.search(r'\b(inc|ltd|corp|plc|llc|group|holdings|co\.|sa|ag|gmbh)\b', name, re.I):
                        valid.append(name)
                        continue

                    if 1 <= len(name.split()) <= 4:
                        valid.append(name)

                if len(valid) >= 6:
                    companies.extend(valid)

        # Remove duplicates
        companies = list(dict.fromkeys(companies))

        # ----------- SAVE -----------
        self.save_to_json({
            "title": title,
            "url": response.url,
            "table_of_contents": toc,
            "company_profiles": companies
        }, response.url)
# deb stop here technavio

# deb start here straits
# ==================== STRAITS RESEARCH ====================
    def parse_straits(self, response):
        toc = []
        main_counter = 0

        # grab ALL main segmentation titles
        main_blocks = response.css("#tocdata li.main-points")

        for block in main_blocks:
            main_title = block.css("strong::text").get()
            if not main_title:
                continue

            main_title = re.sub(r"\(.*?\)", "", main_title).strip()

            # skip regional
            if "regional" in main_title.lower():
                continue

            match = re.search(r"\bmarket,\s*(.+)", main_title, re.I)
            if not match:
                continue

            segment_name = re.sub(r"^by\s+", "", match.group(1), flags=re.I).strip()

            main_counter += 1
            toc.append(f"{main_counter}. {segment_name}")

            # ---------- SUB SEGMENTS ----------
            sub_counter = 0

            # find nearest card-body after this main block
            card = block.xpath("./ancestor::div[contains(@class,'toc-head')]/following-sibling::div//li[contains(@class,'clr-black')]")

            for li in card:
                sub_text = li.xpath("normalize-space(text()[1])").get()
                if not sub_text:
                    continue

                # skip nested regional junk
                if "market" in sub_text.lower():
                    continue

                sub_counter += 1
                toc.append(f"{main_counter}.{sub_counter}. {sub_text}")

        yield {"table_of_contents": toc}

        # 🔁 Now request MAIN page to get company_profiles
        yield scrapy.Request(
            response.meta["main_url"],
            callback=self.parse_straits_companies,
            meta={
                "toc": toc,
                "main_url": response.meta["main_url"]
            }
        )
# ==================== COMPANY PROFILES (FINAL STRAITS RESEARCH) -----------
    def parse_straits_companies(self, response):
        toc = response.meta["toc"]
        company_profiles = []

        # Select all <li> under "List of Key and Emerging Players"
        company_nodes = response.xpath(
            "//h2[contains(translate(., 'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'players')]"
            "/following-sibling::ol[1]/li"
        )

        for li in company_nodes:
            # Extract ALL text inside <li> (works for <li>text</li> and <li><a>text</a></li>)
            company = li.xpath("normalize-space(.//text())").get()

            if company:
                company_profiles.append(company.strip())

        # Remove duplicates while keeping order
        company_profiles = list(dict.fromkeys(company_profiles))

        # Save JSON
        self.save_to_json({
            "title": self.clean_title(response.css("title::text").get()),
            "url": response.meta["main_url"],
            "table_of_contents": toc,
            "company_profiles": company_profiles
        }, response.meta["main_url"])

# deb stop here straits

#===================== Precedence Research ===================== 
    def parse_toc(self, response):
        report_code = response.css(
            'a[href*="/table-of-content/"]::attr(href)'
        ).re_first(r'/table-of-content/(\d+)')
        if not report_code:
            self.logger.warning("Report code not found")
            return
 
        toc_url = f"https://www.precedenceresearch.com/table-of-content/{report_code}"
        yield scrapy.Request(
            toc_url,
            callback=self.parse_precedence,
            meta={
                "title": self.clean_title(response.css('title::text').get()),
                "url": response.url,
            }
        )
 
    def parse_precedence(self, response):
        table_of_contents = []
        company_profiles = []
        segment_index = 0
        sub_index = 0
        current_segment = None
        in_company_section = False  # Flag for company profiles section
 
        for elem in response.css("div.toc-content > *"):
            tag = elem.root.tag
            text = elem.xpath("string(.)").get().strip()
            # ---------- SEGMENTS ----------
            if tag == "h2":
                if " By " in text:
                    # New segment
                    segment_index += 1
                    sub_index = 0
                    segment = text.split(" By ", 1)[1].strip().rstrip(",").upper()
                    current_segment = segment_index
                    in_company_section = False
                    table_of_contents.append(f"{segment_index}. {segment}")
 
                elif "Company Profiles" in text:
                    current_segment = None
                    in_company_section = True
 
                else:
                    current_segment = None
                    in_company_section = False
 
            # ---------- SUB-SEGMENTS ----------
            elif tag == "p":
                # Collect sub-segments under current segment
                if current_segment and re.match(r"^\d+\.\d+\.\d+\.?\s+", text):
                    # Remove numbering
                    sub_name = re.sub(r"^\d+\.\d+\.\d+\.?\s*", "", text).strip().upper()
                    # Skip noisy entries
                    if any(
                        bad in sub_name
                        for bad in [
                            "MARKET REVENUE",
                            "FORECAST",
                            "BY ",
                            "PRODUCT OFFERINGS",
                            "FINANCIAL",
                            "RECENT INITIATIVES",
                        ]
                    ):
                        continue
 
                    sub_index += 1
                    table_of_contents.append(f"{current_segment}.{sub_index}. {sub_name}")
                # Fallback: single <p> containing multiple subsegments
                elif current_segment and not re.match(r"^\d+\.\d+\.\d+\.?\s+", text):
                    all_matches = re.findall(
                        r"\d+\.\d+\.\d+\.\s*([A-Za-z &\-\(\)]+?)(?=\d+\.\d+\.|\Z)",
                        text
                    )
                    if len(all_matches) > 1:
                        for sub_name in all_matches:
                            sub_name = sub_name.strip().upper()
                            if not sub_name or sub_name == ".":
                                continue
                            if any(
                                bad in sub_name
                                for bad in [
                                    "MARKET REVENUE",
                                    "FORECAST",
                                    "BY ",
                                    "PRODUCT OFFERINGS",
                                    "FINANCIAL",
                                    "RECENT INITIATIVES",
                                ]
                            ):
                                continue
                            sub_index += 1
                            table_of_contents.append(
                                f"{current_segment}.{sub_index}. {sub_name}"
                            )
 
                # Collect company names
                elif in_company_section:
                    # Detect how many company-level patterns exist
                    company_number_matches = re.findall(r"\b\d+\.\d+\.", text)
 
                    # ---------------- COMPACT STRUCTURE ----------------
                    if len(company_number_matches) > 1:
                        cleaned_text = re.sub(
                            r"\d+\.\d+\.\d+\.\s*[A-Za-z &]+",
                            "",
                            text
                        )
                        # Extract only company-level (2-level)
                        matches = re.findall(
                            r"\b\d+\.\d+\.\s*([A-Za-z0-9 &\-\.,()]+)",
                            cleaned_text
                        )
                        for company_block in matches:
                            # Insert separator before next numbering
                            company_block = re.sub(r"(\d+\.\d+\.)", r"|\1", company_block)
                            split_companies = company_block.split("|")
                            for company_name in split_companies:
                                company_name = re.sub(r"^\d+\.\d+\.\s*", "", company_name).strip()
                                if company_name:
                                    company_profiles.append(company_name)
                    # ---------------- NORMAL STRUCTURE ----------------
                    elif re.match(r"^\d+\.\d+\.?\s+", text):
                        # Skip deeper levels like 10.1.1.
                        if re.match(r"^\d+\.\d+\.\d+\.?", text):
                            continue
                        company_name = re.sub(r"^\d+\.\d+\.?\s*", "", text).strip()
                        company_name = re.sub(r"<.*?>", "", company_name)
                        if company_name:
                            company_profiles.append(company_name)
 
        result_data = {
            "title": response.meta["title"],
            "url": response.meta["url"],
            "table_of_contents": table_of_contents,
            "company_profiles": company_profiles
        }
 
        self.save_to_json(result_data, response.meta["url"])

#===================== Varified Market Research =====================
    def parse_verified(self, response):
        segments_and_companies = self.parse_toc_company(response)
        result_data = {
            "title": self.clean_title(response.css('title::text').get()),
            "url": response.url,
            "table_of_contents": segments_and_companies.get("table_of_contents", []),
            "company_profiles": segments_and_companies.get("company_profiles", [])
        }
        self.save_to_json(result_data, response.url)
   
    def parse_toc_company(self, response):
        table_of_contents = []
        company_profiles = []
        segment_counter = 0
        toc_paragraphs = response.css("#vmr-ptable p")
        for p in toc_paragraphs:
            parts = [t.strip() for t in p.css("::text").getall() if t.strip()]
            if not parts:
                continue
 
            raw_heading = re.sub(r"<.*?>", "", parts[0]).strip()
            heading_upper = raw_heading.upper().replace("\xa0", " ")
            heading_upper = re.sub(r"\s+", " ", heading_upper)
            # ================= COMPANY PROFILES =================
            if "COMPANY PROFILES" in heading_upper:
                strong_tags = p.xpath(".//strong")
                # normal structure
                sec_match = re.match(r"^(\d+)", heading_upper)
                company_section_no = sec_match.group(1) if sec_match else None
                for text in parts:
                    text = text.replace("\xa0", " ")
                    text = re.sub(r"\s+", " ", text).strip()
                    text_upper = text.upper()
                    if "OVERVIEW" in text_upper:
                        continue
                    if company_section_no:
                        if not re.match(rf"^{company_section_no}\s*\.\s*\d+", text):
                            continue
                        if re.match(rf"^{company_section_no}\s*\.\s*\d+\s*\.\s*\d+", text):
                            continue
                    clean_company = re.sub(
                        rf"^{company_section_no}\s*\.\s*\d+\s*",
                        "",
                        text
                    ).strip()
                    clean_company = re.sub(r"^[\.\-\•\u2022]+\s*", "", clean_company).strip()
                    if clean_company and len(clean_company.split()) <= 8:
                        company_profiles.append(clean_company)
                # check next <p> for bullet company names =====
                if not company_profiles:
                    next_p = p.xpath("following-sibling::p[1]")
                    if next_p:
                        next_parts = [
                            t.replace("\xa0", " ").strip()
                            for t in next_p.css("::text").getall()
                            if t.strip()
                        ]
                        for text in next_parts:
                            text = re.sub(r"\s+", " ", text).strip()
                            text_upper = text.upper()
                            if "OVERVIEW" in text_upper:
                                continue
                            clean_company = re.sub(
                                r"^[•\-\u2022]?\s*",
                                "",
                                text
                            ).strip()
                            clean_company = re.sub(r"^[\.\-\•\u2022]+\s*", "", clean_company).strip()
                            if clean_company and len(clean_company.split()) <= 8:
                                company_profiles.append(clean_company)
 
                # Fallback for messy structure with multiple strong tags
                if not company_profiles and len(strong_tags) > 1:
                    parts = [
                        t.replace("\xa0", " ").strip()
                        for t in p.css("::text").getall()
                        if t.strip()
                    ]
                    for text in parts:
                        text = re.sub(r"\s+", " ", text).strip()
                        text_upper = text.upper()
                        if "COMPANY PROFILES" in text_upper:
                            continue
                        if "OVERVIEW" in text_upper:
                            continue
                        if re.match(r"^\d+\s+[A-Z]", text_upper) and not text_upper.startswith(company_section_no):
                            break
                        num_match = re.match(
                            rf"^{company_section_no}\s*\.\s*(\d+)\s+(.*)",
                            text
                        )
                        if num_match:
                            clean_company = num_match.group(2).strip()
                            clean_company = re.sub(r"^[\.\-\•\u2022]+\s*", "", clean_company).strip()
                            company_profiles.append(clean_company)
                continue
            # ================= SEGMENTS =================
            # Detect messy structure
            strong_tags = p.xpath(".//strong")
            if len(strong_tags) > 1:
                parts = [
                    t.replace("\xa0", " ").strip()
                    for t in p.css("::text").getall()
                    if t.strip()
                ]
                for idx, text in enumerate(parts):
                    text_upper = re.sub(r"\s+", " ", text.upper()).strip()
                    if "BY GEOGRAPHY" in text_upper:
                        break
                    # Detect segment like: 5 MARKET, BY TYPE
                    seg_match = re.match(r"^(\d+)\s+.*?,\s*BY\s+(.*)", text_upper)
                    if seg_match:
                        segment_no = seg_match.group(1)
                        segment_name = seg_match.group(2).strip()
                        segment_counter += 1
                        sub_counter = 0
                        table_of_contents.append(f"{segment_counter}. {segment_name}")
 
                        # Extract subsegments inside same p
                        for sub_text in parts[idx + 1:]:
                            sub_text_upper = re.sub(r"\s+", " ", sub_text.upper()).strip()
                            # Stop when next main section starts
                            if re.match(r"^\d+\s+", sub_text_upper) and not sub_text_upper.startswith(segment_no + "."):
                                break
                            sub_match = re.match(rf"^{segment_no}\s*\.\s*(\d+)\s+(.*)", sub_text_upper)
                            if sub_match:
                                sub_name = sub_match.group(2).strip()
                                if "OVERVIEW" in sub_name:
                                    continue
                                if len(sub_name.split()) >= 8:
                                    continue
                                sub_counter += 1
                                table_of_contents.append(
                                    f"{segment_counter}.{sub_counter}. {sub_name}"
                                )
                continue                    
            # Normal structure with single strong tag                    
            if " BY " not in heading_upper:
                continue
 
            if "BY GEOGRAPHY" in heading_upper:
                continue
            segment_name = heading_upper.split(" BY ", 1)[1].strip()
            segment_counter += 1
            sub_counter = 0
            table_of_contents.append(f"{segment_counter}. {segment_name}")
            for text in parts[1:]:
                text_upper = text.upper()
 
                if "OVERVIEW" in text_upper:
                    continue
 
                clean_text = re.sub(r"^[•\-\u2022]?\s*", "", text_upper)
                clean_text = re.sub(r"^\d+(\.\d+)*\s*", "", clean_text)
                if not clean_text or len(clean_text.split()) > 8:
                    continue
 
                sub_counter += 1
                table_of_contents.append(
                    f"{segment_counter}.{sub_counter}. {clean_text}"
                )
            # check next <p> for bullet items =====
            if sub_counter == 0:
                next_p = p.xpath("following-sibling::p[1]")
                if next_p:
                    next_parts = [
                        t.strip()
                        for t in next_p.css("::text").getall()
                        if t.strip()
                    ]
                    for text in next_parts:
                        text_upper = text.upper()
 
                        if "OVERVIEW" in text_upper:
                            continue
 
                        clean_text = re.sub(r"^[•\-\u2022]?\s*", "", text_upper)
                        clean_text = re.sub(r"^\d+(\.\d+)*\s*", "", clean_text)
                        if not clean_text or len(clean_text.split()) > 8:
                            continue
 
                        sub_counter += 1
                        table_of_contents.append(
                            f"{segment_counter}.{sub_counter}. {clean_text}"
                        )    
        # ================= FALLBACK LOGIC =================
        if not table_of_contents and not company_profiles:
            segment_counter = 0
            sub_counter = 0
            current_segment = None
            in_company_section = False
            company_section_no = None  
            for node in response.css("#vmr-ptable ::text"):
                text = node.get().strip()
                if not text:
                    continue
                text = re.sub(r"^[•·\-\u2022\u00B7]+\s*", "", text)
                text_upper = text.upper()
                # ---------- DETECT COMPANY PROFILES HEADER ----------
                m = re.match(r"^(\d+)\.?\s+COMPANY\s+PROFILE(S)?", text_upper)
                if m:
                    company_section_no = m.group(1)
                    in_company_section = True
                    continue
                # ---------- INSIDE COMPANY SECTION ----------
                if in_company_section:
 
                    if re.match(rf"^(?!{company_section_no}(\.|$))\d+(\.|$|\s)", text):
                        in_company_section = False
                        company_section_no = None
                        continue
 
                    if re.match(rf"^{company_section_no}\.\d+\.\d+", text):
                        continue
 
                    if re.match(rf"^{company_section_no}\.\d+\.?\s+", text):
                        clean_company = re.sub(
                            rf"^{company_section_no}\.\d+\s*", "", text
                        ).strip()
                        clean_company = re.sub(r"^\.\s*", "", clean_company)
                        if clean_company:
                            company_profiles.append(clean_company)
                        continue
 
                    bullet_clean = text.strip()
 
                    if re.match(r"^\d+(\.|$|\s)", bullet_clean):
                        continue
 
                    if re.search(
                        r"(OVERVIEW|FINANCIAL|OUTLOOK|DEVELOPMENT|PROFILE|APPENDIX|SUMMARY)",
                        bullet_clean.upper()
                    ):
                        continue
 
                    if len(bullet_clean.split()) > 6:
                        continue
 
                    company_profiles.append(bullet_clean)
                    continue
 
                # ---------- FALLBACK SEGMENTS ----------
                segment_match = re.match(
                    r"^(\d+)\.?\s+.*?\s*,?\s+BY\s+(.+)",
                    text_upper
                )
                if segment_match:
                    segment_name = segment_match.group(2).strip()
                    # Stop geography section
                    if "GEOGRAPHY" in segment_name:
                        current_segment = None
                        continue
                    segment_counter += 1
                    sub_counter = 0
                    current_segment = segment_counter
                    table_of_contents.append(f"{segment_counter}. {segment_name}")
                    continue
                if " BY GEOGRAPHY" in text_upper:
                    current_segment = None
                    continue
 
                if current_segment:
                    if re.match(r"^\d+\.\s+", text) and " BY " not in text_upper:
                        current_segment = None
                        continue
 
                    if "OVERVIEW" in text_upper:
                        continue
 
                    clean_text = re.sub(r"^\d+(\.\d+)*\s*", "", text_upper)
                    clean_text = re.sub(r"^\.\s*", "", clean_text)
                    clean_text = re.sub(r"^[•·\u2022\u00B7]\s*", "", clean_text)
 
                    if len(clean_text.split()) > 8:
                        continue
 
                    sub_counter += 1
                    table_of_contents.append(
                        f"{current_segment}.{sub_counter}. {clean_text}"
                    )
                   
        return {
            "table_of_contents": table_of_contents,
            "company_profiles": company_profiles
        }
    
    #===================== Verified Market Reports =====================
    def parse_verified_reports(self, response):
        segments = self.extract_segments(response).get("table_of_contents", [])
        company_profiles = self.extract_top_companies(response).get("company_profiles", [])
        result_data = {
            "title": self.clean_title(response.css('title::text').get()),
            "url": response.url,
            "table_of_contents": segments,
            "company_profiles": company_profiles
        }
        self.save_to_json(result_data, response.url)
 
    def extract_segments(self, response):
        table_of_contents = []
        segment_counter = 0
        # -------- PRIMARY METHOD (inside container) --------
        segmentation_h2 = response.xpath(
            '//div[contains(@class,"container")]'
            '//h2[contains(translate(.,"ABCDEFGHIJKLMNOPQRSTUVWXYZ","abcdefghijklmnopqrstuvwxyz"),"market segmentation")]'
        )
        # -------- FALLBACK METHOD (global search) --------
        if not segmentation_h2:
            print("Fallback triggered for Market Segmentation")
            segmentation_h2 = response.xpath(
                '//h2[contains(translate(.,"ABCDEFGHIJKLMNOPQRSTUVWXYZ","abcdefghijklmnopqrstuvwxyz"),"market segmentation")]'
            )
        if not segmentation_h2:
            return {"table_of_contents": table_of_contents}
        # Get all h3 after segmentation h2
        h3_segments = segmentation_h2.xpath("following::h3")
        for h3 in h3_segments:
            heading = h3.xpath("normalize-space(.)").get()
            if not heading:
                continue
            heading_lower = heading.lower()
            # STOP condition
            if "by geography" in heading_lower:
                break
            # Only process headings containing "by"
            if re.search(r"\bby\b", heading, re.IGNORECASE):
                segment_counter += 1
                segment_title = re.sub(
                    r".*\bby\b", "", heading,
                    flags=re.IGNORECASE
                ).strip().upper()
                table_of_contents.append(
                    f"{segment_counter}. {segment_title}"
                )
                # Get first ul after this h3
                sub_items = h3.xpath("following-sibling::ul[1]/li")
                sub_counter = 0
                for li in sub_items:
                    sub_text = li.xpath("normalize-space(.)").get()
                    if not sub_text:
                        continue
 
                    sub_counter += 1
                    table_of_contents.append(
                        f"{segment_counter}.{sub_counter}. {sub_text.upper()}"
                    )
 
        return {"table_of_contents": table_of_contents}
 
    def extract_top_companies(self, response):
        companies = []
        container = response.xpath('//div[contains(@class,"container")]')
        # Find h2 that contains "Top" and "Companies" inside container
        heading = container.xpath(
            './/h2[contains(translate(.,"ABCDEFGHIJKLMNOPQRSTUVWXYZ","abcdefghijklmnopqrstuvwxyz"),"top") '
            'and contains(translate(.,"ABCDEFGHIJKLMNOPQRSTUVWXYZ","abcdefghijklmnopqrstuvwxyz"),"companies")]'
        )
        # -------- FALLBACK (global search if not found inside container) --------
        if not heading:
            print("Company fallback triggered")
            heading = response.xpath(
                '//h2[contains(translate(.,"ABCDEFGHIJKLMNOPQRSTUVWXYZ","abcdefghijklmnopqrstuvwxyz"),"top") '
                'and contains(translate(.,"ABCDEFGHIJKLMNOPQRSTUVWXYZ","abcdefghijklmnopqrstuvwxyz"),"companies")]'
            )
        if heading:
            # Get first ul after that h2 (ignores the p tags automatically)
            company_list = heading.xpath('following::ul[1]/li')
            for li in company_list:
                name = li.xpath('normalize-space(.)').get()
                if name:
                    companies.append(name.strip())
        return {"company_profiles": companies}

# ==================== Global Mraket Insights ====================
    def parse_global_market(self, response):
        segments = self.extract_segment_toc(response).get("table_of_contents", [])
        company = self.extract_company_profiles(response).get("company_profiles", [])
        result_data = {
            "title": self.clean_title(response.css('title::text').get()),
            "url": response.url,
            "table_of_contents": segments,
            "company_profiles": company
        }
        self.save_to_json(result_data, response.url)
 
    def extract_segment_toc(self, response):
        table_of_contents = []
        segment_counter = 0
        main_div = response.css("div.tab-content")
 
        for h4 in main_div.css("h4"):
            heading = " ".join(h4.css("::text").getall()).strip()
            if not heading:
                continue
 
            # Only headings that contain "by"
            if re.search(r"\bby\b", heading, re.IGNORECASE):
                segment_counter += 1
                cleaned_heading = re.split(r"\bby\b", heading, flags=re.IGNORECASE)[-1]
                cleaned_heading = cleaned_heading.strip(" ,:-").strip().upper()
 
                toc_entry = f"{segment_counter}. {cleaned_heading}"
                table_of_contents.append(toc_entry)
 
                ul = h4.xpath("following::ul[1]")
                if not ul:
                    continue
 
                sub_counter = 0
 
                for li in ul.xpath("./li"):
                    sub_text = "".join(li.xpath("./text()").getall()).strip()
                    if not sub_text:
                        continue
 
                    sub_counter += 1
                    sub_entry = f"{segment_counter}.{sub_counter}. {sub_text.upper()}"
                    table_of_contents.append(sub_entry)
 
                    nested_ul = li.xpath("./ul")
                    if nested_ul:
                        sub_sub_counter = 0
 
                        for sub_li in nested_ul.xpath("./li"):
                            sub_sub_text = "".join(
                                sub_li.xpath("./text()").getall()
                            ).strip()
 
                            if not sub_sub_text:
                                continue
 
                            sub_sub_counter += 1
                            sub_sub_entry = (
                                f"{segment_counter}.{sub_counter}.{sub_sub_counter}. "
                                f"{sub_sub_text.upper()}"
                            )
                            table_of_contents.append(sub_sub_entry)
 
                            # ---------------- NEW: SUB-SUB-SUB ----------------
                            nested_ul_3 = sub_li.xpath("./ul")
                            if nested_ul_3:
                                sub_sub_sub_counter = 0
 
                                for sub_sub_li in nested_ul_3.xpath("./li"):
                                    sub_sub_sub_text = "".join(
                                        sub_sub_li.xpath(".//text()").getall()
                                    ).strip()
 
                                    if not sub_sub_sub_text:
                                        continue
 
                                    sub_sub_sub_counter += 1
                                    sub_sub_sub_entry = (
                                        f"{segment_counter}.{sub_counter}."
                                        f"{sub_sub_counter}.{sub_sub_sub_counter}. "
                                        f"{sub_sub_sub_text.upper()}"
                                    )
 
                                    table_of_contents.append(sub_sub_sub_entry)
 
        # ---------------- FALLBACK: P > STRONG ----------------
        if not table_of_contents:
            for tag in main_div.css("p strong"):
                heading = " ".join(tag.css("::text").getall()).strip()
                if not heading:
                    continue
                if re.search(r"\bby\b", heading, re.IGNORECASE):
                    segment_counter += 1
                    cleaned_heading = re.split(r"\bby\b", heading, flags=re.IGNORECASE)[-1]
                    cleaned_heading = cleaned_heading.strip(" ,:-").strip().upper()
                    toc_entry = f"{segment_counter}. {cleaned_heading}"
                    table_of_contents.append(toc_entry)
                    ul = tag.xpath("following::ul[1]")
                    if not ul:
                        continue
 
                    sub_counter = 0
                    for li in ul.xpath("./li"):
                        sub_text = "".join(li.xpath("./text()").getall()).strip()
                        if not sub_text:
                            continue
 
                        sub_counter += 1
                        sub_entry = f"{segment_counter}.{sub_counter}. {sub_text.upper()}"
                        table_of_contents.append(sub_entry)
                        nested_ul = li.xpath("./ul")
                        if nested_ul:
                            sub_sub_counter = 0
                            for sub_li in nested_ul.xpath("./li"):
                                sub_sub_text = "".join(
                                    sub_li.xpath(".//text()").getall()
                                ).strip()
 
                                if not sub_sub_text:
                                    continue
 
                                sub_sub_counter += 1
                                sub_sub_entry = (
                                    f"{segment_counter}.{sub_counter}.{sub_sub_counter}. "
                                    f"{sub_sub_text.upper()}"
                                )
                                table_of_contents.append(sub_sub_entry)
 
        return {"table_of_contents": table_of_contents}
 
    def extract_company_profiles(self, response):
        company_profiles = []
        main_div = response.css("div.tab-content")
 
        # ---------------- PRIMARY: H3 ----------------
        for h3 in main_div.css("h3"):
            heading = " ".join(h3.css("::text").getall()).strip()
            if not heading:
                continue
            if "companies" in heading.lower() or "company" in heading.lower():
                ul = h3.xpath("following::ul[1]")
                if not ul:
                    continue
 
                for li in ul.xpath(".//li"):
                    company = "".join(
                        li.xpath(".//text()").getall()
                    ).replace("\xa0", " ").strip()
 
                    if not company:
                        continue
 
                    # Word count filter
                    if len(company.split()) <= 8:
                        company_profiles.append(company)
 
        # ---------------- FALLBACK 1: H2 ----------------
        if not company_profiles:
            for h2 in main_div.css("h2"):
                heading = " ".join(h2.css("::text").getall()).strip()
                if not heading:
                    continue
 
                if "companies" in heading.lower() or "company" in heading.lower():
                    ul = h2.xpath("following::ul[1]")
                    if not ul:
                        continue
 
                    for li in ul.xpath(".//li"):
                        company = "".join(
                            li.xpath(".//text()").getall()
                        ).replace("\xa0", " ").strip()
                        if not company:
                            continue
 
                        if len(company.split()) <= 8:
                            company_profiles.append(company)
 
        # ---------------- FALLBACK 2: P TEXT ----------------
        if not company_profiles:
            for p in main_div.css("p"):
                paragraph_text = " ".join(p.css("::text").getall()).strip()
 
                if not paragraph_text:
                    continue
 
                if re.search(r"\b(major players|major companies|key players)\b",
                            paragraph_text,
                            re.IGNORECASE):
                    ul = p.xpath("following::ul[1]")
                    if not ul:
                        continue
 
                    for li in ul.xpath(".//li"):
                        company = "".join(
                            li.xpath(".//text()").getall()
                        ).replace("\xa0", " ").strip()
 
                        if company and len(company.split()) <= 8:
                            company_profiles.append(company)
 
                    break  # stop after first valid paragraph section                
        # Remove duplicates
        company_profiles = list(dict.fromkeys(company_profiles))
        return {"company_profiles": company_profiles}
    # ==================== COMMON HELPER FUNCTIONS ==================== #
    def is_likely_company_name(self, text):
        """Check if text is likely a company name"""
        if not text or len(text) < 2:
            return False
 
        text_clean = text.strip()
        text_lower = text_clean.lower()
 
        if text_lower.startswith('by '):
            return False
 
        # NEW: block headers like "Region: Description"
        if ':' in text_clean:
            return False
       
        # HARD STOP: sentences always contain verbs
        sentence_verbs = r'\b(is|are|was|were|be|been|being|has|have|had|driven|growing|expected|leveraging|adopting|using)\b'
        if re.search(sentence_verbs, text_lower) and len(text_clean.split()) > 5:
            return False
 
        # Allow numeric brand names like "3M"
        if re.fullmatch(r'\d+[A-Z]+', text_clean):
            return True
       
        # 1️⃣ STRONG INCLUDE FIRST (company indicators)
        company_indicators = [
            r'\b(inc|incorporated|llc|ltd|limited|corp|corporation|co|company|group|international|holdings|technologies|industries)\b',
            # Only match 2–6 capitalized words for names (avoid full sentences)
            r'^(?:[A-Z][A-Za-z&.\-]*\.?)(?:\s+(?:[A-Z][A-Za-z&.\-]*\.?)){1,5}$',
            r'\b([A-Z][a-z]+ & [A-Z][a-z]+)\b',        # Ampersand names
        ]
 
        for pattern in company_indicators:
            if re.search(pattern, text_clean, re.IGNORECASE):
                return True   # accept immediately
 
        # 2️⃣ STRONG EXCLUDE NEXT
        exclude_patterns = [
            r'.*%.*',
            r'.*CAGR.*',
            r'.*share.*',
            r'.*forecast.*',
            r'.*region.*',
            r'.*segment.*',
            r'.*application.*',
            r'.*market.*',
            r'.*revenue.*',
            r'.*volume.*',
            r'^\d',
            r'.*©.*',
            r'.*privacy.*',
            r'.*cookie.*',
            r'.*terms.*',
            r'.*condition.*',
            r'^[^a-zA-Z]*$',
            r'^the\s',
            r'.*\bis\b.*',
            r'.*\bare\b.*',
            r'.*\bwill\b.*',
            r'^by\s',
            r'.*\bexpected\b.*',
            r'.*\bgrow\b.*',
            r'.*\bgrowth\b.*',
            r'.*\bgrowing\b.*',
            r'.*\bfrom\s+\d{4}\b.*',
            r'.*\bto\s+\d{4}\b.*',
            r'.*\bindustry\b.*',
        ]
        for pattern in exclude_patterns:
            if re.search(pattern, text_lower):
                return False
           
        # 3️⃣ WEAK FALLBACK (short, non-sentence text)
        if len(text_clean.split()) > 6:
            return False
 
        return len(text_clean) < 50 and any(char.isalpha() for char in text_clean)
 
    def clean_title(self, title):
        """Clean the title by removing everything after 'Market'"""
        if not title:
            return title
         
        # Find the position of "Market" (case insensitive)
        market_match = re.search(r"^(.*\bMarket\b)", title, re.IGNORECASE) ##Modified##
        if market_match:
            # Return everything up to and including "Market"
            return market_match.group(1).strip()
     
        # If "Market" not found, return the original title
        return title
 
    def clean_toc_entry(self, text):
        """Clean TOC entry by removing everything after 'By' and parentheses content"""
        # Remove content within parentheses including the parentheses themselves
        cleaned_text = re.sub(r'\s*\([^)]*\)', '', text)
     
        # Remove everything after "By" (case insensitive)
        cleaned_text = re.sub(r'\s+By\s+.*$', '', cleaned_text, flags=re.IGNORECASE)
     
        return cleaned_text.strip()
 
    ##Modified function##
    def clean_toc_line(self, line):
        """Clean a complete TOC line (with numbering)"""
        match = re.match(r'^(\d+(?:\.\d+)*)\.\s*(.*)$', line)
        if not match:
            return line
        number, content = match.groups()
        # Base cleaning
        cleaned_content = self.clean_toc_entry(content)
        # ---- ONLY for segment titles (1., 2., 3.) ----
        if number.isdigit() and hasattr(self, "market_name") and self.market_name:
            # Remove market name from beginning
            cleaned_content = re.sub(
                rf'^{re.escape(self.market_name)}\s+',
                '',
                cleaned_content,
                flags=re.IGNORECASE
            )
            # Remove trailing "Outlook"
            cleaned_content = re.sub(
                r'\s+Outlook$',
                '',
                cleaned_content,
                flags=re.IGNORECASE
            )
        return f"{number}. {cleaned_content.strip()}"
 
    def is_regional_section(self, text):
        """Check if a section title indicates regional content"""
        # More specific patterns to avoid false positives
        regional_patterns = [
            r'regional outlook.*volume.*revenue', # Pattern like "Regional Outlook (Volume, Kilotons; Revenue, USD Million)"
            r'.*regional outlook$', # Ends with "Regional Outlook"
            r'by region',
            r'by country',
            r'country outlook',
            r'Regional',
            r'north america$', 'europe$', 'asia pacific$', 'latin america$', 'middle east$', 'africa$',
            r'emea$', 'apac$',
            r'u\.s\.$', 'us$', 'united states$', 'canada$', 'mexico$', 'germany$', 'uk$', 'united kingdom$',
            r'france$', 'italy$', 'spain$', 'china$', 'india$', 'japan$', 'south korea$', 'brazil$',
            r'australia$', 'russia$', 'saudi arabia$', 'uae$',
            r'region outlook' ## Modification added here ##
        ]
     
        text_lower = text.lower().strip()
     
        # Check if it's specifically a regional section (not just containing "outlook")
        for pattern in regional_patterns:
            if re.search(pattern, text_lower):
                return True
        return False
 
    def is_regional_line(self, line):
        """Check if a line contains regional content"""
        # Remove the numbering part to check just the content
        content = re.sub(r'^\d+(?:\.\d+)*\.\s*', '', line)
        return self.is_regional_section(content)
 
    def is_valid_company_name(self, name):
        """Check if the name is likely a valid company name"""
        # Exclude names that are too short
        if len(name) < 3:
            return False
         
        # Exclude names that are clearly not companies
        excluded_terms = [
            'by ','of','at', 'tier', 'size', 'standard', 'absorption', 'utilized',
            'colocation', 'end-user', 'other', 'small', 'medium', 'large',
            'mega', 'hyperscale', 'retail', 'wholesale', 'threat', 'bargaining',
            'power', 'substitutes', 'entrants', 'suppliers', 'buyers',
            'expansion', 'fragmented', 'government', 'high', 'labor',
            'mechanized', 'rising', 'valorization', 'rest of', 'counterfeit',
            'intensity of competitive rivalry', 'convenience stores',
            'conventional', 'organic', 'processed', 'raw', 'specialty stores',
            'supermarkets and hypermarkets', 'online retail', 'competitive rivalry',
            'edge-ready', 'metro fibre', 'densification', 'mmr', 'incentives',
            'maharashtra', 'policy', 'monsoon-driven', 'flooding risk', 'crz clearances',
            'redevelopment', 'brown-field', 'textile mills', 'dc campuses',
            'scarcity', 'contiguous', 'acre parcels', 'competitive', 'rivalry'
        ]
     
        for term in excluded_terms:
            if term in name.lower():
                return False
             
        # Exclude country and region names
        countries_regions = [
            'africa', 'asia', 'europe', 'america', 'australia', 'china', 'india',
            'japan', 'germany', 'france', 'uk', 'usa', 'canada', 'brazil', 'russia',
            'mexico', 'italy', 'spain', 'korea', 'egypt', 'saudi', 'uae', 'argentina',
            'new zealand', 'south africa', 'middle east', 'north america', 'south america'
        ]
     
        for region in countries_regions:
            if region in name.lower():
                return False
             
        # Valid company names typically start with a capital letter
        if not re.match(r'^[A-Z]', name):
            return False
         
        # Should contain at least one letter
        if not re.search(r'[A-Za-z]', name):
            return False
         
        # Should not be a single common word (like "Organic", "Raw", etc.)
        common_words = [
            'organic', 'raw', 'processed', 'conventional', 'online', 'retail',
            'stores', 'hypermarkets', 'supermarkets', 'convenience', 'specialty',
            'competitive', 'rivalry', 'edge', 'metro', 'fibre', 'mmr', 'incentives',
            'policy', 'monsoon', 'flooding', 'risk', 'crz', 'clearances',
            'redevelopment', 'brown', 'field', 'textile', 'mills', 'dc', 'campuses',
            'scarcity', 'contiguous', 'acre', 'parcels'
        ]
        if name.lower() in common_words:
            return False
         
        # Should contain company indicators (LLC, Ltd, Inc, Corp, etc.)
        company_indicators = [
            'llc', 'ltd', 'inc', 'corp', 'corporation', 'company', 'co.',
            'group', 'holdings', 'international', 'pvt', 'limited', 'services',
            'technologies', 'infrastructure', 'data centers', 'datacenters', 'web services'
        ]
     
        has_company_indicator = any(indicator in name.lower() for indicator in company_indicators)
     
        # If it doesn't have a company indicator, check if it's a multi-word name
        # that's likely a company (not a common phrase)
        if not has_company_indicator:
            words = name.split()
            if len(words) < 2:
                return False
             
            # Check if it's a common phrase that shouldn't be included
            common_phrases = [
                'convenience stores', 'specialty stores', 'supermarkets and hypermarkets',
                'online retail', 'intensity of competitive rivalry', 'counterfeit imports',
                'edge-ready metro fibre densification across mmr',
                'incentives under maharashtra itites policy 2026',
                'monsoon-driven flooding risk and mandatory crz clearances',
                'redevelopment of brown-field textile mills into dc campuses',
                'scarcity of contiguous 50-acre parcels inside mmr'
            ]
         
            if any(phrase in name.lower() for phrase in common_phrases):
                return False
             
        return True
 
    # ==================== COMMON SAVE FUNCTION ====================
    def save_to_json(self, data, url):
        # Use the cleaned title from the data dictionary
        cleaned_market_name = self.sanitize_filename(data['title'])
        
        # Extract domain from URL for unique file naming
        domain = self.extract_domain_from_url(url)
        
        # Format: <market_name>__<domain>.json
        filename = f"{cleaned_market_name}__{domain}.json"
        output_path = os.path.join(self.output_dir, filename)
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=4, ensure_ascii=False)
        self.logger.info(f"✅ Saved: {output_path}")
 
    def closed(self, reason):
        global CURRENT_MARKET_INPUTS
        if self.no_docx:
            return
        #--------------------------- DOCX GENERATION --------------------
        # Generate DOCX for each JSON file after scraping is done
        for filename in os.listdir(self.output_dir):
            if filename.endswith('.json'):
                json_path = os.path.join(self.output_dir, filename)
                with open(json_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                # Set market inputs from data
                CURRENT_MARKET_INPUTS = data.get('market_inputs', {})
                
                docx_path = os.path.join(self.output_dir, filename.replace('.json', '.docx'))
                generate_docx_from_data(data, docx_path)
                self.logger.info(f"✅ Generated DOCX: {docx_path}")
        
        # Print dominating regions after DOCX generation
        self.print_dominating_regions()
    
    def print_dominating_regions(self):
        """Print the first, second, and third dominating regions"""
        global REGIONS_MAPPING
        
        if not REGIONS_MAPPING or REGIONS_MAPPING == {}:
            self.logger.info("❌ No regions data available")
            return
        
        print("\n" + "="*70)
        print("🌍 DOMINATING REGIONS SUMMARY 🌍".center(70))
        print("="*70 + "\n")
        
        # First Region
        first_region = REGIONS_MAPPING.get("dominant", {})
        if first_region:
            region_name = first_region.get("region", "N/A")
            countries = first_region.get("countries", [])
            if isinstance(countries, list) and countries and isinstance(countries[0], dict):
                countries_list = [c["name"] for c in countries]
            elif isinstance(countries, list):
                countries_list = countries
            else:
                countries_list = [countries]
            
            print(f"📍 FIRST REGION (DOMINANT):")
            print(f"   Region: {region_name}")
            print(f"   Countries: {', '.join(countries_list)}\n")
        
        # Second Region
        second_region = REGIONS_MAPPING.get("second", {})
        if second_region:
            region_name = second_region.get("region", "N/A")
            countries = second_region.get("countries", [])
            if isinstance(countries, list) and countries and isinstance(countries[0], dict):
                countries_list = [c["name"] for c in countries]
            elif isinstance(countries, list):
                countries_list = countries
            else:
                countries_list = [countries]
            
            print(f"📍 SECOND REGION:")
            print(f"   Region: {region_name}")
            print(f"   Countries: {', '.join(countries_list)}\n")
        
        # Third Region
        third_region = REGIONS_MAPPING.get("third", {})
        if third_region:
            region_name = third_region.get("region", "N/A")
            countries = third_region.get("countries", [])
            if isinstance(countries, list) and countries and isinstance(countries[0], dict):
                countries_list = [c["name"] for c in countries]
            elif isinstance(countries, list):
                countries_list = countries
            else:
                countries_list = [countries]
            
            print(f"📍 THIRD REGION:")
            print(f"   Region: {region_name}")
            print(f"   Countries: {', '.join(countries_list)}\n")
        
        print("="*70 + "\n")
# rd_doc.py functions integrated below #market insights para
def get_gpt5_insight(market_name, tech="AI"):
    """Fetches GPT-5-mini response for how AI/IoT is transforming the market"""
    prompt = f"""
    Write a detailed 2-paragraph analysis (100-100 words for each paragraph) about {market_name.lower()}.
    Act like a research analyst, Ensure the paragraphs flow smoothly with natural transitions rather than listing concepts using commas.
    Do NOT use markdown symbols (#, *, -, •) anywhere.
    Integrate key factors, explanations, and real implications instead of simply stacking statements.
    Maintain logical progression between paragraphs so each leads into the next.
    Paragraph 1: Explain in depth primary driver , Provide a contextual overview, describing what the market is, why it matters, and how it has developed over time with examples.
    Paragraph 2: Explain in depth key factor of the Global {market_name} Market the major growth factors and opportunities with clear reasoning, real-world use cases, and supporting details.
    The explanation should be in a cause-and-effect manner. 
    Do not use any subheadings and avoid repetitive sentence structures.
    Write in a human-like analytical tone with smooth transitions.
    """
    response = client.chat.completions.create(
        model="gpt-5-mini",
        messages=[{"role": "user", "content": prompt}],
    )
    return response.choices[0].message.content.strip()

def remove_urls(text):
    """Remove URLs from text"""
    return re.sub(r'\(?https?://[^\s)]+\)?|[\(\[]\s*[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\s*[\)\]]','', text)

def remove_empty_parentheses(text):
    """Remove empty parentheses from text while preserving content in non-empty parentheses"""
    # Remove empty parentheses: () with optional spaces inside
    text = re.sub(r'\s*\(\s*\)\s*', ' ', text)
    # Clean up multiple spaces
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def clean_sentence_end(text):
    """Clean trailing empty parentheses and brackets from text"""
    return text.rstrip().rstrip('()[] ').strip()

def clean_all_paragraphs(text):
    """Apply clean_sentence_end and remove empty parentheses from every paragraph in text"""
    paragraphs = text.split('\n')
    cleaned_paragraphs = [remove_empty_parentheses(clean_sentence_end(para)) if para.strip() else para for para in paragraphs]
    return '\n'.join(cleaned_paragraphs)


def capitalize_first_alpha(s: str) -> str:
    """Capitalize the first alphabetical character in a string without changing the rest.

    Examples:
      'how are you?' -> 'How are you?'
      ' "what now?"' -> ' "What now?"'
    """
    if not s:
        return s
    for i, ch in enumerate(s):
        if ch.isalpha():
            return s[:i] + ch.upper() + s[i+1:]
    return s


def capitalize_questions_in_text(text: str) -> str:
    """Find questions (substrings ending with '?') and capitalize their first alphabetical character.

    This targets only questions and leaves non-question sentences unchanged.
    """
    if not text:
        return text

    pattern = re.compile(r'[^?]*\?')
    parts = []
    last = 0
    for m in pattern.finditer(text):
        start, end = m.span()
        parts.append(text[last:start])
        q = m.group(0)
        # Capitalize first alpha in q
        parts.append(capitalize_first_alpha(q))
        last = end
    parts.append(text[last:])
    return ''.join(parts)

def get_gpt5_question_and_answer(market_name):
    """Generates a GPT-generated question and a 150-word justified answer"""
 
    # Step 1: Generate question  #market insight question
    q_prompt = f"""
    Act like a research analyst, Generate one **short, simple, and professional** question (max 20 words)
    about the {market_name.title().capitalize()}.
 
    Example: "How is AI reshaping the European digital signage market?"
     "What Makes AI Integration Crucial for Database Automation Market?"
     "How are Emerging Technologies like AI and IOT Transforming the Folate Market?"
     "How is Artificial Intelligence Transforming Operational Precision in the Deep Water Drilling Market?"

 
    Do not include multiple clauses or long lists . Mention specific technologies (AI, IoT, blockchain, automation, etc.) in the question.
    
    Return only the question, nothing else.
    """
    q_response = client.chat.completions.create(
        model="gpt-5-mini",
        messages=[{"role": "user", "content": q_prompt}],
    )
    question = q_response.choices[0].message.content.strip()
    # Ensure the question starts with a capitalized first letter
    question = capitalize_first_alpha(question)
    
    # Step 2: Generate answer using the actual GPT-generated question #market insight answer
    a_prompt = f"""
    Write a professional answer (~150 words total) to the following question:
    "{question}"
   
    Structure the response into exactly 2 paragraphs:
 
    Paragraph 1 (~100 words): Provide a structured answer with clear organization. Define key aspects, explain current state, provide context about the market situation, add some instances and make it more engaging and contextually relevant.
 
    Paragraph 2 (~50 words): Focus on one clear real world development related to the {market_name} market. . Reference a recent development using a specific company name  month and year, followed by a comma. Emphasize how this innovation {question} supports market growth or efficiency with the help of the technolgy mentioned in the question.
 
    - Do not use bullets, numbering, or "Executive summary".
    - Do not use unnecessary hyphens, ':', ';'.
    - Write in a human-like tone with smooth transitions.
    - Avoid too long sentences.
    - Use qualitative data, not numerical or percentages.
    - use numberic data only for dates as given in the example part.
    - Ensure the answer is relevant to the question and market.
    - Dont add any links or references.
    """
 
    a_response = client.responses.create(
        model="gpt-5-mini",
        tools=[{
        "type": "web_search_preview",
        "search_context_size": "low",
        }
    ],
        input=[
        {
            "role": "system",
                "content": (
                    "Write ONLY plain text. "
                    "Do not use markdown and this '()' symbols anywhere."
                    "look up a verified recent company development online so you can include an accurate company name and month/year in the second paragraph"
                )
        },
        {
            "role": "user",
            "content": a_prompt
        }
    ]
)

    answer = clean_all_paragraphs(clean_sentence_end(remove_urls(a_response.output_text.strip())))
    return question, answer


        #------------------------Regional Insights (3 regions + countries)-----------------------
def build_regional_prompt(market_name):
    # Extract regions and countries from REGIONS_MAPPING dictionary
    dominant_region = REGIONS_MAPPING["dominant"]["region"]
    dominant_countries_list = REGIONS_MAPPING["dominant"]["countries"]
    dominant_countries = ", ".join(dominant_countries_list) if isinstance(dominant_countries_list[0], str) else ", ".join([c["name"] for c in dominant_countries_list])
    
    second_region = REGIONS_MAPPING["second"]["region"]
    second_countries_list = REGIONS_MAPPING["second"]["countries"]
    # Handle both list and dict formats for Europe countries
    if isinstance(second_countries_list[0], dict):
        second_countries = ", ".join([c["name"] for c in second_countries_list])
        second_countries_info = "\n   ".join([f"- {c['name']} ({c['type']})" for c in second_countries_list])
    else:
        second_countries = ", ".join(second_countries_list)
        second_countries_info = "\n   ".join([f"- {c}" for c in second_countries_list])
    
    third_region = REGIONS_MAPPING["third"]["region"]
    third_countries_list = REGIONS_MAPPING["third"]["countries"]
    third_countries = ", ".join(third_countries_list) if isinstance(third_countries_list[0], str) else ", ".join([c["name"] for c in third_countries_list])
    
    return f"""
Write ONLY plain text, no markdown, no titles, no bullet points, no extra blank lines beyond single newlines between sections/paragraphs.

Generate a Regional Insights section for the global {market_name} using the following fixed regional hierarchy:

Dominant Region: {dominant_region} ({dominant_countries})
Second Region: {second_region} ({second_countries})
   Countries in {second_region}:
   {second_countries_info}
Third Region: {third_region} ({third_countries})

Use professional, qualitative market-research tone. All content must be qualitative — NO numbers, percentages, years, statistics, forecasts.

STRUCTURE FOR EACH REGION - MUST FOLLOW THIS EXACT FORMAT:

[QUESTION ABOUT REGION]
[ANSWER ABOUT REGION - 120-150 words]
[COUNTRY NAME] {market_name}
[PARAGRAPH ABOUT COUNTRY - 70 words, starting with {market_name}]

SPECIFIC INSTRUCTIONS:

Region 1: Why does {dominant_region} Dominate the Global {market_name}?
   Answer paragraph (120–150 words) on dominance drivers and strengths
   Countries: {dominant_countries}
   For each country in {dominant_countries}: [CountryName] {market_name} (subheading) + paragraph

Region 2: What is Driving the Rapid Expansion of {market_name} in {second_region}?
   Answer paragraph (120–150 words) on growth drivers and position
   Key Countries:
   {second_countries_info}
   For each country listed: [CountryName] {market_name} (subheading) + paragraph

Region 3: How is {third_region} Strengthening its Position in {market_name}?
   Answer paragraph (120–150 words) on how region is advancing its role
   Countries: {third_countries}
   For each country in {third_countries}: [CountryName] {market_name} (subheading) + paragraph

Output only the generated regional insights text — nothing else.

Mandatory rules:
- Do NOT use numbering, bullet points, or markdown symbols (#, *, -, •)
- Only lines ending with "(subheading)" are treated as subheadings
- NO numbers, percentages, years, statistics, or CAGR
- Do NOT use possessive "'s" after country names (e.g., use "Canada renewable diesel" NOT "Canada's renewable diesel")
- Professional, qualitative, market-research tone
- Clean plain text suitable for Word document insertion
- CRITICAL: Each region must have format: QUESTION then ANSWER then COUNTRY sections with [CountryName]{market_name} format
"""
def get_regional_insights_text(market_name, market_inputs=None, segments_data=None):
    # Get dynamic REGIONS_MAPPING based on market analysis
    global REGIONS_MAPPING
    REGIONS_MAPPING = get_regions_mapping_from_gpt(market_name, market_inputs, segments_data)
    
    # Build prompt with dynamically determined regions
    prompt = build_regional_prompt(market_name)
    # Call GPT
    response = client.chat.completions.create(
        model="gpt-5-mini",
        messages=[{"role": "user", "content": prompt}],
    )
    text = response.choices[0].message.content.strip()
    
    # Remove possessive 's from regional names
    regions = ["North America", "Europe", "Asia Pacific", "Latin America", "Middle East & Africa"]
    for region in regions:
        text = text.replace(f"{region}'s", region)
    
    return text
def add_regional_insights(document, regional_text):
    lines = [l.strip() for l in regional_text.split("\n") if l.strip()]
    # ---- Heading: Regional Insights ----
    heading = document.add_paragraph()
    heading.style = document.styles["Heading 1"]
    run = heading.add_run("Regional Insights")
    run.font.name = "Poppins"
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)

    expect_answer = False
    for line in lines:
        # If previous line was a question, this line is the answer paragraph
        if expect_answer:
            p = document.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Poppins"
            run.font.size = Pt(12)
            # ensure answer is NOT bold
            run.bold = False
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            expect_answer = False
            continue

        # Case 1: Question (ends with ?)
        if line.endswith("?"):
            p = document.add_paragraph()
            run = p.add_run(line.strip() + " |@12")
            run.font.name = "Poppins"
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.bold = False
            run.italic = False
            expect_answer = True
            continue

        # Case 2: Country subheading marked explicitly with "(subheading)"
        m = re.search(r"\(subheading\)\s*$", line, re.IGNORECASE)
        if m:
            # remove the marker and any trailing punctuation/spaces
            clean = re.sub(r"\(subheading\)\s*$", "", line, flags=re.IGNORECASE).strip()
            p = document.add_paragraph()
            run = p.add_run(clean + " |@12")
            run.font.name = "Poppins"
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.bold = False
            run.italic = False
            continue

        # Default: plain paragraph (qualitative answer or other text)
        p = document.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Poppins"
        run.font.size = Pt(12)
        run.bold = False
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
def build_dynamics_prompt(market_name):
    return f"""
Write ONLY plain text, no markdown, no titles, no asterisks, no __ symbols, no <b> tags, no ** symbols.
Generate a Market Dynamics section for the Global {market_name}. Market Dynamics heading not required as i have already added it manually.
Structure:

Drivers
Subheading 1
Write the subheading on a new line. Maximum 6 words. Use title case for all words EXCEPT the specific sentence verbs listed below. Do NOT use __, <b>, **, or any formatting characters.
The first word of the subheading must always be capitalized, even if it appears in the sentence verbs list.
Allowed lowercase sentence verbs only:
is, and, are, was, were, be, on, been, being, has, have, had, driven, growing, expected, leveraging, adopting, using

Bullet explanation
On the next line, add one bullet point.
Indent the bullet with exactly one tab so it appears offset from the margin.
The explanation text must start immediately after the bullet, not indented further.
The explanation must be 80 to 90 words and justified.
Explain clearly why or how this factor contributes to market growth using qualitative discussion only.
Do not use analytical or numerical data.
Maintain a factual and professional tone.

Subheading 2
Follow the same rules as Subheading 1.
Maximum 6 words.
Use title case for all words EXCEPT the allowed lowercase sentence verbs listed above.
Do NOT use formatting characters.

Bullet explanation
On the next line, add one bullet point, indented by one tab.
The explanation must be 80 to 90 words and justified.
Follow the same formatting, content, and tone rules as above.

Restraints
Use the same structure, capitalization rules, bullet formatting, explanation length, and tone as defined for Drivers.
Restraints
Same structure.
Subheading 1
Write the subheading on a new line. Maximum 6 words. Use title case for all words EXCEPT the specific sentence verbs listed below. Do NOT use __, <b>, **, or any formatting characters. The first word of the subheading must always be capitalized, even if it appears in the sentence verbs list.

Bullet explanation
On the next line, add one bullet point. Indent the bullet with exactly one tab so it appears offset from the margin. The explanation text must start immediately after the bullet, not indented further. The explanation must be 80 to 90 words and justified. Explain clearly why or how this factor restricts or slows market growth using qualitative discussion only. Do not use analytical or numerical data. Maintain a factual and professional tone.

Subheading 2
Follow the same rules as Subheading 1. Maximum 6 words. Use title case for all words EXCEPT the allowed lowercase sentence verbs listed above. Do NOT use formatting characters.

Bullet explanation
On the next line, add one bullet point, indented by one tab. The explanation must be 80 to 90 words and justified. Follow the same formatting, content, and tone rules as above. Ensure this is a SECOND DISTINCT restraint factor affecting the market.
Rules:

- Output must be plain text only.
- NO markdown, NO asterisks, NO numbered lists, NO underscores, NO <b> tags, NO ** symbols.
- Drivers must explain why or how a factor supports growth, not merely describe a trend.
- Use qualitative language only.
- Section headers must appear as heading-style text on their own line.
- Do not place subheadings in bullets.
- Do not use unnecessary hyphens.
- Bullet lines must be indented with one tab, with text starting immediately after the bullet.
- Maintain consistent capitalization: title case for all words except the specified sentence verbs only.
- Tone must remain factual, neutral, and professional.
"""
def get_market_dynamics(market_name):
    prompt = build_dynamics_prompt(market_name)
    response = client.chat.completions.create(
        model="gpt-5-mini",
        messages=[{"role": "user", "content": prompt}],
    )
    return response.choices[0].message.content.strip()
def add_market_dynamics(doc, dynamics_text):
    # Market Dynamics Heading (H2)
    run = doc.add_heading("Market Dynamics", level=1).runs[0]
    run.font.size = Pt(16)
    run.font.name = "poppins"
    run.font.color.rgb = RGBColor(0, 0, 0)
    section_header = False
    for para in dynamics_text.split("\n"):
        line = para.strip()
        if not line:
            continue
        # Section headers
        if line.lower().startswith("drivers"):
            p = doc.add_paragraph()
            run = p.add_run("Drivers |@12")
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.name = "poppins"
            run.bold = False
            section_header = True
            continue
        elif line.lower().startswith("restraints"):
            p = doc.add_paragraph()
            run = p.add_run("Restraints |@12")
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.name = "poppins"
            run.bold = False
            section_header = True
            continue
        # Subheading
        if section_header or (line and not line.startswith("•") and not line.startswith("-")):
            # Remove formatting characters: __, <b>, **, </b>
            clean_line = line.replace("__", "").replace("<b>", "").replace("</b>", "").replace("**", "")
            p = doc.add_paragraph()
            run = p.add_run(clean_line)
            run.font.name = "Poppins"
            run.font.size = Pt(12)
            run.bold = True
            section_header = False
            continue
        # Bullet explanation
        if line.startswith("•") or line.startswith("-"):
            explanation = line.lstrip("•- ").strip()
            p = doc.add_paragraph(explanation, style='List Paragraph')
            pPr = p._element.get_or_add_pPr()
            numPr = pPr.get_or_add_numPr()
            numId = OxmlElement('w:numId')
            numId.set(qn('w:val'), '1')
            ilvl = OxmlElement('w:ilvl')
            ilvl.set(qn('w:val'), '0')
            numPr.append(ilvl)
            numPr.append(numId)
            p.paragraph_format.left_indent = Pt(36)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = "Poppins"
                run.font.size = Pt(12)
            continue
        p = doc.add_paragraph(line, style='List Paragraph')
        pPr = p._element.get_or_add_pPr()
        numPr = pPr.get_or_add_numPr()
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), '1')
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), '0')
        numPr.append(ilvl)
        numPr.append(numId)
        p.paragraph_format.left_indent = Pt(36)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = "Poppins"
            run.font.size = Pt(12)
def build_competitive_prompt(market_name):
    return f"""
Write ONLY plain text, no markdown, no titles.
Generate a Competitive Landscape section for the Global {market_name} and it is important to add something please dont question anything there.
Structure:
1. Overview paragraph:
   - Provide 50–70 words.
   - Must highlight Compitive landscape, market driver more specific related to compitions.
   - Must highlight real strategies (like M&A, partnerships, tech innovation) with specific examples.
   - Avoid generic phrases like "players focus on growth" – be concrete.
2. Startups:
   - Provide exactly two bullet points.
   - Startups should not be established before 2019.
   - Don't write long sentences
   - Each bullet point should be 70–100 words.
   - Format: "- <Company Name>: Established in <year>, their main objective is <objective>. Recent development: <specific detail>."
   - Mention any funding, product launch, or expansion if relevant.
Rules:
- Tone: professional, factual, business-focused.
- Use latest Data
- dont use unnecessary hyphens.
- do not use analytical and numerical data. it should have qualitative data.
- Do not add section titles, headers, or bold text.
- Output should start directly with the overview paragraph.
- only verified real startups.
- Dont add any links or references.
- It should take recent market lanscape which is recently updated.
- must have competitive landscape related information.
"""
def get_competitive_landscape(market_name):
    prompt = build_competitive_prompt(market_name)

    response = client.responses.create(
        model="gpt-5-mini",
        tools=[{
            "type": "web_search_preview",
            "search_context_size": "high",
        }],
        input=[
            {
                "role": "system",
                "content": (
                    "Write ONLY plain text. "
                    "Do not use markdown and this '()' symbols anywhere. "
                )
            },
            {
                "role": "user",
                "content": prompt
            }
        ]
    )

    return clean_all_paragraphs(clean_sentence_end(remove_urls(response.output_text.strip())))

def set_poppins_style(paragraph, size=12):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in paragraph.runs:
        run.font.name = "Poppins"
        run.font.size = Pt(size)
    if not paragraph.runs:
        run = paragraph.add_run()
        run.font.name = "Poppins"
        run.font.size = Pt(size)
    return paragraph.runs[0]
def add_competitive_landscape(doc, landscape_text):
    # Add heading
    run = doc.add_heading("Competitive Landscape", level=1).runs[0]
    run.font.size = Pt(16)
    run.font.name = "Poppins"
    run.font.color.rgb = RGBColor(0, 0, 0)
    lines = [line.strip() for line in landscape_text.split("\n") if line.strip()]
    overview_lines = []
    bullet_lines = []
    for line in lines:
        if line.startswith("-"):
            bullet_lines.append(line)
        else:
            overview_lines.append(line)
    if overview_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run()
        set_poppins_style(run, size=12)
        run.text = " ".join(overview_lines)
    # Add bullet points
    for bullet in bullet_lines:
        bullet_text = bullet.lstrip("-").strip()
        p = doc.add_paragraph(bullet_text, style='List Paragraph')
        pPr = p._element.get_or_add_pPr()
        numPr = pPr.get_or_add_numPr()
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), '1')
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), '0')
        numPr.append(ilvl)
        numPr.append(numId)
        p.paragraph_format.left_indent = Pt(36)
        p.paragraph_format.first_line_indent = Pt(-18)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = "Poppins"
            run.font.size = Pt(12)
def generate_segmental_insights(market_name, segments_data):
    # Build prompt dynamically from parsed JSON segments
    prompt = f"""You are SkyQuest’s ABIRAW team. Write a segmental insights section for the {market_name.lower()}.
 
Segments and their sub-segments:
"""
    for seg, subs in segments_data.items():
        # Filter out geographic segments
        geographic_keywords = ['region', 'country', 'geography', 'by region', 'by country', 'geographic']
        if not any(keyword in seg.lower() for keyword in geographic_keywords):
            prompt += f"\n- {seg}: {', '.join(subs)}"
    prompt += f"""
Instructions:
Generate exactly TWO questions strictly based on sub-segments ONLY (NO geographic references).

FOR EACH QUESTION:
Structure the answer into exactly 2 paragraphs:

PARAGRAPH 1 - DOMINATING SUB-SEGMENT (~70 words):
- MUST START WITH: "[Sub-Segment Name] segment [dominates/leads/stands out] because"
- Example openings: "Power infrastructure segment dominates because", "Cloud service providers lead because", "Enterprise solutions segment stand out because"
- After the opening, explain why the dominant sub-segment leads the {market_name} market
- Explain the factors driving its dominance
- Use qualitative reasoning with cause-and-effect explanation
- Do not use numbers or percentages

PARAGRAPH 2 - FASTEST GROWING SUB-SEGMENT (~50 words):
- MUST START WITH a transition connector such as:
  "However,", "On the other hand,", "Meanwhile,", "Conversely,", or another appropriate contrast/transition phrase relevant to the market context
- After the connector, immediately introduce a DIFFERENT sub-segment and describe it as the most rapidly expanding / witnessing the strongest growth momentum / emerging as the key high-growth area in the {market_name}
- Identify a DIFFERENT sub-segment from the same segment other than the dominating one in Paragraph 1
- Explain why this sub-segment is growing the fastest in the {market_name}
- Focus on growth drivers such as rising adoption, innovation, expanding applications, regulatory momentum, or emerging demand
- Explain how this sub-segment is accelerating future market expansion and opportunity creation
- Do NOT mention or refer to the dominating sub-segment from Paragraph 1
- Do not use numbers, statistics, comparisons, or percentages

QUESTION FORMAT:
- Use lowercase sentence case for questions. Start with a lowercase letter after the question word (e.g., "What role do...", "How is..." in What "W" is capital and in How "H" is capital)
- Questions MUST be specific to the {market_name} market and its unique characteristics
- Questions should address real market dynamics specific to the sub-segments in {market_name}
- Questions should sound like professional research analyst questions that would only apply to this specific market
- Don't use generic templates - make each question directly relevant to {market_name}
- Don't use the word "segment" and "sub-segment" in the question
- Ask about sub-segment characteristics, roles, or market impact that are specific to {market_name}
- Example patterns: "What role do [Sub-Segment] play in transforming [market_name]?", "How is [Sub-Segment] addressing key challenges in [market_name]?", "Which [Sub-Segment] offers the most competitive advantage in [market_name]?"
- and make the question like a research analyst, specific to {market_name} only.
- and make the question concise (max 15 words), avoiding multiple clauses or long lists.
Rules:
- Generate 2 complete questions with 2-paragraph answers each
- Each question MUST be unique and tailored to {market_name} specifically
- NO generic or templated questions
- NO geographic/regional references in any question or answer
- NO markdown symbols (#, *, -,[],())
- Professional, analytical market-research tone
- Clear, coherent language
- Avoid unnecessary hyphens
- Each paragraph must strictly adhere to word count (~70 words for Para 1, ~50 words for Para 2 and try to avoid this "-" in the answer)
"""
 
    response = client.chat.completions.create(
        model="gpt-5-mini",
        messages=[{"role": "user", "content": prompt}],
    )
    raw_text = response.choices[0].message.content.strip()
    # Ensure questions in segmental insights start with capital letters
    return capitalize_questions_in_text(raw_text)

def load_company_profiles_from_data(data):
    return data.get("company_profiles") or data.get("companies") or []

def add_company_profiles(doc, market_name, data):
    # Company Profiles heading in Poppins, 12pt, bold (as paragraph, no heading)
    p_heading = doc.add_paragraph()
    run_heading = p_heading.add_run("Top Player’s Company Profile")
    run_heading.font.name = "Poppins"
    run_heading.font.size = Pt(12)
    run_heading.bold = True
    p_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Load company profiles from data
    company_profiles = load_company_profiles_from_data(data)

    # Normalise: accept string (newline-separated) or list
    if isinstance(company_profiles, str):
        company_profiles = [c.strip() for c in company_profiles.splitlines() if c.strip()]

    # Loop through the company profiles and add them to the document
    for company in company_profiles:
        p = doc.add_paragraph(style='Normal')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pPr = p._element.get_or_add_pPr()
        numPr = pPr.get_or_add_numPr()
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), '0')
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), '1')
        numPr.append(ilvl)
        numPr.append(numId)
        run = p.add_run(company.strip())
        run.font.name = "Poppins"
        run.font.size = Pt(12)
        run.bold = False
def build_recent_developments_prompt(market_name):
    return f"""
Write ONLY plain text, no markdown, no titles.
Generate a Recent Developments section for the Global {market_name} focusing on the years 2025 to present.
Instructions:
- Provide exactly three recent developments from 2025 to present.
- The Recent Developments section should focus on key players, rather than covering the market holistically.(for example: Schneider Electric introduced the Modicon M660, in February 2025, a state-of-the-art Industrial PC (IPC) Motion controller that combines precision motion control, safety features, and edge computing capabilities in a single device. The combined solution allows for simple integration of control, data processing, and communication within applications with real-time control and low latency, enhancing the efficiency and reliability of MEP systems.)
- Each development should be in a bullet point.
- do not use analytical and numerical data. it should have qualitative data.
- Each bullet point should be 50-60 words.
- include year and month.(for ex dont use : Earlier this year Visa rolled out an enterprise. instead use : Visa rolled out an enterprise in February 2025. )
- Present developments in reverse chronological order (most recent first).
- Bullets should start on the next line with one tab indent.
- Text after the bullet should start immediately after the bullet (no extra space).
- Use professional and factual tone.
- dont use unnecessary hyphens.
- Output only bullet points, no headers or extra titles.
- Dont add any links or references.
"""
def get_recent_developments(market_name):
    prompt = build_recent_developments_prompt(market_name)

    response = client.responses.create(
        model="gpt-5-mini",
        tools=[{
            "type": "web_search_preview",
            "search_context_size": "medium",
        }],
        input=[
            {
                "role": "system",
                "content": (
                    "Write ONLY plain text. "
                    "Do not use markdown and this '()' symbols anywhere."
                )
            },
            {
                "role": "user",
                "content": prompt
            }
        ]
    )

    return clean_all_paragraphs(clean_sentence_end(remove_urls(response.output_text.strip())))

def add_recent_developments(doc, recent_text):
    p_heading = doc.add_paragraph()
    run_heading = p_heading.add_run("Recent Developments")
    run_heading.font.name = "Poppins"
    run_heading.font.size = Pt(16)
    run_heading.bold = True
    p_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lines = [line.strip() for line in recent_text.split("\n") if line.strip()]
    for line in lines:
        clean_text = line.lstrip("-•. ").strip()
        p = doc.add_paragraph(clean_text, style='List Paragraph')
        pPr = p._element.get_or_add_pPr()
        numPr = pPr.get_or_add_numPr()
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), '1')
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), '0')
        numPr.append(ilvl)
        numPr.append(numId)
        p.paragraph_format.left_indent = Pt(36)
        p.paragraph_format.first_line_indent = Pt(-18)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = "Poppins"
            run.font.size = Pt(12)
            run.bold = False
def extract_restraints_headings(dynamics_text):
    """Extract Restraints subheadings from Market Dynamics text"""
    restraints_headings = []
    in_restraints = False
    
    for line in dynamics_text.split("\n"):
        line = line.strip()
        
        # Check if we're entering Restraints section
        if line.lower() == "restraints":
            in_restraints = True
            continue
        
        # Stop collecting if we hit another section (in case there's one after)
        if in_restraints and line and not line.startswith("•") and not line.startswith("-"):
            # This is a restraint heading (not a bullet)
            if line and line[0].isupper():
                restraints_headings.append(line.lower())
    
    return restraints_headings

def build_key_trends_prompt(market_name, excluded_headings=None):
    if excluded_headings is None:
        excluded_headings = []
    
    excluded_text = ""
    if excluded_headings:
        excluded_text = f"\n\nIMPORTANT: DO NOT use these headings (from Restraints section) for your trends:\n- {', '.join(excluded_headings)}\nInstead, use completely different, fresh topic headings that represent emerging opportunities."
    
    return f"""
Write ONLY plain text, no markdown, no titles.
Generate a Key Market Trends section for the Global {market_name}.
Instructions:
- Provide exactly 2 key market trends. Each trend must begin with a short, bold subheading (maximum 6 words, for ex- Electrification Momentum Transforming Automotive Value Chains), followed by a colon and an 80-word explanation in plain text.
- Each trend should be about 80 words.
- Each trend must have a clear subheading (max 5 words).
- Do NOT include any bullet characters in the output.
- Format each trend starting with the subheading (bold), followed by the explanation, on its own line.
- Use professional, business tone.
- do not use analytical and numerical data. it should have qualitative data.
- dont use any markdown symbols (*, -, •,+). - The code will handle bullets automatically.
- Do not use any symbols of html like <b>,<i>, etc.
- Output starts directly with the subheading, no extra headings or titles.{excluded_text}
"""

def get_key_market_trends(market_name, dynamics_text=None):
    excluded_headings = []
    
    # Extract Restraints headings to avoid duplication
    if dynamics_text:
        excluded_headings = extract_restraints_headings(dynamics_text)
    
    prompt = build_key_trends_prompt(market_name, excluded_headings)
    response = client.chat.completions.create(
        model="gpt-5-mini",
        messages=[{"role": "user", "content": prompt}],
    )
    return response.choices[0].message.content.strip()
def set_poppins_style(run, size=12, bold=False):
    run.font.name = "Poppins"
    run.font.size = Pt(size)
    run.bold = bold
def add_key_market_trends(doc, trends_text):
    heading_run = doc.add_heading("Key Market Trends", level=1).runs[0]
    heading_run.font.name = "Poppins"
    heading_run.font.size = Pt(16)
    heading_run.font.color.rgb = RGBColor(0, 0, 0)
    lines = [line.strip() for line in trends_text.split("\n") if line.strip()]
    for line in lines:
        if ":" in line:
            subhead, explanation = line.split(":", 1)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Pt(36)
            p.paragraph_format.first_line_indent = Pt(-18)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run1 = p.add_run(subhead.strip().title() + ": ")
            set_poppins_style(run1, size=12, bold=True)
            run2 = p.add_run(explanation.strip())
            set_poppins_style(run2, size=12, bold=False)
        else:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if p.runs:
                run = p.runs[0]
                set_poppins_style(run, size=12, bold=False)
def extract_doc_text(doc):
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
def build_skyquest_analysis_prompt_from_doc(market_name, doc):
    doc_text = extract_doc_text(doc)
    return f"""
Summarize the following report content into a 'SkyQuest Analysis' section for the Global {market_name}.
Content to summarize:
{doc_text}
Requirements:
- Length: 100–120 words.
- Summarize this in a professional, user-friendly tone, just like SkyQuest’s ABIRAW team would, keeping it clear and approachable without analytical jargon.
- Include explicitly:
  • One key driver
  • One restraint
  • Dominating region
  • Dominating segment
  • Second driver
- Cohesive paragraph, no bullet points, no markdown.
- Do not use unnecessary colons in the text and dont give any extra space. Avoid formats like "SkyQuest Analysis:" or "Key Findings:". Instead, rewrite sentences so they flow naturally, for example: "As per SkyQuest analysis, ..."

"""
def get_skyquest_analysis_from_doc(market_name, doc):
    prompt = build_skyquest_analysis_prompt_from_doc(market_name, doc)
    response = client.chat.completions.create(
        model="gpt-5-mini",
        messages=[{"role": "user", "content": prompt}],
    )
    return response.choices[0].message.content.strip() if response.choices else "No SkyQuest Analysis generated."
def add_skyquest_analysis(doc, market_name):
    heading = doc.add_heading("SkyQuest Analysis", level=1)
    run = heading.runs[0]
    run.font.size = Pt(16)
    run.font.name = "Poppins"
    run.font.color.rgb = RGBColor(0, 0, 0)
    analysis = get_skyquest_analysis_from_doc(market_name, doc)
    # Prepend ABIRAW description
    abiraw_line = ("SkyQuest’s ABIRAW (Advanced Business Intelligence, Research & Analysis Wing) "
                   "is our Business Information Services team that Collects, Collates, Correlates, "
                   "and Analyses the Data collected by means of Primary Exploratory Research backed "
                   "by robust Secondary Desk research. \n\n")
    full_analysis = abiraw_line + analysis
    p = doc.add_paragraph()
    run = p.add_run(full_analysis)
    run.font.size = Pt(12)
    run.font.name = "Poppins"
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
def set_poppins_style(run, size=12, bold=False, color=RGBColor(0, 0, 0)):
 
    run.font.name = "Poppins"
 
    try:
        rPr = run._element.rPr
        if rPr is None:
            from docx.oxml import OxmlElement
            rPr = OxmlElement('w:rPr')
            run._element.insert(0, rPr)
        rFonts = rPr.rFonts
        if rFonts is None:
            from docx.oxml import OxmlElement
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), 'Poppins')
    except Exception:
        pass
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    return run
def parse_segments_hierarchy(toc_list):
    hierarchy = []
    stack = [(0, hierarchy)]
    for line in toc_list:
        s = line.strip()
        if not s or s.lower().startswith(("revenue", "market")):
            continue
        match = re.match(r"^([\d\.]+)\s+(.*)$", s)
        if match:
            num, name = match.groups()
            num = num.strip('.')
            level = len(num.split('.'))
        else:
            name = s
            level = 1
        node = {"name": name.strip(), "children": []}
        while stack and stack[-1][0] >= level:
            stack.pop()
        stack[-1][1].append(node)
        stack.append((level, node["children"]))
    return hierarchy
BULLETS = ["•", "○"]  # can add more like ":black_small_square:", "–", etc.
def add_segments_to_doc(doc, segments, level=0):
    for seg in segments:
        if level == 0:
            # Top-level (no bullet)
            p = doc.add_paragraph()
            run = p.add_run(seg["name"])
            set_poppins_style(run, size=12, bold=True)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.left_indent = Cm(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            # Pick bullet based on level (cycling)
            bullet_char = BULLETS[(level - 1) % len(BULLETS)]
            p = doc.add_paragraph()
            run = p.add_run(f"{bullet_char} {seg['name']}")
            set_poppins_style(run, size=12, bold=False)
            # Indent grows with level
            p.paragraph_format.left_indent = Cm(0.5 * level)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Recurse for children up to second level only (exclude third level)
        if seg.get("children") and level < 1:
            add_segments_to_doc(doc, seg["children"], level + 1)
def transform_market_data(data, market_name):
    segments = {}
    current_segment = None
    for item in data:
        if isinstance(item, tuple):
            item, level = item
        else:
            level = 0 # Assume level 0 if not tuple
        if level == 0:
            if f"Global {market_name} Size by" in item:
                segment_name = item.replace(
                    f"Global {market_name} Size by ", ""
                ).split(" & ")[0]
                current_segment = segment_name
                segments[current_segment] = []
            continue
        if current_segment and level > 0:
            if item != "Market Overview":
                if level == 1:
                    segments[current_segment].append(item)
                elif level == 2:
                    if (
                        segments[current_segment]
                        and isinstance(segments[current_segment][-1], list)
                    ):
                        segments[current_segment][-1].append(item)
                    else:
                        segments[current_segment].append([item])
    formatted_output = []
    for segment, sub_segments in segments.items():
        formatted_subs = []
        for sub in sub_segments:
            if isinstance(sub, list):
                formatted_subs[-1] += f" ({', '.join(sub)})"
            else:
                formatted_subs.append(sub)
        segment_line = f"Segment {segment}: Sub-Segments {', '.join(formatted_subs)}"
        formatted_output.append(segment_line)
    return formatted_output, segments
def generate_segmental_analysis(market_name, data):
    # Load from data instead of file
    toc_list = data.get("table_of_contents", [])
 
    segments_data = {}
    current_segment = None
 
    for entry in toc_list:
        match = re.match(r"^(\d+)\.\s+(.*)$", entry)
        sub_match = re.match(r"^(\d+\.\d+)\.\s+(.*)$", entry)
        if match:
            current_segment = match.group(2).strip()
            segments_data[current_segment] = []
        elif sub_match and current_segment:
            segments_data[current_segment].append(sub_match.group(2).strip())
 
    # Build analysis text
    text = f"Global {market_name.lower()} is segmented by "
    segment_names = []
    segment_details = []
    
    # Keywords to identify region-related segments (to filter from names and details)
    region_keywords = ["region", "geographic", "geography", "geographical"]
 
    for segment, sub_segments in segments_data.items():
        # Convert segment name to lowercase
        segment_lower = segment.lower()
        
        # Skip region-related segments from segment_names and segment_details
        if any(keyword in segment_lower for keyword in region_keywords):
            continue
            
        segment_names.append(segment_lower)
        sub_details = [s for s in sub_segments if isinstance(s, str)]
     
        if not sub_details:
            segment_details.append(
                f"Based on {segment_lower}, no specific sub-segments were identified."
            )
        elif len(sub_details) > 1:
            joined = ", ".join(sub_details[:-1]) + " and " + sub_details[-1]
            segment_details.append(f"Based on {segment_lower}, the market is segmented into {joined}.")
        else:
            segment_details.append(f"Based on {segment_lower}, the market is segmented into {sub_details[0]}.")
 
    # Build segment names list (include region in opening sentence)
    if segment_names:
        text += ", ".join(segment_names) + " and region. "
    else:
        text += "region. "
 
    text += " ".join(segment_details)
    text += " Based on region, the market is segmented into North America, Europe, Asia Pacific, Latin America and Middle East & Africa."
 
    return text, segments_data

def build_meta_description(market_name, value_2024, value_2025, value_2033, currency, cagr):
    return (
        f"{market_name.lower().capitalize()} size is poised to grow from USD {value_2025} {currency} in 2025 to USD {value_2033} {currency} by 2033, "
        f"growing at a CAGR of {cagr}% from (2026-2033)."
    )

def title_h1(segments_data, market_name):
    # Build segment parts with ALL sub-segments
    segment_parts = []  # list of (segment_name, [sub_segments])
    for segment, sub_segments in segments_data.items():
        top_level_subsegments = [s for s in sub_segments if isinstance(s, str)]
        segment_parts.append((segment, list(top_level_subsegments)))

    def build_title(parts):
        pieces = []
        for seg_name, subs in parts:
            if subs:
                pieces.append(f"By {seg_name} ({', '.join(subs)})")
            else:
                pieces.append(f"By {seg_name}")
        pieces.append("By Region")
        text_seg = ", ".join(pieces)
        return f"{market_name} Size, Share, Growth Analysis, {text_seg} - Industry Forecast 2026-2033"

    title = build_title(segment_parts)

    # If >40 words, remove sub-segments from the LAST segment first, then work backwards
    while len(title.split()) > 40 and segment_parts:
        trimmed = False
        for i in range(len(segment_parts) - 1, -1, -1):
            if segment_parts[i][1]:  # has sub-segments
                segment_parts[i] = (segment_parts[i][0], [])  # remove them
                trimmed = True
                break
        if not trimmed:
            break  # nothing left to trim
        title = build_title(segment_parts)

    return title
def add_rd_metadata_section(doc, metadata_pairs):
    for key, value in metadata_pairs:
        # Special handling for Report Type
        if key == "Report Type":
            # Add Report Type heading
            k = doc.add_heading(level=1)
            run_k = k.runs[0] if k.runs else k.add_run()
            set_poppins_style(run_k, size=16, bold=True).text = str(key)
            
            # Hardcoded pricing structure - each item on separate line
            pricing_lines = [
                "Single",
                "PPT",
                "5300",
                "Single",
                "Word",
                "5300",
                "Single",
                "Excel",
                "5300",
                "Single",
                "PowerBI",
                "6200",
                "Multiple",
                "PPT",
                "6200",
                "Multiple",
                "Word",
                "6200",
                "Multiple",
                "Excel",
                "6200",
                "Multiple",
                "PowerBI",
                "7100",
                "Enterprise",
                "PPT",
                "7100",
                "Enterprise",
                "Word",
                "7100",
                "Enterprise",
                "Excel",
                "7100",
                "Enterprise",
                "PowerBI",
                "8000",
            ]
            
            # Add each line as normal paragraph text
            for line in pricing_lines:
                p = doc.add_paragraph()
                run = p.add_run(line)
                set_poppins_style(run, size=12)
        else:
            # Regular handling for all other metadata
            k = doc.add_heading(level=1)
            run_k = k.runs[0] if k.runs else k.add_run()
            set_poppins_style(run_k, size=16, bold=True).text = str(key)
         
            # Text content: size 12, normal, Poppins
            v = doc.add_paragraph()
            run_v = v.add_run(str(value))
            set_poppins_style(run_v, size=12)
def add_analyst_support_section(doc, market_name):
    h = doc.add_heading(level=1)
    run = h.add_run("Analyst Support")
    run.font.name = "Poppins"
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Poppins")
    p = doc.add_paragraph()
    run = p.add_run("Customization Options")
    run.font.name = "Poppins"
    run.font.size = Pt(12)
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Poppins")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p = doc.add_paragraph()
    intro = (
        f"With the given market data, our dedicated team of analysts can offer you the following customization options are available for the {market_name.lower()}:"
    )
    run = p.add_run(intro)
    run.font.name = "Poppins"
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Poppins")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    bullets = [
        ("Product Analysis:", "Product matrix, which offers a detailed comparison of the product portfolio of companies."),
        ("Regional Analysis:", f"Further analysis of the global {market_name.lower()} for additional countries."),
        ("Competitive Analysis:", "Detailed analysis and profiling of additional market players & comparative analysis of competitive products."),
        ("Go to Market Strategy:", "Find the high growth end-users to invest your marketing efforts and increase your customer base."),
        ("Innovation Mapping:", "Identify racial solutions and innovation, connected to deep ecosystems of innovators, start-ups, academics, and strategic partners."),
        ("Category Intelligence:", "Customized intelligence that is relevant to their supply markets which will enable them to make smarter sourcing decisions and improve their category management."),
        ("Public Company Transcript Analysis:", "To improve investment performance by generating new alpha and making better-informed decisions."),
        ("Social Media Listening:", "To Analyse the conversations and trends happening not just around your brand, but around your industry, and using those insights to make better marketing decisions."),
    ]
    for subhead, explanation in bullets:
        p = doc.add_paragraph()
        run1 = p.add_run(subhead)
        run1.font.name = "Poppins"
        run1.font.size = Pt(12)
        run1.bold = True
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), "Poppins")
        run2 = p.add_run(" " + explanation)
        run2.font.name = "Poppins"
        run2.font.size = Pt(12)
        run2._element.rPr.rFonts.set(qn('w:eastAsia'), "Poppins")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(2)
def fetch_ai_contents(market_name, tech="AI", segments_data=None, market_inputs=None):
    """Fetch all AI-generated contents in parallel"""
    with ThreadPoolExecutor(max_workers=10) as executor:
        # Submit independent tasks (all except trends which depends on dynamics)
        future_to_content = {
            executor.submit(get_gpt5_insight, market_name, tech): 'insight',
            executor.submit(get_gpt5_question_and_answer, market_name): 'qa',
            executor.submit(get_regional_insights_text, market_name, market_inputs, segments_data): 'regional',
            executor.submit(get_market_dynamics, market_name): 'dynamics',
            executor.submit(get_competitive_landscape, market_name): 'competitive',
            executor.submit(generate_segmental_insights, market_name, segments_data): 'segmental_insights',
            executor.submit(get_recent_developments, market_name): 'recent',
        }
        
        contents = {}
        for future in as_completed(future_to_content):
            key = future_to_content[future]
            try:
                if key == 'qa':
                    contents['question'], contents['answer'] = future.result()
                elif key == 'insight':
                    contents['insight'] = future.result()
                elif key == 'regional':
                    contents['regional'] = future.result()
                elif key == 'dynamics':
                    contents['dynamics'] = future.result()
                elif key == 'competitive':
                    contents['competitive'] = future.result()
                elif key == 'segmental_insights':
                    contents['segmental_insights'] = future.result()
                elif key == 'recent':
                    contents['recent'] = future.result()
            except Exception as exc:
                print(f'{key} generated an exception: {exc}')
        
        # Now generate trends based on dynamics (to avoid duplicate headings)
        if 'dynamics' in contents:
            try:
                contents['trends'] = get_key_market_trends(market_name, contents['dynamics'])
            except Exception as exc:
                print(f'trends generated an exception: {exc}')
    
    return contents
def export_to_word_RD(data_dict, value_2024, value_2025, value_2033, currency, cagr, companies, rd_meta, output_path=None, tech="AI", doc=None):
    if doc is None:
        doc = Document()
    market_name = data_dict.get("title", "")
    toc_list = data_dict.get("table_of_contents", [])
    
    # Add companies to data_dict for add_company_profiles function
    if companies:
        data_dict["company_profiles"] = companies
    
    # Define market input variables from parameters (passed from generate_docx_from_data)
    value_2024_use = value_2024
    value_2025_calc = value_2025
    value_2033_calc = value_2033
    currency_use = currency
    cagr_use = cagr
    formatted_output, segments = transform_market_data(toc_list, market_name) # Adjusted to use toc_list directly

    # Report Name
    title = doc.add_heading(level=1)
    run = title.runs[0] if title.runs else title.add_run()
    set_poppins_style(run, size=16, bold=True).text = "Report Name"
    p = doc.add_paragraph()
   
    run = p.add_run(f"{market_name}")
    set_poppins_style(run, size=12)

     # Upcoming
    h = doc.add_heading(level=1)
    run = h.runs[0] if h.runs else h.add_run()
    set_poppins_style(run, size=16, bold=True).text = "Upcoming"
    p = doc.add_paragraph()
    run_p = p.add_run("No")
    set_poppins_style(run_p, size=12)
 
    # Segments (new format with Segment and Sub-Segments)
    segments_dict = {}
    current_main = ""
    current_sub = ""
 
    for s in toc_list:
        if not isinstance(s, str):
            continue
         
        match = re.match(r"^([\d\.]+)\s+(.*)$", s)
        if match:
            num, name = match.groups()
            num = num.strip('.')
            num_parts = num.split('.')
            level = len(num_parts)
         
            # Level 1 segments
            if level == 1:
                main_segment = name.strip()
                segments_dict[main_segment] = {}
                current_main = main_segment
             
            # Level 2 segments
            elif level == 2:
                sub_segment = name.strip()
                if current_main not in segments_dict:
                    segments_dict[current_main] = {}
                segments_dict[current_main][sub_segment] = []
                current_sub = sub_segment
             
            # Level 3 segments
            elif level == 3:
                sub_sub_segment = name.strip()
                if current_main in segments_dict and current_sub in segments_dict[current_main]:
                    segments_dict[current_main][current_sub].append(sub_sub_segment)
 
    #output the segments
    h = doc.add_heading(level=1)
    run_h = h.runs[0] if h.runs else h.add_run()
    set_poppins_style(run_h, size=16, bold=True).text = "Segments"
 
    for main_segment, sub_segments in segments_dict.items():
        # Add "Segment" heading
        p_segment_heading = doc.add_paragraph()
        p_segment_heading.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_segment_heading = p_segment_heading.add_run("Segment")
        set_poppins_style(run_segment_heading, size=16)
        p_segment_heading.style = "Heading 2"
        run_segment_heading.font.color.rgb = RGBColor(0, 0, 0)
     
        # Add main segment name
        p_main_segment = doc.add_paragraph()
        p_main_segment.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_main_segment = p_main_segment.add_run(main_segment)
        set_poppins_style(run_main_segment, size=12)
     
        # Add "Sub-Segments" heading
        p_sub_heading = doc.add_paragraph()
        p_sub_heading.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_sub_heading = p_sub_heading.add_run("Sub-Segments")
        set_poppins_style(run_sub_heading, size=16)
        p_sub_heading.style = "Heading 2"
        run_sub_heading.font.color.rgb = RGBColor(0, 0, 0)
     
        # Format and add sub-segments only (exclude sub-sub-segments in brackets)
        formatted_sub_segments = list(sub_segments.keys())
     
        # Add the formatted sub-segments
        p_formatted = doc.add_paragraph()
        p_formatted.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_formatted = p_formatted.add_run(", ".join(formatted_sub_segments))
        set_poppins_style(run_formatted, size=12)

        # Add "Sub-Segments-For-{name}" heading + children for each sub-segment with level-3 items
        for sub_segment, sub_sub_segments in sub_segments.items():
            if sub_sub_segments:
                p_sub_sub_heading = doc.add_paragraph()
                p_sub_sub_heading.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run_sub_sub_heading = p_sub_sub_heading.add_run(f"Sub-Segments-For-{sub_segment}")
                set_poppins_style(run_sub_sub_heading, size=16)
                p_sub_sub_heading.style = "Heading 2"
                run_sub_sub_heading.font.color.rgb = RGBColor(0, 0, 0)

                p_sub_sub = doc.add_paragraph()
                p_sub_sub.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run_sub_sub = p_sub_sub.add_run(", ".join(sub_sub_segments))
                set_poppins_style(run_sub_sub, size=12)
 
    # RD metadata block
    add_rd_metadata_section(doc, rd_meta)
    # --- Methodologies Section ---
    # Heading
    heading = doc.add_paragraph()
    heading.style = doc.styles["Heading 1"]
    run = heading.add_run("Methodologies")
    run.font.name = "Poppins"
    run.font.size = Pt(16)
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Poppins")
    run.font.color.rgb = RGBColor(0, 0, 0)
    # Intro paragraph
    intro= doc.add_paragraph()
    run = intro.add_run(
        f"For the {market_name.lower()}, our research methodology involves a "
        "mixture of primary and secondary data sources. Key steps involved in the "
        "research process are listed below:"
    )
    run.font.name = "Poppins"
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Poppins")
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Information Procurement
    p= doc.add_paragraph()
    run = p.add_run("Information Procurement: ")
    run.font.name = "Poppins"
    run.font.size = Pt(12)
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Poppins")
    run= p.add_run(
        "This stage involved the procurement of market data or related information via primary "
        "and secondary sources. The various secondary sources used included various company "
        "websites, annual reports, trade databases, and paid databases such as Hoover’s, "
        "Bloomberg Business, Factiva, and Avention. Our team did 45 primary interactions Globally "
        "which included several stakeholders such as manufacturers, customers, key opinion leaders, etc. "
        "Overall, information procurement was one of the most extensive stages in our research process."
    )
    run.font.name = "Poppins"
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Poppins")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Information Analysis
    p= doc.add_paragraph()
    run = p.add_run("Information Analysis: ")
    run.font.name = "Poppins"
    run.font.size = Pt(12)
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Poppins")
    run= p.add_run(
        f"This step involved triangulation of data through bottom-up and top-down approaches "
        f"to estimate and validate the total size and future estimate of the {market_name.lower()}."
    )
    run.font.name = "Poppins"
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Poppins")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Report Formulation
    p = doc.add_paragraph()
    run = p.add_run("Report Formulation: ")
    run.font.name = "Poppins"
    run.font.size = Pt(12)
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Poppins")
    run= p.add_run(
        "The final step entailed placing data points in appropriate market spaces to draw viable conclusions."
    )
    run.font.name = "Poppins"
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Poppins")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Validation & Publishing
    p= doc.add_paragraph()
    run = p.add_run("Validation & Publishing: ")
    run.font.name = "Poppins"
    run.font.size = Pt(12)
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Poppins")
    run= p.add_run(
        "Validation is the most important step in the process. Validation & re-validation via an intricately "
        "designed process helped us finalize data points to be used for final calculations. The final market "
        "estimates and forecasts were then aligned and sent to our panel of industry experts for validation "
        "of data. Once the validation was done the report was sent to our Quality Assurance team to ensure "
        "adherence to style guides, consistency & design."
    )
    run.font.name = "Poppins"
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Poppins")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Add Analyst Support section after methodology
    add_analyst_support_section(doc, market_name)
    # Market Insights
    h = doc.add_heading(level=1)
    run = h.runs[0] if h.runs else h.add_run()
    set_poppins_style(run, size=16, bold=True).text = "Market Insights"
    p = doc.add_paragraph()
    run_p = p.add_run(
        f"Global {market_name.title()} size was valued at USD {value_2024_use} {currency_use} in 2024 "
        f"and is poised to grow from USD {value_2025_calc} {currency_use} in 2025 to USD {value_2033_calc} {currency_use} by 2033, "
        f"growing at a CAGR of {cagr_use}% during the forecast period (2026-2033)."
    )
    set_poppins_style(run_p, size=12)
    # Fetch all AI contents in parallel
    seg_text, segments_data = generate_segmental_analysis(market_name, data_dict) # Use data_dict
    market_inputs = data_dict.get('market_inputs', {})
    ai_contents = fetch_ai_contents(market_name, tech, segments_data, market_inputs)
    # Add GPT-5-mini insight
    insight_text = ai_contents['insight']
    # Remove em dashes (—) and replace with spaces
    insight_text = insight_text.replace("—", " ").replace("–", " ")
    p = doc.add_paragraph(insight_text)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY_LOW
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = "Poppins"
    
    # Add question in bold, font size 12
    question = ai_contents['question']
    answer = ai_contents['answer']
    q_para = doc.add_paragraph()
    run = q_para.add_run(question)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Poppins"
    # justified
    a_para = doc.add_paragraph(answer)
    a_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY_LOW
    for run in a_para.runs:
        run.font.size = Pt(12)
        run.font.name = "Poppins"
    # Segmental Analysis
    h = doc.add_heading(level=1)
    run_h = h.runs[0] if h.runs else h.add_run()
    set_poppins_style(run_h, size=16, bold=True).text = "Segmental Analysis"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_p = p.add_run(seg_text)
    set_poppins_style(run_p, size=12)
    
    insights_text = ai_contents['segmental_insights']
    for line in insights_text.split("\n"):
        line = line.strip()
        if not line:
            continue
 
        if line.endswith("?"):
            q_para = doc.add_paragraph(line + " |@12")
            q_para.alignment = WD_ALIGN_PARAGRAPH.LEFT # prevent spacing issues
            run_q = q_para.runs[0]
            set_poppins_style(run_q, size=12, bold=False)
        else:
            ans_para = doc.add_paragraph()
            ans_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run_ans = ans_para.add_run(line)
            set_poppins_style(run_ans, size=12)
    add_regional_insights(doc, ai_contents['regional'])
    add_market_dynamics(doc, ai_contents['dynamics'])
    add_competitive_landscape(doc, ai_contents['competitive'])
    add_company_profiles(doc, market_name, data_dict) # Use data_dict
    add_recent_developments(doc, ai_contents['recent'])
    add_key_market_trends(doc, ai_contents['trends'])
    add_skyquest_analysis(doc, market_name)

#whats included
    h = doc.add_heading(level=1)
    run_h = h.runs[0] if h.runs else h.add_run()
    set_poppins_style(run_h, size=16, bold=True).text = "What’s Included"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # NA
    inc = doc.add_paragraph()
    run = inc.add_run("NA")
    run.font.name = "Poppins"
    run.font.size = Pt(12)
    run.bold = False
    run.italic = False



 
    # H1 Title
 
    h = doc.add_heading(level=1)
    run_h = h.runs[0] if h.runs else h.add_run()
    set_poppins_style(run_h, size=16, bold=True).text = "H1 Title"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Build author/segment info safely
    # H1 Title requirement: show ONLY sub-segments; exclude third-level (sub-sub) items and REGION segments
    region_keywords = ["region", "geographic", "geography", "geographical", "by region", "by country", "country"]
    
    segments_dict_filtered = {}
    for main_segment, sub_segments in segments_dict.items():
        # Skip region-related segments
        if not any(keyword in main_segment.lower() for keyword in region_keywords):
            segments_dict_filtered[main_segment] = sub_segments
    
    # Build segment parts with ALL sub-segments
    segment_parts = []  # list of (segment_name, [sub_segment_names])
    for main_segment, sub_segments in segments_dict_filtered.items():
        sub_names = list(sub_segments.keys()) if sub_segments else []
        segment_parts.append((main_segment, sub_names))

    def build_h1(parts):
        pieces = []
        for seg_name, subs in parts:
            if subs:
                pieces.append(f"By {seg_name} ({', '.join(subs)})")
            else:
                pieces.append(f"By {seg_name}")
        pieces.append("By Region")
        segs_str = ", ".join(pieces)
        return f"{data_dict['title']} {segs_str} - Industry Forecast 2026-2033"

    title_text = build_h1(segment_parts)

    # If >40 words, remove sub-segments from the LAST segment first, then work backwards
    while len(title_text.split()) > 40 and segment_parts:
        trimmed = False
        for i in range(len(segment_parts) - 1, -1, -1):
            if segment_parts[i][1]:  # still has sub-segments
                segment_parts[i] = (segment_parts[i][0], [])  # remove them
                trimmed = True
                break
        if not trimmed:
            break
        title_text = build_h1(segment_parts)
    
    run_p = p.add_run(title_text)
    color = RGBColor(0, 0, 0)
    set_poppins_style(run_p, size=12, color=color)
    
    if output_path:
        doc.save(output_path)
    return doc
# ------------------ Constants ------------------
#value_2024 = 100.0
#cagr = 12.0
# currency = "million"
# Formatting: Headings will be Heading 1, bold, font size 16, Poppins
# Text content will be size 12, normal, Poppins
rd_metadata_pairs = [
    ("Image", "Null"),
    ("Market", "Global {}"),
    ("Role", "Global {}"),
    ("Country", "Global"),
    ("Report Data", "Null"),
    ("Product ID", "Null"),
    ("Download", "13"),
    ("Image Alt", "{}"),
    ("Report Slug", "{}"),
    ("Meta Title", "RANDOM_META_TITLE"),
    ("Meta Description", "Null"),
    ("Pages", "157"),
    ("Report Type", "Single"),
]

def get_random_meta_title(market_name):
    """Randomly select a meta title template and format it with market name"""
    meta_title_templates = [
        "{} Size, Share | Industry Growth [2033]",
        "{} Size, Share | Industry Forecast [2033]",
        "{} Size, Share | Growth Analysis [2033]",
        "{} Segment Analysis | Industry Growth [2033]",
        "{} Size, Share | Industry Report [2033]",
        "{} Size, Share | Forecast Report [2033]",
    ]
    selected_template = random.choice(meta_title_templates)
    return selected_template.format(market_name)

def clean_slug(text):
    """Clean and format text for use as a URL slug
    
    Replace special characters, commas, parentheses with hyphens or remove them.
    Convert to lowercase and remove consecutive hyphens.
    """
    # Replace parentheses with spaces (keeping content inside)
    text = text.replace('(', ' ').replace(')', ' ')
    # Replace commas with hyphens
    text = text.replace(',', '-')
    # Replace special characters (keep only alphanumeric, spaces, and hyphens)
    text = re.sub(r'[^a-zA-Z0-9\s-]', '-', text)
    # Replace spaces with hyphens
    text = text.replace(' ', '-')
    # Convert to lowercase
    text = text.lower()
    # Remove consecutive hyphens
    text = re.sub(r'-+', '-', text)
    # Remove leading/trailing hyphens
    text = text.strip('-')
    # Print the cleaned slug
    print(f"Report Slug: {text}")
    return text

def generate_docx_from_data(data, output_path=None, doc=None):
    market_name = data.get("title", "Aortic Endografts Market")
    
    # Build data_tuples from ToC
    segments_map = {"Type": [], "Procedure": [], "End User": []}
    current_chapter = None
    for line in data.get("table_of_contents", []):
        s = line.strip()
        if "Chapter 4" in s and "BY TYPE" in s.upper():
            current_chapter = "Type"
            continue
        if "Chapter 5" in s and "BY PROCEDURE" in s.upper():
            current_chapter = "Procedure"
            continue
        if "Chapter 6" in s and "BY END USER" in s.upper():
            current_chapter = "End User"
            continue
        if current_chapter and s and s[0].isdigit() and ". " in s:
            after_num = s.split(". ", 1)[1].strip()
            if after_num and not after_num.lower().startswith(("key ", "market ")):
                if after_num not in segments_map[current_chapter]:
                    segments_map[current_chapter].append(after_num)
    
    # Convert to data tuples
    data_tuples = []
    for seg_name, subitems in segments_map.items():
        data_tuples.append((f"Global {market_name.lower()} Market Size by {seg_name} (2026-2033)", 0))
        data_tuples.append(("Market Overview", 1))
        for it in subitems[:6]:
            data_tuples.append((it, 1))
    
    rd_meta = []
    
    # Add industry classification data FIRST (right after Image)
    industry_classification = data.get("industry_classification", {})
    if industry_classification:
        sector = industry_classification.get("sector", "Not specified")
        industry_group = industry_classification.get("industry_group", "Not specified")
        industry = industry_classification.get("industry", "Not specified")
        sub_industry = industry_classification.get("sub_industry", "Not specified")
        
        # Process rd_metadata_pairs and insert industry classification after Image
        for i, (k, v) in enumerate(rd_metadata_pairs):
            if k == "Image":
                # Add Image first
                if k == "Meta Title" and v == "RANDOM_META_TITLE":
                    formatted_value = get_random_meta_title(market_name)
                    rd_meta.append((k, formatted_value))
                elif '{}' in v:
                    formatted_value = v.format(market_name)
                    if k == "Report Slug":
                        formatted_value = clean_slug(formatted_value)
                    rd_meta.append((k, formatted_value))
                else:
                    rd_meta.append((k, v))
                
                # Add industry classification right after Image
                rd_meta.append(("Sector", sector))
                rd_meta.append(("Industry Group", industry_group))
                rd_meta.append(("Industry", industry))
                rd_meta.append(("Sub-Industry", sub_industry))
                break
        
        # Add remaining rd_metadata_pairs
        for k, v in rd_metadata_pairs[1:]:
            if k == "Meta Title" and v == "RANDOM_META_TITLE":
                formatted_value = get_random_meta_title(market_name)
                rd_meta.append((k, formatted_value))
            elif '{}' in v:
                formatted_value = v.format(market_name)
                if k == "Report Slug":
                    formatted_value = clean_slug(formatted_value)
                rd_meta.append((k, formatted_value))
            else:
                rd_meta.append((k, v))
    else:
        # If no industry classification, just process normally
        for k, v in rd_metadata_pairs:
            if k == "Meta Title" and v == "RANDOM_META_TITLE":
                # Special case for Meta Title - randomly select from templates
                formatted_value = get_random_meta_title(market_name)
                rd_meta.append((k, formatted_value))
            elif '{}' in v:
                formatted_value = v.format(market_name)
                # Special case for Report Slug - clean special characters, commas, and parentheses
                if k == "Report Slug":
                    formatted_value = clean_slug(formatted_value)
                rd_meta.append((k, formatted_value))
            else:
                rd_meta.append((k, v))

    # --- Market input overrides (unit, value_2024, value_2025, value_2033, CAGR) ---
    mi = (data.get("market_inputs") or {})
    
    # Extract CAGR
    try:
        user_cagr = float(mi.get("cagr")) if mi.get("cagr") is not None else None
    except (ValueError, TypeError):
        user_cagr = None
    
    # Extract currency/unit
    currency_override = (mi.get("unit") or None)
    
    # Extract value_2024
    try:
        user_value_2024 = float(mi.get("value_2024")) if mi.get("value_2024") is not None else None
    except (ValueError, TypeError):
        user_value_2024 = None
    
    # Extract value_2025
    try:
        user_value_2025 = float(mi.get("value_2025")) if mi.get("value_2025") is not None else None
    except (ValueError, TypeError):
        user_value_2025 = None
    
    # Extract value_2033
    try:
        user_value_2033 = float(mi.get("value_2033")) if mi.get("value_2033") is not None else None
    except (ValueError, TypeError):
        user_value_2033 = None

    # Choose values (fallback to constants when not provided)
    cagr_use = user_cagr if user_cagr is not None else cagr
    currency_use = currency_override if currency_override else currency

    # Prioritize user-provided value_2024
    if user_value_2024 is not None:
        value_2024_use = user_value_2024
    elif user_value_2025 is not None and cagr_use is not None:
        # Derive 2024 from 2025 and CAGR if 2024 not provided
        value_2024_use = round(user_value_2025 / (1 + cagr_use / 100), 2)
    else:
        value_2024_use = value_2024

    # Calculate/use value_2025
    if user_value_2025 is not None:
        value_2025_calc = user_value_2025
    else:
        value_2025_calc = round(value_2024_use * (1 + cagr_use / 100) ** 1, 2)

    # Calculate/use value_2033
    if user_value_2033 is not None:
        value_2033_calc = user_value_2033
    else:
        value_2033_calc = round(value_2024_use * (1 + cagr_use / 100) ** 9, 2)

    # Build Meta Description using provided pattern and selected values
    meta_desc = build_meta_description(
        market_name=market_name,
        value_2024=value_2024_use,
        value_2025=value_2025_calc,
        value_2033=value_2033_calc,
        currency=currency_use,
        cagr=cagr_use,
    )
    rd_meta = [(k, meta_desc if k == "Meta Description" else v) for k, v in rd_meta]
    
    # Print all metadata being saved (including industry classification)
    print("\n=== Metadata being saved to document ===")
    for key, value in rd_meta:
        if key in ["Sector", "Industry Group", "Industry", "Sub-Industry", "Report Slug"]:
            print(f"{key}: {value}")
    print("========================================\n")
    
    # Get companies from data, not the global variable
    companies_list = data.get("companies", [])
    
    return export_to_word_RD(
        data_dict=data,
        value_2024=value_2024_use,
        value_2025=value_2025_calc,
        value_2033=value_2033_calc,
        currency=currency_use,
        cagr=cagr_use,
        companies=companies_list,
        rd_meta=rd_meta,
        output_path=output_path,
        doc=doc
    )